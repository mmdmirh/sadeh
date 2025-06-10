import os
import io
import tempfile
import datetime
import json
import wave
import logging
import glob
import time
import threading

from flask import Flask, render_template, request, redirect, url_for, flash, Response, send_file, jsonify, session, stream_with_context
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, login_user, logout_user, current_user, login_required, UserMixin
from flask_migrate import Migrate
from flask_mail import Mail, Message
from werkzeug.security import generate_password_hash, check_password_hash
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
from flask_wtf.csrf import CSRFProtect
import subprocess
from ollama import RequestError, ResponseError

# Import our new LLM service abstraction
from llm_service import LLMServiceFactory

# Speech service functionality has been removed

# Add logging setup near the top of your file, after imports
import logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Define application pages/endpoints that can have access controlled
MANAGED_PAGE_ENDPOINTS = [
    {'endpoint': 'index', 'display_name': 'Home Page', 'description': 'The main landing page of the application.'},
    {'endpoint': 'chat', 'display_name': 'Chat Interface', 'description': 'The primary chat functionality.'},
    {'endpoint': 'admin_rbac_page', 'display_name': 'Access Control', 'description': 'This admin page for managing roles, users, and permissions.'},
    # {'endpoint': 'profile', 'display_name': 'User Profile', 'description': 'User profile viewing and editing page.'}, # Example for future
    # {'endpoint': 'settings', 'display_name': 'User Settings', 'description': 'User-specific application settings.'} # Example for future
]

# Load environment variables from .env file
load_dotenv(dotenv_path='deploy/.env')

# Ensure static directories exist
if not os.path.exists('static'):
    os.makedirs('static')
if not os.path.exists('static/css'): # Corrected line
    os.makedirs('static/css')
if not os.path.exists('static/js'):
    os.makedirs('static/js')
if not os.path.exists('static/uploads'):
    os.makedirs('static/uploads')

app = Flask(__name__, static_folder='static')
app.secret_key = os.environ.get('SECRET_KEY', 'you-will-never-guess')
csrf = CSRFProtect(app)

# Retrieve individual MySQL settings from .env
mysql_user = os.environ.get('MYSQL_USER')
mysql_password = os.environ.get('MYSQL_PASSWORD')
mysql_database = os.environ.get('MYSQL_DATABASE')
# Use the Docker service name 'mysql' or the env var if set
mysql_host = os.environ.get('MYSQL_HOST', 'mysql')
mysql_port = os.environ.get('MYSQL_PORT', '3306') # Internal Docker port

# Build the SQLAlchemy connection string dynamically.
logger.info(f"Attempting to connect to MySQL with host: '{mysql_host}' and port: '{mysql_port}'")
connection_str = f"mysql+pymysql://{mysql_user}:{mysql_password}@{mysql_host}:{mysql_port}/{mysql_database}"
app.config['SQLALCHEMY_DATABASE_URI'] = connection_str
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Initialize SQLAlchemy AFTER app is created and configured
db = SQLAlchemy(app)
migrate = Migrate(app, db, compare_type=True)

# Configure Flask-Mail (read SMTP settings from environment)
app.config['MAIL_SERVER'] = os.environ.get('MAIL_SERVER', 'smtp.example.com')
app.config['MAIL_PORT'] = int(os.environ.get('MAIL_PORT', 587))
app.config['MAIL_USE_TLS'] = os.environ.get('MAIL_USE_TLS', 'true').lower() in ['true', '1']
app.config['MAIL_USERNAME'] = os.environ.get('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.environ.get('MAIL_PASSWORD')
app.config['MAIL_DEFAULT_SENDER'] = os.environ.get('MAIL_DEFAULT_SENDER')
mail = Mail(app)

# Configure Flask-Login
login_manager = LoginManager(app)
login_manager.login_view = 'index'


# === Initialize LLM Service and Manage Models ===
llm_service_type = os.environ.get('LLM_SERVICE', 'ollama').lower()
logger.info(f"Initializing LLM service of type: {llm_service_type}")
llm_service = LLMServiceFactory.create_service() # This might call test_connection

# Store successfully managed models and the effective default model in app.config
app.config['MANAGED_OLLAMA_MODELS'] = []
app.config['EFFECTIVE_DEFAULT_MODEL_NAME'] = None
DEFAULT_MODEL_NAME = None # Will be determined by the logic below

if llm_service_type == 'ollama':
    if llm_service and hasattr(llm_service, 'test_connection') and llm_service.test_connection(): # Ensure service is responsive
        ollama_models_env = os.environ.get('OLLAMA_MODELS')
        if ollama_models_env:
            managed_model_names_from_env = [name.strip() for name in ollama_models_env.split(',') if name.strip()]
            logger.info(f"Target Ollama models from OLLAMA_MODELS env var: {managed_model_names_from_env}")

            current_ollama_server_models = llm_service.list_models() # Get models currently on server
            logger.info(f"Models currently on Ollama server: {current_ollama_server_models}")

            successfully_managed_models_list = []
            for model_name in managed_model_names_from_env:
                is_on_server = False
                for server_model in current_ollama_server_models:
                    if server_model == model_name or server_model.startswith(model_name + ":"):
                        is_on_server = True
                        break
                
                if is_on_server:
                    logger.info(f"Model '{model_name}' is already available on Ollama server.")
                    successfully_managed_models_list.append(model_name)
                else:
                    logger.info(f"Model '{model_name}' not found on Ollama server. Attempting to pull...")
                    if llm_service.pull_model(model_name): # pull_model returns True on success
                        logger.info(f"Successfully pulled model '{model_name}'.")
                        successfully_managed_models_list.append(model_name)
                    else:
                        logger.warning(f"Failed to pull model '{model_name}'. It will not be available through this managed list.")
            
            app.config['MANAGED_OLLAMA_MODELS'] = successfully_managed_models_list
            logger.info(f"Successfully managed Ollama models (available/pulled): {app.config['MANAGED_OLLAMA_MODELS']}")
        else:
            logger.warning("OLLAMA_MODELS environment variable not set. No specific Ollama models will be pre-managed/pulled.")
            # If OLLAMA_MODELS is not set, MANAGED_OLLAMA_MODELS remains empty.
            # initialize_rbac_data will then not add any models from this list.

        # Determine DEFAULT_MODEL_NAME for Ollama
        env_default_model_from_env = os.environ.get('DEFAULT_MODEL_NAME')
        effective_default = None

        if env_default_model_from_env:
            if env_default_model_from_env in app.config['MANAGED_OLLAMA_MODELS']:
                effective_default = env_default_model_from_env
            else:
                logger.warning(f"DEFAULT_MODEL_NAME '{env_default_model_from_env}' from .env is not in the list of successfully managed models ({app.config['MANAGED_OLLAMA_MODELS']}).")
                if app.config['MANAGED_OLLAMA_MODELS']:
                    effective_default = app.config['MANAGED_OLLAMA_MODELS'][0]
                    logger.warning(f"Falling back to the first managed model as default: '{effective_default}'.")
                else:
                    logger.error("No managed Ollama models available. Cannot set a default model from managed list.")
        elif app.config['MANAGED_OLLAMA_MODELS']:
            effective_default = app.config['MANAGED_OLLAMA_MODELS'][0]
            logger.info(f"DEFAULT_MODEL_NAME not set in .env. Using the first managed model as default: '{effective_default}'.")
        else:
            logger.warning("DEFAULT_MODEL_NAME not set in .env and no managed Ollama models available. Default model not set from managed list.")
        
        app.config['EFFECTIVE_DEFAULT_MODEL_NAME'] = effective_default
        DEFAULT_MODEL_NAME = effective_default 
    else:
        logger.error("Ollama service is not available or failed connection test. Cannot manage Ollama models or set default.")
        app.config['MANAGED_OLLAMA_MODELS'] = [] # Ensure it's empty
        app.config['EFFECTIVE_DEFAULT_MODEL_NAME'] = None
        DEFAULT_MODEL_NAME = None

elif llm_service_type == 'llamacpp':
    DEFAULT_MODEL_NAME = os.environ.get('LLAMACPP_MODEL', "llama-2-7b-chat.Q4_K_M.gguf")
    app.config['EFFECTIVE_DEFAULT_MODEL_NAME'] = DEFAULT_MODEL_NAME
    # LlamaCPP models are not currently managed via OLLAMA_MODELS env var
    app.config['MANAGED_OLLAMA_MODELS'] = [] # Ensure this is empty for non-ollama services
else:
    logger.error(f"Unsupported LLM_SERVICE_TYPE '{llm_service_type}'. No default model configured through this logic.")
    DEFAULT_MODEL_NAME = None
    app.config['EFFECTIVE_DEFAULT_MODEL_NAME'] = None
    app.config['MANAGED_OLLAMA_MODELS'] = []

if DEFAULT_MODEL_NAME:
    logger.info(f"Global DEFAULT_MODEL_NAME set to: {DEFAULT_MODEL_NAME}")
else:
    logger.warning("Global DEFAULT_MODEL_NAME could not be determined based on configuration and service availability.")

# Add a template filter for converting newlines to <br> tags
@app.template_filter('nl2br')
def nl2br(value):
    # First trim the string to remove leading/trailing whitespace
    if not value:
        return value
    value = value.strip()
    # Replace newlines with <br> tags
    value = value.replace('\n', '<br>')
    # Remove any double <br> tags that might cause excessive spacing
    while '<br><br><br>' in value:
        value = value.replace('<br><br><br>', '<br><br>')
    return value

# Print all registered routes for debugging (always runs)
print("\n=== Registered Flask Routes ===")
try:
    for rule in app.url_map.iter_rules():
        print(f"{rule.endpoint}: {rule}")
except Exception as e:
    print(f"Error printing routes: {e}")
print("==============================\n")

# ===========================
# Association Tables for RBAC
# ===========================
# Association table for User and Role (many-to-many)
user_roles = db.Table('user_roles',
    db.Column('user_id', db.Integer, db.ForeignKey('user.id'), primary_key=True),
    db.Column('role_id', db.Integer, db.ForeignKey('role.id'), primary_key=True)
)

# Association table for Role and Model (many-to-many)
role_models = db.Table('role_models',
    db.Column('role_id', db.Integer, db.ForeignKey('role.id'), primary_key=True),
    db.Column('model_id', db.Integer, db.ForeignKey('model.id'), primary_key=True)
)

# No explicit association table needed for PagePermission as it's a direct model with a ForeignKey to Role (one-to-many from Role's perspective)

# ===========================
# Database Models
# ===========================
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(64), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    firstname = db.Column(db.String(100), nullable=True)
    lastname = db.Column(db.String(100), nullable=True)
    # Increase length from 256 to 512 to accommodate modern hashes like scrypt
    password_hash = db.Column(db.String(512))
    confirmed = db.Column(db.Boolean, default=False)
    is_active = db.Column(db.Boolean, nullable=False, default=True)  # New field for user active status
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    conversations = db.relationship('Conversation', backref='user', lazy=True)

    # Relationship for roles (many-to-many)
    # Roles assigned to this user
    roles = db.relationship('Role', secondary=user_roles,
                            lazy='subquery', backref=db.backref('users_in_role', lazy=True))

    def can_access_page(self, page_endpoint):
        """Check if the user can access a specific page based on their roles."""
        if not self.is_active:
            return False
        # Admins have access to all pages by default (conventionally)
        if self.has_role('admin'):
            return True
        for role in self.roles:
            if role.has_page_access(page_endpoint):
                return True
        return False

    def has_role(self, role_name):
        """Check if the user has a specific role."""
        return any(role.name == role_name for role in self.roles)

    def can_access_model(self, model_id):
        """Check if the user can access a specific model based on their roles."""
        if self.has_role('admin'): # Admins can access all models
            return True
        for role in self.roles:
            if role.has_model_access(model_id):
                return True
        return False

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

class Conversation(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    title = db.Column(db.String(100), default="New Conversation")
    selected_model = db.Column(db.String(64))
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    document_mode = db.Column(db.Boolean, default=False)  # Add this line
    messages = db.relationship('ChatMessage', backref='conversation', cascade="all, delete-orphan", lazy=True)
    documents = db.relationship('Document', backref='conversation', cascade="all, delete-orphan", lazy=True)

class ChatMessage(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    conversation_id = db.Column(db.Integer, db.ForeignKey('conversation.id'), nullable=False)
    sender = db.Column(db.String(10))  # 'user' or 'ai'
    content = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)

class Document(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    conversation_id = db.Column(db.Integer, db.ForeignKey('conversation.id'), nullable=False)
    filename = db.Column(db.String(256))
    data = db.Column(db.LargeBinary)  # store file as BLOB
    mime_type = db.Column(db.String(128))
    uploaded_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)

# ===========================
# Flask-Login loader
# ===========================

class PagePermission(db.Model):
    __tablename__ = 'page_permission'
    id = db.Column(db.Integer, primary_key=True)
    role_id = db.Column(db.Integer, db.ForeignKey('role.id'), nullable=False)
    page_endpoint = db.Column(db.String(255), nullable=False) # e.g., 'chat', 'admin_rbac_page'

    # Unique constraint to prevent duplicate permissions for the same role and page
    __table_args__ = (db.UniqueConstraint('role_id', 'page_endpoint', name='_role_page_uc'),)

    def __repr__(self):
        return f"<PagePermission role_id={self.role_id} page='{self.page_endpoint}'>"


class Role(db.Model):
    __tablename__ = 'role'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(80), unique=True, nullable=False)
    description = db.Column(db.String(255), nullable=True)

    # Relationship to Model (many-to-many)
    # Models that this role has access to
    models = db.relationship('Model', secondary=role_models,
                             lazy='subquery', backref=db.backref('roles_having_access', lazy=True))

    # Relationship to PagePermission (one-to-many: one Role can have many PagePermissions)
    page_permissions = db.relationship('PagePermission', backref='role', lazy='dynamic', cascade="all, delete-orphan")

    def has_page_access(self, page_endpoint):
        """Check if this role has permission for a specific page endpoint."""
        return self.page_permissions.filter_by(page_endpoint=page_endpoint).first() is not None

    def has_model_access(self, model_id):
        """Check if this role has access to a specific model by its ID."""
        # self.models is the relationship to Model (many-to-many)
        if not hasattr(self, 'models') or not self.models:
            return False
        return any(model.id == model_id for model in self.models)

    def __repr__(self):
        return f'<Role {self.name}>'

class Model(db.Model):
    __tablename__ = 'model'
    id = db.Column(db.Integer, primary_key=True)
    # The exact name Ollama uses, e.g., "llama3:8b-instruct-q5_K_M"
    ollama_model_name = db.Column(db.String(128), unique=True, nullable=False, index=True)
    # A user-friendly name for UIs, e.g., "Llama 3 8B Instruct"
    display_name = db.Column(db.String(128), nullable=False)
    description = db.Column(db.Text, nullable=True)
    # Useful for grouping or identifying the family of the model, e.g., "llama3", "gemma"
    base_model_identifier = db.Column(db.String(128), nullable=True)
    # Flag to enable/disable in the app without deleting from DB or Ollama
    is_active = db.Column(db.Boolean, default=True, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    # 'roles_having_access' backref is created by Role.models relationship

    def __repr__(self):
        return f'<Model {self.display_name} ({self.ollama_model_name})>'


@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))

@app.route('/', methods=['GET', 'POST'])
def index():
    if current_user.is_authenticated:
        return redirect(url_for('chat'))

    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        user = User.query.filter_by(email=email).first()

        if not user or not user.check_password(password):
            flash("Invalid email or password.", 'danger')
            return redirect(url_for('index'))

        if not user.confirmed:
            flash("Please confirm your email before logging in.", 'warning')
            return redirect(url_for('index'))
        
        if not user.is_active:
            flash("Your account is inactive. Please contact your administrator.", 'danger')
            return redirect(url_for('index'))

        login_user(user)
        next_page = request.args.get('next')
        if next_page and next_page.startswith('/'):
            return redirect(next_page)
        else:
            return redirect(url_for('chat'))
            
    return render_template('index.html')

@app.route('/ping')
def ping():
    return 'pong'

# ===========================
# Ollama model listing at startup
# ===========================
def load_and_ensure_llm_models():
    """Loads model list, ensures default (and specified) models are pulled if missing."""
    with app.app_context(): # Ensure all operations run within app context
        # Use the centrally managed list of Ollama models from app.config
        managed_ollama_models = app.config.get('MANAGED_OLLAMA_MODELS', [])
        logger.info(f"LOAD_AND_ENSURE: Using MANAGED_OLLAMA_MODELS from app.config: {managed_ollama_models}")

        models_to_ensure = set(managed_ollama_models) # Start with all managed models
        
        # DEFAULT_MODEL_NAME should already be in managed_ollama_models if valid,
        # but adding it here ensures it's considered if somehow missed or if it's a LlamaCPP model not in OLLAMA_MODELS.
        if DEFAULT_MODEL_NAME: 
            models_to_ensure.add(DEFAULT_MODEL_NAME)
        
        logger.info(f"LOAD_AND_ENSURE: Final models_to_ensure (after adding default if needed): {list(models_to_ensure)}")

        available_models = []
        try:
            all_active_db_models = Model.query.filter_by(is_active=True).all()
            active_db_model_names = {model.ollama_model_name for model in all_active_db_models}
            logger.info(f"Active models from DB: {active_db_model_names}")

            if llm_service_type == 'ollama':
                try:
                    initial_server_models_list = llm_service.list_models()
                    initial_server_models_set = set(initial_server_models_list)
                    logger.info(f"LOAD_AND_ENSURE: Initial models on Ollama server: {initial_server_models_set}")

                    models_pulled_this_run = set()
                    # Ensure managed models are pulled if not on server
                    for model_name_to_pull in managed_ollama_models: # managed_ollama_models is from app.config
                        if model_name_to_pull not in initial_server_models_set:
                            try:
                                logger.info(f"LOAD_AND_ENSURE: Model '{model_name_to_pull}' (managed) not on server. Attempting to pull...")
                                llm_service.pull_model(model_name_to_pull)
                                logger.info(f"LOAD_AND_ENSURE: Successfully pulled model '{model_name_to_pull}'.")
                                models_pulled_this_run.add(model_name_to_pull)
                            except Exception as e:
                                logger.error(f"LOAD_AND_ENSURE: Failed to pull managed model '{model_name_to_pull}': {e}")
                    
                    final_server_models_set = initial_server_models_set.union(models_pulled_this_run)
                    logger.info(f"LOAD_AND_ENSURE: Final models on Ollama server (after pulls): {final_server_models_set}")

                    # Populate available_models from all models now on the server that are also active in DB
                    for server_model_name in final_server_models_set:
                        if server_model_name in active_db_model_names:
                            db_model_obj = next((m for m in all_active_db_models if m.ollama_model_name == server_model_name), None)
                            if db_model_obj: # Should be true if server_model_name in active_db_model_names
                                available_models.append({
                                    'id': db_model_obj.id,
                                    'name': db_model_obj.display_name or db_model_obj.ollama_model_name,
                                    'ollama_model_name': db_model_obj.ollama_model_name,
                                    'is_default': db_model_obj.ollama_model_name == DEFAULT_MODEL_NAME
                                })
                except RequestError as re:
                    logger.error(f"Ollama RequestError when ensuring models: {re}. This might happen if Ollama is not running or not reachable.")
                    # Fallback: only use models already in DB if Ollama is down and they were in models_to_ensure
                    for db_model in all_active_db_models:
                        if db_model.ollama_model_name in models_to_ensure:
                            available_models.append({
                                'id': db_model.id,
                                'name': db_model.display_name or db_model.ollama_model_name,
                                'ollama_model_name': db_model.ollama_model_name,
                                'is_default': db_model.ollama_model_name == DEFAULT_MODEL_NAME
                            })

            elif llm_service_type == 'llamacpp':
                if DEFAULT_MODEL_NAME in active_db_model_names:
                    db_model_obj = next((m for m in all_active_db_models if m.ollama_model_name == DEFAULT_MODEL_NAME), None)
                    if db_model_obj:
                        available_models.append({
                            'id': db_model_obj.id,
                            'name': db_model_obj.display_name, # Corrected to display_name
                            'ollama_model_name': db_model_obj.ollama_model_name,
                            'is_default': True
                        })
                else:
                    logger.warning(f"LlamaCPP model '{DEFAULT_MODEL_NAME}' is not marked active in the database.")

            # Consolidate fallback for empty available_models
            if not available_models and DEFAULT_MODEL_NAME and DEFAULT_MODEL_NAME in active_db_model_names:
                logger.info(f"No specific models made it to available_models list, but default '{DEFAULT_MODEL_NAME}' is active. Adding it.")
                db_model_obj = next((m for m in all_active_db_models if m.ollama_model_name == DEFAULT_MODEL_NAME), None)
                if db_model_obj:
                    available_models.append({
                        'id': db_model_obj.id,
                        'name': db_model_obj.display_name, # Corrected to display_name
                        'ollama_model_name': db_model_obj.ollama_model_name,
                        'is_default': True
                    })
            
            if DEFAULT_MODEL_NAME:
                available_models.sort(key=lambda x: x['ollama_model_name'] != DEFAULT_MODEL_NAME)

            return available_models

        except Exception as e:
            logger.exception(f"Error during model loading and pulling process (within app_context): {e}")
            # Fallback logic if an error occurs even within the app_context
            if DEFAULT_MODEL_NAME:
                # Check if DEFAULT_MODEL_NAME exists in the database as a last resort
                try:
                    default_db_model = Model.query.filter_by(ollama_model_name=DEFAULT_MODEL_NAME, is_active=True).first()
                    if default_db_model:
                        logger.warning(f"Proceeding with fallback default model for UI due to error: {default_db_model.ollama_model_name} (from DB)")
                        available_models.append({
                            'id': default_db_model.id,
                            'name': default_db_model.display_name,  
                            'ollama_model_name': default_db_model.ollama_model_name, 
                            'is_default': True,
                            'description': default_db_model.description
                        })
                except Exception as db_e:
                    logger.error(f"Could not even fetch default model from DB during fallback: {db_e}")
                
                logger.warning(f"Proceeding with fallback default model name (string only) for UI due to error: {DEFAULT_MODEL_NAME}")
                # Return a structure consistent with what the chat route expects if possible, even if it's just the name
                return [{'name': DEFAULT_MODEL_NAME, 'ollama_model_name': DEFAULT_MODEL_NAME, 'is_default': True, 'id': None}]
            return []

llm_models = load_and_ensure_llm_models()
app.config['LLM_MODELS'] = llm_models if llm_models else ([DEFAULT_MODEL_NAME] if DEFAULT_MODEL_NAME else [])
logger.info(f"Using models for dropdown: {app.config['LLM_MODELS']}")

# ===========================
# Audio Processing Placeholders (functionality removed)
# ===========================

def check_whisper_model_exists(model_name="base"):
    """Placeholder function - Whisper functionality has been removed"""
    logger.warning("Whisper functionality has been removed from this version")
    return False

def check_ffmpeg_installed():
    """Placeholder function - FFmpeg check"""
    logger.warning("FFmpeg check - audio processing functionality has been removed")
    return False

def recognize_audio(file_path, language=None):
    """Placeholder function - Audio recognition has been removed"""
    logger.warning("Audio recognition functionality has been removed")
    return "Audio recognition is not available in this version"

def convert_audio_format(input_path):
    """Placeholder function - Audio conversion has been removed"""
    logger.warning("Audio conversion functionality has been removed")
    return None

def detect_language(audio_file_path):
    """Placeholder function - Language detection has been removed"""
    logger.warning("Language detection functionality has been removed")
    return 'en'  # Default to English

# ===========================
# Routes for Authentication
# ===========================
    
@app.route('/confirm/<token>')
def confirm_email(token):
    try:
        user_id = int(token.split('-')[0])
    except Exception:
        flash("Invalid confirmation token.")
        return redirect(url_for('index'))
    user = db.session.get(User, user_id)
    if user:
        user.confirmed = True
        db.session.commit()
        flash("Your account has been confirmed!")
    return redirect(url_for('index'))



@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('You have been logged out.', 'info')
    return redirect(url_for('index'))
    
@app.route('/reset', methods=['GET', 'POST'])
def reset_request():
    if request.method == 'POST':
        email = request.form['email']
        user = User.query.filter_by(email=email).first()
        if user:
            token = f"{user.id}-reset-token"  # Replace with secure token generation.
            reset_url = url_for('reset_password', token=token, _external=True)
            msg = Message("Password Reset", recipients=[email])
            msg.body = f"Reset your password by clicking on the link: {reset_url}"
            mail.send(msg)
            flash("Password reset email sent.")
        else:
            flash("Email not found.")
    return render_template('reset_password_coming_soon.html')

@app.route('/reset_password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    try:
        user_id = int(token.split('-')[0])
    except Exception:
        flash("Invalid reset token.")
        return redirect(url_for('reset_request'))
    user = db.session.get(User, user_id)
    if request.method == 'POST':
        new_password = request.form['password']
        user.set_password(new_password)
        db.session.commit()
        flash("Your password has been updated.")
        return redirect(url_for('login'))
    return render_template('reset_password.html')

# ===========================
# Routes for Chat and Conversations
# ===========================
@app.route('/chat', methods=['GET', 'POST'])
@login_required
def chat():
    if not current_user.can_access_page('chat'):
        flash('You do not have permission to access this page.', 'danger')
        return redirect(url_for('index'))
    # Check if user has a preferred LLM service
    user_service = session.get('user_llm_service')
    global llm_service
    global llm_service_type
    
    # If user has a preference that's different from current service, switch
    if user_service and user_service != llm_service_type:
        try:
            logger.info(f"Switching to user's preferred LLM service: {user_service}")
            
            # Try to create the service and test it
            temp_service = LLMServiceFactory.create_service_by_type(user_service)
            available_models = temp_service.list_models()
            
            # If we get here, the service is working
            llm_service = temp_service
            llm_service_type = user_service
            
            # Update the application's cached models list
            app.config['LLM_MODELS'] = available_models if available_models else [DEFAULT_MODEL_NAME]
            logger.info(f"Switched to user's preferred service {user_service} with models: {app.config['LLM_MODELS']}")
        except Exception as e:
            logger.exception(f"Failed to switch to user's preferred service {user_service}: {e}")
            # If we can't use the preferred service, clear the preference
            session.pop('user_llm_service', None)
    
    # Continue with the existing chat route logic
    conversation_id = request.args.get('conversation_id', None)
    llm_models_config = app.config.get('LLM_MODELS', [])
    logger.debug(f"Chat Route: Initial llm_models_config from app.config: {llm_models_config}")
    processed_ollama_names = []

    if isinstance(llm_models_config, list):
        for item in llm_models_config:
            if isinstance(item, str):
                processed_ollama_names.append(item)
            elif isinstance(item, dict):
                model_name = item.get('ollama_model_name')
                if model_name and isinstance(model_name, str):
                    processed_ollama_names.append(model_name)
                else:
                    logger.warning(f"CHAT_ROUTE_PROCESSING: Skipping malformed dict item in LLM_MODELS config: {item}")
            else:
                logger.warning(f"CHAT_ROUTE_PROCESSING: Skipping item of unexpected type in LLM_MODELS config: {item}")
    else:
        logger.warning(f"CHAT_ROUTE_PROCESSING: LLM_MODELS config is not a list: {llm_models_config}")

    raw_available_models_names = processed_ollama_names # This is the variable your existing logging and query use
    logger.debug(f"Chat Route: Extracted raw_available_models_names (for DB query): {raw_available_models_names}")

    # Filter these models based on database entries and user permissions
    accessible_models_for_template = []
    if current_user.is_authenticated and raw_available_models_names:
        # Fetch Model objects from DB that are active and match the names from config
        # We match based on ollama_model_name as that's what app.config['LLM_MODELS'] seems to store
        logger.info(f"CHAT_ROUTE_DEBUG: Type of raw_available_models_names: {type(raw_available_models_names)}")
        if isinstance(raw_available_models_names, list) and raw_available_models_names:
            logger.info(f"CHAT_ROUTE_DEBUG: Type of first element in raw_available_models_names: {type(raw_available_models_names[0])}")
            logger.info(f"CHAT_ROUTE_DEBUG: Content of raw_available_models_names: {raw_available_models_names}")
        elif isinstance(raw_available_models_names, list):
            logger.info(f"CHAT_ROUTE_DEBUG: raw_available_models_names is an empty list.")
        else:
            logger.info(f"CHAT_ROUTE_DEBUG: raw_available_models_names is not a list. Content: {raw_available_models_names}")

        db_models_to_check = Model.query.filter(
            Model.is_active==True,
            Model.ollama_model_name.in_(raw_available_models_names)
        ).all()
        logger.debug(f"Chat Route: DB models to check (active and on server): {[m.ollama_model_name for m in db_models_to_check]}")

        for model_obj in db_models_to_check:
            if current_user.can_access_model(model_obj.id):
                accessible_models_for_template.append({
                    'id': model_obj.id, 
                    'name': model_obj.display_name, 
                    'ollama_model_name': model_obj.ollama_model_name, 
                    'description': model_obj.description
                })
        logger.debug(f"Chat Route: Accessible models for template (after permission check, before fallback): {accessible_models_for_template}")
    
    # Ensure there's always at least one model if possible, or an empty list
    # The 'name' field in this list is what the user sees in the dropdown.
    # The 'ollama_model_name' is what's sent to the backend/LLM service.
    if not accessible_models_for_template and DEFAULT_MODEL_NAME:
        # Fallback to default model if user has access to it, or if no specific permissions are set for it (implicitly allowed)
        # This part might need refinement based on how default model access is handled for users without explicit permissions
        default_db_model = Model.query.filter_by(ollama_model_name=DEFAULT_MODEL_NAME, is_active=True).first()
        if default_db_model and current_user.can_access_model(default_db_model.id):
            available_models_for_template = [{
                'id': default_db_model.id, 
                'name': default_db_model.display_name, 
                'ollama_model_name': default_db_model.ollama_model_name, 
                'description': default_db_model.description
            }]
            logger.info(f"User {current_user.id} has no specific models, falling back to accessible default: {DEFAULT_MODEL_NAME}")
        else:
            available_models_for_template = [] # No accessible models
            logger.warning(f"User {current_user.id} has no accessible models, including the default. Model selection will be empty.")
    elif not accessible_models_for_template:
        available_models_for_template = []
        logger.warning(f"User {current_user.id} has no accessible models. Model selection will be empty.")
    else:
        available_models_for_template = accessible_models_for_template

    # This 'available_models' will be passed to the template
    # It's a list of dicts, e.g., [{'id':1, 'name':'GPT-4', 'ollama_model_name':'gpt-4'}, ...]
    available_models = available_models_for_template 

    # Create a set of ollama_model_names that are currently available and accessible to the user
    accessible_ollama_model_names = {model_dict['ollama_model_name'] for model_dict in available_models if 'ollama_model_name' in model_dict}
    logger.debug(f"Chat Route: Accessible ollama_model_names (set, after fallback logic): {accessible_ollama_model_names}")
    logger.debug(f"Chat Route: Configured EFFECTIVE_DEFAULT_MODEL_NAME: {app.config.get('EFFECTIVE_DEFAULT_MODEL_NAME')}")

    if conversation_id:
        conversation = Conversation.query.filter_by(id=conversation_id, user_id=current_user.id).first_or_404()
        
        model_needs_fallback = False
        if not conversation.selected_model: # No model selected yet for this existing conversation
            model_needs_fallback = True
            logger.info(f"Conversation {conversation.id} has no model selected.")
        elif conversation.selected_model not in accessible_ollama_model_names: # Selected model is no longer available/accessible
            model_needs_fallback = True
            logger.warning(f"Conversation {conversation.id} had model '{conversation.selected_model}' which is not available or accessible.")

        if model_needs_fallback:
            if available_models: # Check if there are any models to fall back to
                fallback_model_dict = available_models[0]
                new_model_name = fallback_model_dict.get('ollama_model_name')
                conversation.selected_model = new_model_name
                logger.info(f"Falling back/setting model for conversation {conversation.id} to '{new_model_name}'.")
                db.session.add(conversation) # Mark for update
            else: # No models available to fall back to
                conversation.selected_model = None
                logger.warning(f"No accessible models to fall back to for conversation {conversation.id}. Setting selected_model to None.")
                db.session.add(conversation) # Mark for update
    else: # No conversation_id, so it's a new session or fetching the latest conversation
        conversation = Conversation.query.filter_by(user_id=current_user.id).order_by(Conversation.created_at.desc()).first()
        if not conversation: # No existing conversations, create a new one
            initial_model_name = None
            if available_models:
                default_conv_model_dict = available_models[0]
                initial_model_name = default_conv_model_dict.get('ollama_model_name')
                logger.info(f"Creating first conversation for user {current_user.id} with model '{initial_model_name}'.")
            else: # No models available for a new conversation
                logger.warning(f"Cannot create new conversation for user {current_user.id} as no models are accessible. Initial model will be None.")
            
            conversation = Conversation(
                user_id=current_user.id,
                title="New Conversation",
                selected_model=initial_model_name # Assign string or None
            )
            db.session.add(conversation)
        elif conversation: # Existing conversation fetched as latest
            model_needs_fallback_for_latest = False
            if not conversation.selected_model:
                model_needs_fallback_for_latest = True
                logger.info(f"Latest conversation {conversation.id} has no model selected.")
            elif conversation.selected_model not in accessible_ollama_model_names:
                model_needs_fallback_for_latest = True
                logger.warning(f"Latest conversation {conversation.id} had model '{conversation.selected_model}' which is not available or accessible.")

            if model_needs_fallback_for_latest:
                if available_models:
                    fallback_model_dict = available_models[0]
                    new_model_name = fallback_model_dict.get('ollama_model_name')
                    conversation.selected_model = new_model_name
                    logger.info(f"Falling back/setting model for latest conversation {conversation.id} to '{new_model_name}'.")
                    db.session.add(conversation)
                else:
                    conversation.selected_model = None
                    logger.warning(f"No accessible models to fall back to for latest conversation {conversation.id}. Setting selected_model to None.")
                    db.session.add(conversation)


    all_conversations = Conversation.query.filter_by(user_id=current_user.id).order_by(Conversation.created_at.desc()).all()
    messages = ChatMessage.query.filter_by(conversation_id=conversation.id).order_by(ChatMessage.created_at).all()

    logger.debug(f"Chat Route: Final 'available_models' for template: {available_models}")
    logger.debug(f"Chat Route: Final 'conversation.selected_model' for UI: {conversation.selected_model if conversation else 'No conversation object'}")
    
    # Pass the LLM service type to the template
    return render_template('chat.html', 
                          conversation=conversation, 
                          all_conversations=all_conversations,
                          messages=messages, 
                          models=available_models, # This is now the filtered list of model dicts
                          llm_service_type=llm_service_type)

@app.route('/conversation/new', methods=['POST'])
@login_required
def new_conversation():
    llm_models_config = app.config.get('LLM_MODELS', [])
    processed_ollama_names = []

    if isinstance(llm_models_config, list):
        for item in llm_models_config:
            if isinstance(item, str):
                processed_ollama_names.append(item)
            elif isinstance(item, dict):
                model_name = item.get('ollama_model_name')
                if model_name and isinstance(model_name, str):
                    processed_ollama_names.append(model_name)
                else:
                    logger.warning(f"NEW_CONV_PROCESSING: Skipping malformed dict item in LLM_MODELS config: {item}")
            else:
                logger.warning(f"NEW_CONV_PROCESSING: Skipping item of unexpected type in LLM_MODELS config: {item}")
    else:
        logger.warning(f"NEW_CONV_PROCESSING: LLM_MODELS config is not a list: {llm_models_config}")

    if not processed_ollama_names:
        logger.error("NEW_CONV: No models available (processed_ollama_names is empty). Cannot create new conversation.")
        flash("Cannot start a new chat: No AI models are currently available or configured correctly.", "danger")
        return redirect(url_for('chat'))

    # Get model name string from form
    form_model_name = request.form.get('model')
    final_selected_model_name = None

    if form_model_name and form_model_name in processed_ollama_names:
        final_selected_model_name = form_model_name
    else:
        if form_model_name: # Model from form was provided but invalid or not in the processed list
            logger.warning(f"NEW_CONV: Model '{form_model_name}' from form is not in available list {processed_ollama_names}. Falling back.")
        else: # No model provided in form, also fall back
            logger.info(f"NEW_CONV: No model provided in form. Falling back.")
        
        # Fallback to the first model in the processed (and validated) list
        final_selected_model_name = processed_ollama_names[0]
        logger.info(f"NEW_CONV: Using fallback model '{final_selected_model_name}'.")

    logger.info(f"Creating new conversation for user {current_user.id} with model '{final_selected_model_name}'")
    conversation = Conversation(
        user_id=current_user.id,
        title="New Conversation", # Title can be set later based on first message
        selected_model=final_selected_model_name # This is now guaranteed to be an ollama_model_name string
    )
    db.session.add(conversation)
    db.session.commit()
    return redirect(url_for('chat', conversation_id=conversation.id))

@app.route('/conversation/<int:conversation_id>/rename', methods=['POST'])
@login_required
def rename_conversation(conversation_id):
    conversation = db.session.get(Conversation, conversation_id)
    # Check ownership
    if conversation.user_id != current_user.id:
        flash("Unauthorized access")
        return redirect(url_for('chat'))
    new_title = request.form.get('title', 'Untitled Conversation')
    conversation.title = new_title
    db.session.commit()
    
    flash("Conversation renamed successfully")
    return redirect(url_for('chat', conversation_id=conversation_id))

@app.route('/conversation/<int:conversation_id>/delete', methods=['POST'])
@login_required
def delete_conversation(conversation_id):
    conversation = db.session.get(Conversation, conversation_id)
    # Check ownership
    if conversation.user_id != current_user.id:
        flash("Unauthorized access")
        return redirect(url_for('chat'))
    else:
        db.session.delete(conversation)
        db.session.commit()
        flash("Conversation deleted")
    return redirect(url_for('chat'))

@app.route('/switch_model', methods=['POST'])
@login_required
def switch_model():
    conversation_id = request.form['conversation_id']
    new_model = request.form['model']
    conversation = db.session.get(Conversation, conversation_id)
    if conversation and conversation.user_id == current_user.id:
        conversation.selected_model = new_model
        db.session.commit()
        flash("Model switched successfully.")
    return redirect(url_for('chat'))

@app.route('/edit_message/<int:message_id>', methods=['POST'])
@login_required
def edit_message(message_id):
    new_content = request.form['content']
    message = db.session.get(ChatMessage, message_id)
    if message and message.conversation.user_id == current_user.id:
        message.content = new_content
        db.session.commit()
        flash("Message updated.")
    return redirect(url_for('chat'))

# ===========================
# File and Voice Upload Routes
# ===========================
def extract_text_from_document(document):
    """Extract text content from a document based on its MIME type"""
    try:
        mime_type = document.mime_type
        
        # For plain text files
        if (mime_type == 'text/plain'):
            return document.data.decode('utf-8')
        # For PDF files
        elif mime_type == 'application/pdf':
            try:
                import io
                from pdfminer.high_level import extract_text
                pdf_file = io.BytesIO(document.data)
                text = extract_text(pdf_file)
                return text
            except ImportError:
                logger.warning("pdfminer.six not installed. Cannot extract PDF text.")
                return "PDF text extraction requires pdfminer.six. Please install it with: pip install pdfminer.six"
                
        # For docx files
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            try:
                import io
                import docx
                docx_file = io.BytesIO(document.data)
                doc = docx.Document(docx_file)
                return "\n".join([para.text for para in doc.paragraphs])
            except ImportError:
                logger.warning("python-docx not installed. Cannot extract DOCX text.")
                return "DOCX text extraction requires python-docx. Please install it with: pip install python-docx"
        
        # For other formats, return a message
        else:
            return f"Content extraction not supported for {mime_type}. Using filename only."
    except Exception as e:
        logger.exception(f"Error extracting text from document: {str(e)}")
        return f"Error extracting text: {str(e)}"

@app.route('/upload_document', methods=['POST'])
@login_required
def upload_document():
    file = request.files['file']
    conversation_id = request.form['conversation_id']
    if file:
        try:
            # Get the conversation
            conversation = db.session.get(Conversation, conversation_id)
            if not conversation or conversation.user_id != current_user.id:
                flash("Unauthorized access")
                return redirect(url_for('chat'))
            
            # Save document to database
            doc = Document(
                conversation_id=conversation_id,
                filename=file.filename,
                data=file.read(),
                mime_type=file.mimetype
            )
            
            # Enable document mode for this conversation
            conversation.document_mode = True
            
            db.session.add(doc)
            # Commit here to ensure doc and conversation mode are saved before system message
            db.session.commit() 
            
            # Extract document text for context
            # file.seek(0)  # Reset file pointer
            # doc_text = extract_text_from_document(doc)
            
            # Create a system message to indicate document context mode
            system_message = ChatMessage(
                conversation_id=conversation_id,
                sender='ai',
                content=f"📄 Document '{file.filename}' has been uploaded. My responses will now be based only on knowledge from this document."
            )
            db.session.add(system_message)
            db.session.commit()
            
            flash("Document uploaded successfully. AI will now respond based on document content.")
        except Exception as e:
            logger.exception(f"Error uploading document: {str(e)}")
            flash(f"Error uploading document: {str(e)}")
            db.session.rollback()
    return redirect(url_for('chat', conversation_id=conversation_id))

@app.route('/upload_voice', methods=['POST'])
@login_required 
def upload_voice():
    if not check_whisper_model_exists():
        return jsonify({'success': False, 'error': 'Voice transcription service not available.'}), 500
        
    if 'voice' not in request.files:
        return jsonify({'success': False, 'error': 'No voice file part'}), 400
        
    file = request.files['voice']
    conversation_id = request.form.get('conversation_id')
    language = request.form.get('language', 'english')  # Default to English if not provided

    if file.filename == '':
        return jsonify({'success': False, 'error': 'No selected voice file'}), 400

    if not conversation_id:
        return jsonify({'success': False, 'error': 'Missing conversation ID'}), 400

    conversation = db.session.get(Conversation, conversation_id)
    if not conversation or conversation.user_id != current_user.id:
        return jsonify({'success': False, 'error': 'Conversation not found or unauthorized'}), 404

    # Use a temporary file for processing
    temp_audio_path = None
    transcription_result = None
    detected_language = None
    error_message = None

    try:
        # Save blob to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".webm") as temp_audio:
            file.save(temp_audio.name)
            temp_audio_path = temp_audio.name
            logger.info(f"Saved temporary voice file to {temp_audio_path}")

        # Transcribe using Whisper
        logger.info(f"Transcribing {temp_audio_path} with language hint: {language}")
        # Let Whisper detect language unless a specific one is strongly needed
        whisper_options = {"language": language if language in ['english', 'persian'] else None} 
        transcription_result = speech_service.transcribe_audio(temp_audio_path, **whisper_options)
        transcribed_text = transcription_result['text'].strip()
        detected_language = transcription_result.get('language', language)  # Use detected or fallback to hint
        logger.info(f"Transcription successful. Detected language: {detected_language}. Text: {transcribed_text}")

        if not transcribed_text:
            raise ValueError("Transcription resulted in empty text.")

        # --- Save User Message (Transcription) ---
        # Prefix with an indicator that it came from voice
        user_message_content = f"🎤: {transcribed_text}"
        user_message = ChatMessage(
            conversation_id=conversation.id,
            sender='user',
            content=user_message_content,
        )
        logger.info(f"Saved user transcription message with ID: {user_message.id}")

        # --- Get AI Response ---
        history = ChatMessage.query.filter_by(conversation_id=conversation.id).order_by(ChatMessage.created_at).all()
        formatted_history = [{"role": msg.sender, "content": msg.content} for msg in history]
        
        # Use the transcribed text as the latest user prompt
        # No need to include the "🎤: " prefix for the AI model context
        latest_prompt = transcribed_text 

        try:
            logger.info(f"Sending prompt to {llm_service_type} model {conversation.selected_model}: {latest_prompt}")
            # Use the new llm_service abstraction
            response = llm_service.chat(
                model=conversation.selected_model,
                messages=formatted_history
            )
            ai_response_text = response['message']['content']
            logger.info(f"Received AI response: {ai_response_text}")

            # --- Save AI Message ---
            ai_message = ChatMessage(
                conversation_id=conversation.id,
                sender='ai',
                content=ai_response_text,
            )
            db.session.add(ai_message)
            db.session.commit()
            logger.info(f"Saved AI response message with ID: {ai_message.id}")

            # --- Prepare JSON Response ---
            return jsonify({
                'success': True,
                'transcription': user_message_content, 
                'message_id': user_message.id,  # ID of the saved user message
                'ai_response': ai_response_text,
                'detected_language': detected_language,
            })

        except Exception as e:
            logger.error(f"Error getting AI response: {e}")
            error_message = f"Error getting AI response: {e}"
            # Still return success=True because transcription worked, but include error for AI part
            return jsonify({
                'success': True,  # Transcription succeeded
                'transcription': user_message_content,
                'message_id': user_message.id,
                'ai_response': None,  # Indicate AI response failed
                'error': error_message,  # Provide error detail
                'detected_language': detected_language,
            })

    except Exception as e:
        logger.error(f"Error processing voice file: {e}")
        error_message = f"Error processing voice: {e}"
        # Return success=False as the core voice processing failed
        return jsonify({'success': False, 'error': error_message, 'transcription': error_message}), 500
    finally:
        # Clean up temporary file
        if temp_audio_path and os.path.exists(temp_audio_path):
            try:
                os.remove(temp_audio_path)
                logger.info(f"Removed temporary voice file: {temp_audio_path}")
            except Exception as e:
                logger.error(f"Error removing temporary file {temp_audio_path}: {e}")

# Add a new function to call the AI model directly from backend
def call_ai_model(model_name, prompt):
    """Call the AI model synchronously and return the full response using the configured LLM service"""
    # Force use of the fixed model for voice assistant
    fixed_model = os.environ.get('DEFAULT_VOICE_MODEL', "llama2") # Consider making this configurable

    if not isinstance(prompt, str):
        logger.error(f"call_ai_model received non-string prompt: {type(prompt)}")
        raise TypeError("Prompt must be a string")

    logger.info(f"Calling {llm_service_type} model {fixed_model} with prompt: {prompt[:50]}...")

    # Detect if the prompt contains Persian text
    is_persian = any('\u0600' <= c <= '\u06FF' for c in prompt)

    # Add language instruction for Persian
    if is_persian and "پاسخ به زبان فارسی" not in prompt:
        prompt = "لطفا به سوال زیر به زبان فارسی پاسخ دهید:\n\n" + prompt
        logger.info("Added Persian language instruction to prompt")

    try:
        logger.info(f"Sending prompt to {llm_service_type} model {fixed_model}: {prompt}")
        # Use the new llm_service abstraction
        response = llm_service.chat(
            model=fixed_model,
            messages=[{'role': 'user', 'content': prompt}]
        )
        ai_response_text = response['message']['content']
        logger.info(f"Received AI response: {ai_response_text[:50]}...")
        return ai_response_text
    except Exception as e:
        logger.exception(f"Error calling {llm_service_type} API: {e}")
        raise Exception(f"Failed to get response from AI model via API: {e}")

# Add a route to get the voice recording
@app.route('/voice_recording/<int:recording_id>', methods=['GET'])
@login_required 
def get_voice_recording(recording_id):
    """Return the voice recording audio file"""
    temp_file_path = None
    try:
        # Get the document
        voice_doc = db.session.get(Document, recording_id)
        
        # Check if the voice belongs to a conversation owned by the current user
        conversation = db.session.get(Conversation, voice_doc.conversation_id)
        if not conversation or conversation.user_id != current_user.id:
            return "Unauthorized", 403
        
        # Create a temporary file to serve
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
            tmp.write(voice_doc.data)
            temp_file_path = tmp.name
        
        # Send the file without using the unsupported after_request parameter
        return send_file(
            temp_file_path,
            mimetype=voice_doc.mime_type,
            as_attachment=False,
            download_name=voice_doc.filename
        )
    except Exception as e:
        logger.exception(f"Error retrieving voice recording: {str(e)}")
        return f"Error retrieving voice recording: {str(e)}", 500
    finally:
        # Clean up the temporary file in a background thread to ensure
        # it happens after the file is served
        if temp_file_path and os.path.exists(temp_file_path):
            def cleanup_temp_file():
                try:
                    # Add a small delay to ensure file serving completes
                    time.sleep(1)
                    if os.path.exists(temp_file_path):
                        os.remove(temp_file_path)
                except Exception as e:
                    logger.warning(f"Failed to remove temporary file {temp_file_path}: {e}")
            
            import threading
            cleanup_thread = threading.Thread(target=cleanup_temp_file)
            cleanup_thread.daemon = True
            cleanup_thread.start()

# Update voice help route if needed, or remove if template is removed
@app.route('/voice_help')
def voice_help():
    """Provides detailed instructions for setting up voice recognition"""
    # Update this template or remove the route if the template is removed
    return render_template('voice_help.html') # Make sure this template exists and is updated

# ===========================
# Endpoint to Call the AI Model and Stream Response
# ===========================
def stream_llm_response(model_name, messages_history):
    """Streams response from the configured LLM service."""
    logger.info(f"--> Entering stream_llm_response for model: {model_name}")
    logger.info(f"--> Messages history count: {len(messages_history)}")
    logger.info(f"--> First message: {str(messages_history[0])[:100]}..." if messages_history else "No messages in history")
    
    stream_generator = None
    sent_any_chunk = False
    try:
        logger.info("--> Calling llm_service.stream_chat...")
        # Use the new llm_service abstraction for streaming
        stream_generator = llm_service.stream_chat(model_name, messages_history)
        logger.info("--> Got generator object from llm_service.stream_chat.")
        logger.info("--> Preparing to yield from stream_generator...")
        logger.info("--> Starting the generator iteration loop")
        try:
            for chunk in stream_generator:
                logger.info(f"Yielding chunk from llm_service.stream_chat: {str(chunk)[:60]}")
                # Wrap every chunk in SSE format (data: ...\n\n)
                try:
                    if isinstance(chunk, str) and chunk.startswith('data: '):
                        yield chunk if chunk.endswith('\n\n') else chunk + '\n\n'
                    else:
                        if not isinstance(chunk, str):
                            chunk = json.dumps(chunk)
                        yield f"data: {chunk}\n\n"
                    sent_any_chunk = True
                except Exception as e:
                    logger.exception(f"Error formatting chunk for SSE: {e}")
                    error_text = f"Error formatting chunk: {e}"
                    escaped_text = json.dumps({"error": error_text, "text": f"⚠️ {error_text}"})
                    yield f"data: {escaped_text}\n\n"
                    sent_any_chunk = True
            logger.info("Exited stream_generator loop in stream_llm_response.")
        except Exception as e:
            logger.exception(f"Exception while iterating stream_generator: {e}")
            error_text = json.dumps({"error": str(e), "text": f"⚠️ {str(e)}"})
            yield f"data: {error_text}\n\n"
            sent_any_chunk = True
    except Exception as e:
        logger.exception(f"--> Error during llm_service.stream_chat call or yield from: {e}")
        error_text = f"Error during {llm_service_type} stream: {e}"
        escaped_text = json.dumps({"error": error_text, "text": f"⚠️ {error_text}"})
        yield f"data: {escaped_text}\n\n"
    finally:
        logger.info(f"--> Exiting stream_llm_response for model: {model_name}")
        if not sent_any_chunk:
            logger.warning("No chunks were yielded from llm_service.stream_chat; sending empty response message.")
            empty_text = json.dumps({"error": "No response from model.", "text": "⚠️ No response from model."})
            yield f"data: {empty_text}\n\n"

@app.route('/call_model', methods=['POST'])
@login_required
def call_model():
    logger.info(f"Received /call_model request from user {current_user.id}")
    conversation_id = request.form['conversation_id']
    prompt = request.form['prompt']
    logger.info(f"Request details - conversation_id: {conversation_id}, prompt: {prompt[:50]}...")
    
    # Get conversation
    conversation = db.session.get(Conversation, conversation_id)
    if not conversation or conversation.user_id != current_user.id:
        logger.warning(f"Unauthorized access attempt to conversation {conversation_id}")
        return jsonify({"error": "Unauthorized"}), 403
    
    # Prepare message history
    messages_history = []
    for msg in conversation.messages:
        messages_history.append({
            'role': 'user' if msg.sender == 'user' else 'assistant',
            'content': msg.content
        })
    messages_history.append({'role': 'user', 'content': prompt})
    
    model_name = conversation.selected_model if hasattr(conversation, 'selected_model') and conversation.selected_model else DEFAULT_MODEL_NAME
    logger.info(f"Using model: {model_name}")
    
    # Save the user message to the DB
    user_message = ChatMessage(
        conversation_id=conversation_id,
        sender='user',
        content=prompt
    )
    db.session.add(user_message)
    db.session.commit()
    logger.info(f"Saved user message with ID {user_message.id}")
    
    def response_wrapper():
        logger.info("Starting response_wrapper generator")
        full_response = ""
        ai_message_id = None
        user_id = current_user.id if hasattr(current_user, 'id') and current_user.id else 0
        conv_id = conversation_id
        generator_key = f"user_{user_id}_conv_{conv_id}"
        active_response_generators[generator_key] = False
        try:
            logger.info("Preparing message history for LLM API")
            logger.info(f"Streaming from ollama model {model_name}")
            logger.info(f"Attempting ollama stream with model: {model_name}")
            chunk_count = 0
            try:
                logger.info("--> Entering stream_llm_response from response_wrapper...")
                yielded_any = False
                for chunk in stream_llm_response(model_name, messages_history):
                    logger.info(f"Yielding chunk #{chunk_count+1}: {str(chunk)[:60]}")
                    # Accumulate bot response text from each chunk
                    # Each chunk should be like: 'data: {"text": "..."}\n\n'
                    if chunk.startswith('data: '):
                        try:
                            data = json.loads(chunk[6:].strip())
                            if 'text' in data:
                                full_response += data['text']
                        except Exception as e:
                            logger.warning(f"Failed to parse streamed chunk for accumulation: {e}")
                    yield chunk
                    yielded_any = True
                    chunk_count += 1
                logger.info(f"Exited streaming loop after {chunk_count} chunks.")
                if not yielded_any:
                    logger.warning("No chunks were yielded from stream_llm_response; yielding fallback error chunk.")
                    error_text = json.dumps({"error": "No response from model.", "text": "⚠️ No response from model."})
                    yield f"data: {error_text}\n\n"
                # After streaming is done, save the bot message to DB
                if full_response.strip():
                    try:
                        ai_message = ChatMessage(
                            conversation_id=conversation_id,
                            sender='ai',
                            content=full_response
                        )
                        db.session.add(ai_message)
                        db.session.commit()
                        logger.info(f"Saved AI message with ID {ai_message.id}")
                        ai_message_id = ai_message.id
                    except Exception as e:
                        logger.error(f"Failed to save AI message to DB: {e}")
                        db.session.rollback()
                # Extract conversation topic after saving AI message
                try:
                    topic_prompt = "Please provide a concise conversation topic (single word or phrase, max 6 characters) for the conversation above."
                    # Build topic messages with full conversation context and AI response
                    topic_messages = messages_history + [
                        {"role": "assistant", "content": full_response},
                        {"role": "user", "content": topic_prompt}
                    ]
                    topic_response = llm_service.chat(model_name, topic_messages, stream=False)
                    logger.info(f"Topic extraction raw response: {topic_response}")
                    # Robustly extract topic
                    topic = None
                    resp = topic_response
                    if isinstance(resp, dict):
                        # Try OpenAI-style choices
                        if 'choices' in resp and isinstance(resp.get('choices'), list) and resp['choices']:
                            choice = resp['choices'][0]
                            msg_obj = choice.get('message') or choice
                            if isinstance(msg_obj, dict):
                                topic = msg_obj.get('content') or msg_obj.get('text')
                        # Fallback to top-level message
                        elif 'message' in resp and isinstance(resp['message'], dict):
                            topic = resp['message'].get('content') or resp['message'].get('text')
                        # Fallback to text
                        elif 'text' in resp:
                            topic = resp.get('text')
                    elif isinstance(resp, str):
                        topic = resp
                    # Normalize topic
                    topic = (topic or "").strip().strip('"').strip()
                    # Update DB title if found
                    if topic:
                        conv = db.session.get(Conversation, conversation_id)
                        conv.title = topic
                        db.session.commit()
                    else:
                        logger.warning("Topic extraction returned empty topic")
                    logger.info(f"Emitting topic SSE: {topic} for conversation {conversation_id}")
                    # Always emit a topic SSE event
                    yield f"data: {json.dumps({'topic': topic})}\n\n"
                except Exception as e:
                    logger.error(f"Error extracting topic: {e}")
                    yield f"data: {json.dumps({'topic': ''})}\n\n"
            except Exception as e:
                logger.exception(f"Exception in streaming loop: {e}")
                error_text = json.dumps({"error": str(e), "text": f"⚠️ {str(e)}"})
                yield f"data: {error_text}\n\n"
        finally:
            logger.info("Exiting response_wrapper generator")
    logger.info("Returning streaming response")
    # Ensure correct mimetype for SSE and use stream_with_context
    return Response(stream_with_context(response_wrapper()), mimetype="text/event-stream")

# Add a dictionary to track active response generators
active_response_generators = {}

@app.route('/stop_response', methods=['POST'])
@login_required
def stop_response():
    """Stop an active AI response for a conversation"""
    conversation_id = request.form.get('conversation_id')
    logger.info(f"Request to stop response for conversation {conversation_id}")
    
    if not conversation_id:
        return jsonify({"success": False, "error": "No conversation ID provided"}), 400
        
    # Check permission (user must own the conversation)
    conversation = db.session.get(Conversation, conversation_id)
    if not conversation or conversation.user_id != current_user.id:
        return jsonify({"success": False, "error": "Unauthorized"}), 403
    
    # Set the flag to stop the generator for this conversation
    generator_key = f"user_{current_user.id}_conv_{conversation_id}"
    if generator_key in active_response_generators:
        active_response_generators[generator_key] = True
        logger.info(f"Set stop flag for generator {generator_key}")
        return jsonify({"success": True, "message": "Response generation stopping"})
    else:
        return jsonify({"success": False, "error": "No active response found"}), 404

@app.route('/toggle_document_mode/<int:conversation_id>', methods=['POST'])
@login_required
def toggle_document_mode(conversation_id):
    conversation = db.session.get(Conversation, conversation_id)
    
    # Check ownership
    if conversation.user_id != current_user.id:
        return jsonify({"success": False, "error": "Unauthorized"}), 403
        
    try:
        # Toggle document mode
        conversation.document_mode = not conversation.document_mode
        
        # Add a system message to indicate the change
        message_text = ""
        if conversation.document_mode:
            # Find the latest document
            latest_doc = Document.query.filter(
                Document.conversation_id == conversation_id,
                Document.mime_type != 'audio/wav' # Exclude voice recordings
            ).order_by(Document.uploaded_at.desc()).first()
            if latest_doc:
                message_text = f"📄 Document mode enabled. My responses will now be based only on document: '{latest_doc.filename}'."
            else:
                message_text = "📄 Document mode enabled, but no documents found. Please upload a document."
        else:
            message_text = "📄 Document mode disabled. I'll now use my general knowledge to answer your questions."
        
        system_message = ChatMessage(
            conversation_id=conversation_id,
            sender='ai',
            content=message_text
        )
        db.session.add(system_message)
        db.session.commit()
        
        return jsonify({
            "success": True,
            "document_mode": conversation.document_mode,
            "message": message_text
        })
    except Exception as e:
        logger.exception(f"Error toggling document mode: {str(e)}")
        db.session.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

# Add a new route to update conversation title directly
@app.route('/conversation/<int:conversation_id>/update_title', methods=['POST'])
@login_required
def update_conversation_title(conversation_id):
    """Update a conversation title and save it to the database"""
    try:
        conversation = db.session.get(Conversation, conversation_id)
        
        # Check ownership
        if conversation.user_id != current_user.id:
            return jsonify({"success": False, "error": "Unauthorized"}), 403
        
        data = request.get_json()
        if not data or 'title' not in data:
            return jsonify({"success": False, "error": "No title provided"}), 400
        
        title = data['title'].strip()
        if not title:
            title = "Untitled Conversation"
        
        # Update the title in the database
        conversation.title = title
        db.session.commit()
        
        logger.info(f"Title updated for conversation {conversation_id}: '{title}'")
        return jsonify({"success": True, "title": title})
    except Exception as e:
        logger.exception(f"Error updating conversation title: {str(e)}")
        db.session.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

# Add routes for text-to-speech capabilities
@app.route('/voice_for_message/<int:message_id>', methods=['GET'])
@login_required
def voice_for_message(message_id):
    """Check if a voice recording exists for a message and return it"""
    try:
        message = db.session.get(ChatMessage, message_id)
        
        # Check if this message belongs to a conversation owned by the current user
        conversation = db.session.get(Conversation, message.conversation_id)
        if not conversation or conversation.user_id != current_user.id:
            return "Unauthorized", 403
        
        # Check if this is a voice response message
        if message.sender == 'ai' and message.content.startswith('VOICE_RESPONSE:'):
            parts = message.content.split(':')
            if len(parts) >= 3:
                language = parts[1]
                content = ':'.join(parts[2:])
                
                # Check if we already have a voice recording for this AI message
                # We'll search for documents with a filename containing the message_id
                voice_doc = Document.query.filter(
                    Document.conversation_id == message.conversation_id,
                    Document.filename.like(f'ai_voice_response_%_{message_id}.wav')
                ).first()
                
                if voice_doc:
                    # Create a temporary file to serve
                    temp_file_path = None
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as tmp:
                        tmp.write(voice_doc.data)
                        temp_file_path = tmp.name
                    
                    # Function to clean up temporary file
                    def cleanup_temp_file():
                        try:
                            # Add a small delay to ensure file serving completes
                            time.sleep(1)
                            if os.path.exists(temp_file_path):
                                os.remove(temp_file_path)
                        except Exception as e:
                            logger.warning(f"Failed to remove temporary file {temp_file_path}: {e}")
                    
                    # Start cleanup thread
                    cleanup_thread = threading.Thread(target=cleanup_temp_file)
                    cleanup_thread.daemon = True
                    cleanup_thread.start()
                    
                    # Return the file
                    return send_file(
                        temp_file_path,
                        mimetype="audio/wav",
                        as_attachment=False
                    )
        
        # If we get here, no voice recording was found
        return "No voice recording found for this message", 404
    
    except Exception as e:
        logger.exception(f"Error retrieving voice for message: {e}")
        return f"Error retrieving voice: {e}", 500

@app.route('/synthesize_for_message', methods=['POST'])
@login_required
def synthesize_for_message():
    """Generate text-to-speech for a message and store it"""
    try:
        text = request.form['text']
        language = request.form.get('language', 'english')
        message_id = request.form.get('message_id')
        
        if not text.strip():
            return "No text provided", 400
        
        # Generate speech
        speech_file = synthesize_speech(text, language) # This function needs to be defined or imported
        if not speech_file:
            return "Failed to generate speech", 500
        
        # Find the associated message if message_id was provided
        if message_id:
            message = db.session.get(ChatMessage, int(message_id))
            if message:
                conversation_id = message.conversation_id
            else:
                return "Message not found", 404
        else:
            # If no specific message, use the active conversation
            conversation_id = request.form.get('conversation_id')
            if not conversation_id:
                return "No conversation ID provided", 400
        
        # Check user authorization for this conversation
        conversation = db.session.get(Conversation, conversation_id)
        if not conversation or conversation.user_id != current_user.id:
            return "Unauthorized", 403
        
        # Read the generated audio file
        with open(speech_file, 'rb') as audio_file:
            audio_data = audio_file.read()
        
        # Save the speech as a document with reference to the message
        filename = f"ai_voice_response_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
        if message_id:
            filename += f"_{message_id}"
        filename += ".wav"
        
        voice_doc = Document(
            conversation_id=conversation_id,
            filename=filename,
            data=audio_data,
            mime_type="audio/wav"
        )
        
        db.session.add(voice_doc)
        db.session.commit()
        
        # Clean up the temporary file
        try:
            os.remove(speech_file)
        except:
            pass
        
        # Return the audio
        return send_file(
            io.BytesIO(audio_data),
            mimetype="audio/wav",
            as_attachment=False
        )
    except Exception as e:
        logger.exception(f"Error synthesizing speech: {e}")
        return f"Error synthesizing speech: {e}", 500

# Add a function to check system dependencies at startup
def check_system_dependencies():
    deps = {
        "ffmpeg": check_ffmpeg_installed(),
        "whisper": check_whisper_model_exists() # Check if base model can load
    }
    return deps

# === Add Placeholder for synthesize_speech ===
def synthesize_speech(text, language):
    """Placeholder function for text-to-speech synthesis."""
    logger.warning(f"Placeholder synthesize_speech called for language '{language}'. Text: {text[:50]}...")
    # In a real implementation, this would call Bark via speech_service
    # and return the path to the generated audio file.
    # For now, return None to indicate failure.
    # Example call (if speech_service had this method):
    # return speech_service.synthesize_speech(text, language)
    return None
# === End Placeholder ===

# Create a session variable to store the user's LLM service preference
@app.route('/switch_llm_service', methods=['POST'])
@login_required
def switch_llm_service():
    """Switch the LLM service (Ollama or Llama.cpp) for the current user"""
    try:
        new_service = request.form.get('llm_service', 'ollama').lower()
        
        # Validate the service type
        if new_service not in ['ollama', 'llamacpp']:
            return jsonify({"success": False, "error": f"Invalid LLM service type: {new_service}"}), 400
        
        logger.info(f"User {current_user.id} switching LLM service to {new_service}")
        
        # Store the user's preference in session
        session['user_llm_service'] = new_service
        
        # Create a new LLM service instance
        global llm_service
        global llm_service_type
        
        # Create temp service to test connectivity
        try:
            # Use the factory to create the new service
            temp_service = LLMServiceFactory.create_service_by_type(new_service)
            # Test listing models to ensure connectivity
            available_models = temp_service.list_models()
            
            if not available_models:
                logger.warning(f"No models found for {new_service} service")
                
            # If we get here, the service is working
            llm_service = temp_service
            llm_service_type = new_service
            
            # Update the application's cached models list
            app.config['LLM_MODELS'] = available_models if available_models else [DEFAULT_MODEL_NAME]
            logger.info(f"Successfully switched to {new_service} service with models: {app.config['LLM_MODELS']}")
            
            return jsonify({
                "success": True, 
                "service": new_service,
                "models": app.config['LLM_MODELS']
            })
            
        except Exception as e:
            logger.exception(f"Error switching to {new_service} service: {e}")
            return jsonify({"success": False, "error": f"Failed to connect to {new_service} service: {str(e)}"}), 500
            
    except Exception as e:
        logger.exception(f"Error in switch_llm_service: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/test_ollama', methods=['GET'])
def test_ollama():
    """
    Test connectivity to Ollama service and return detailed diagnostics.
    """
    results = {
        "status": "Running diagnostics...",
        "ollama_host": os.environ.get("OLLAMA_HOST", "http://ollama:11434"),
        "llm_service_type": os.environ.get("LLM_SERVICE", "ollama"),
        "tests": []
    }
    
    try:
        # Test 1: Basic connectivity via requests
        results["tests"].append({"name": "Basic connectivity test"})
        try:
            import requests
            ollama_url = os.environ.get("OLLAMA_HOST", "http://ollama:11434")
            health_url = f"{ollama_url}/api/tags"
            logger.info(f"Testing basic connectivity to Ollama at {health_url}")
            response = requests.get(health_url, timeout=5)
            results["tests"][-1]["status"] = f"Success ({response.status_code})"
            results["tests"][-1]["details"] = f"Connected to {health_url}"
            
            # Include the first 500 chars of the response for verification
            response_data = response.json()
            results["tests"][-1]["response_preview"] = str(response_data)[:500]
        except Exception as e:
            results["tests"][-1]["status"] = "Failed"
            results["tests"][-1]["details"] = f"Error: {str(e)}"
            logger.exception(f"Basic connectivity test failed: {e}")
        
        # Test 2: Try to list models
        results["tests"].append({"name": "List models test"})
        try:
            models = llm_service.list_models()
            results["tests"][-1]["status"] = "Success"
            results["tests"][-1]["details"] = f"Found {len(models)} models"
            results["tests"][-1]["models"] = models
        except Exception as e:
            results["tests"][-1]["status"] = "Failed"
            results["tests"][-1]["details"] = f"Error: {str(e)}"
            logger.exception(f"List models test failed: {e}")
        
        # Test 3: Simple completion without streaming
        results["tests"].append({"name": "Simple completion test"})
        try:
            import requests
            import json
            
            ollama_url = os.environ.get("OLLAMA_HOST", "http://ollama:11434")
            api_url = f"{ollama_url}/api/chat"
            payload = {
                "model": DEFAULT_MODEL_NAME,
                "messages": [{"role": "user", "content": "Hello, say hi in one word"}],
                "stream": False
            }
            logger.info(f"Testing simple completion to {api_url}")
            response = requests.post(api_url, json=payload, timeout=10)
            
            if response.status_code == 200:
                response_data = response.json()
                results["tests"][-1]["status"] = "Success"
                results["tests"][-1]["details"] = "Completion received"
                results["tests"][-1]["response"] = str(response_data)[:500]
            else:
                results["tests"][-1]["status"] = "Failed"
                results["tests"][-1]["details"] = f"Error: Status {response.status_code}"
                results["tests"][-1]["response"] = response.text[:500]
        except Exception as e:
            results["tests"][-1]["status"] = "Failed"
            results["tests"][-1]["details"] = f"Error: {str(e)}"
            logger.exception(f"Simple completion test failed: {e}")
        
        # Overall status
        failed_tests = [t for t in results["tests"] if t.get("status", "").startswith("Failed")]
        if failed_tests:
            results["status"] = f"Failed ({len(failed_tests)}/{len(results['tests'])} tests failed)"
        else:
            results["status"] = "Success (all tests passed)"
            
    except Exception as e:
        results["status"] = "Error running diagnostics"
        results["error"] = str(e)
        logger.exception(f"Error in /test_ollama endpoint: {e}")
    
    return jsonify(results)

# Main entry point
def initialize_rbac_data():
    logger.info("FUNC_INIT_RBAC: Entered initialize_rbac_data function.") # Cascade Temp Log
    """
    Initializes default roles and populates the Model table
    from available Ollama models if they don't already exist.
    Also assigns all found models to the 'admin' role.
    """
    with app.app_context():
        logger.info("FUNC_INIT_RBAC: Attempting to process roles and models...") # Cascade Temp Log (was: Initializing RBAC data (roles and models)...)

        # 1. Create default roles
        default_roles_data = {
            "admin": "Administrator with full access to all models and system settings.",
            "user": "Standard user with access to a default set of models."
        }
        admin_role_obj = None
        user_role_obj = None

        for role_name, role_desc in default_roles_data.items():
            role = Role.query.filter_by(name=role_name).first()
            if not role:
                role = Role(name=role_name, description=role_desc)
                db.session.add(role)
                logger.info(f"Created role: {role_name}")
            if role_name == "admin":
                admin_role_obj = role
            elif role_name == "user":
                user_role_obj = role
        
        try:
            db.session.commit() # Commit roles first
        except Exception as e:
            db.session.rollback()
            logger.error(f"Error committing roles: {e}", exc_info=True)
            return # Cannot proceed without roles

        if not admin_role_obj:
            logger.error("Admin role could not be found or created. Cannot proceed.")
            return

        # 2. Create and assign 'Admin' role to the default admin user (admin@admin.com)
        default_admin_email = "admin@admin.com"
        admin_user_obj = User.query.filter_by(email=default_admin_email).first()
        if not admin_user_obj:
            logger.info(f"Default admin user '{default_admin_email}' not found. Creating new admin user.")
            hashed_password = generate_password_hash("admin")
            admin_user_obj = User(
                username="admin",
                email=default_admin_email,
                firstname="admin",
                lastname="admin",
                password_hash=hashed_password,
                confirmed=True, 
                is_active=True
            )
            db.session.add(admin_user_obj)
            try:
                db.session.commit() # Commit new admin user first
                logger.info(f"Successfully created default admin user '{default_admin_email}'.")
            except Exception as e:
                db.session.rollback()
                logger.error(f"Error creating default admin user '{default_admin_email}': {e}", exc_info=True)
                # Don't return, try to proceed with other admin if configured
        
        # Ensure the default admin user is active if they exist
        if admin_user_obj and not admin_user_obj.is_active:
            logger.info(f"Ensuring default admin user '{default_admin_email}' is active.")
            admin_user_obj.is_active = True
            # Commit this change before proceeding with role assignment
            try:
                db.session.commit()
                logger.info(f"Default admin user '{default_admin_email}' set to active.")
            except Exception as e:
                db.session.rollback()
                logger.error(f"Error setting default admin user '{default_admin_email}' to active: {e}", exc_info=True)

        # Assign Admin role to the default admin user if they exist and don't have it
        if admin_user_obj and admin_role_obj not in admin_user_obj.roles:
            admin_user_obj.roles.append(admin_role_obj)
            try:
                db.session.commit()
                logger.info(f"Assigned 'admin' role to default admin user '{default_admin_email}'.")
            except Exception as e:
                db.session.rollback()
                logger.error(f"Error assigning 'admin' role to default admin user '{default_admin_email}': {e}", exc_info=True)
        elif admin_user_obj:
            logger.info(f"Default admin user '{default_admin_email}' already has 'admin' role or was just created with it.")

        # 3. Assign 'Admin' role to the admin user from .env (if different from default and exists)
        admin_username_env = app.config.get('ADMIN_USERNAME')
        if admin_username_env and admin_username_env != "admin": # Check if .env admin is set and different from default 'admin'
            env_admin_user = User.query.filter_by(username=admin_username_env).first()
            if env_admin_user:
                if admin_role_obj not in env_admin_user.roles:
                    env_admin_user.roles.append(admin_role_obj)
                    try:
                        db.session.commit()
                        logger.info(f"Assigned 'admin' role to .env admin user '{admin_username_env}'.")
                    except Exception as e:
                        db.session.rollback()
                        logger.error(f"Error assigning 'admin' role to .env admin user '{admin_username_env}': {e}", exc_info=True)
                else:
                    logger.info(f".env admin user '{admin_username_env}' already has 'admin' role.")
            else:
                logger.warning(f"Admin user '{admin_username_env}' specified in .env not found in database. Cannot assign 'admin' role.")
        elif not admin_username_env:
            logger.info("ADMIN_USERNAME not set in .env file. Skipping .env admin role assignment.")
        elif admin_username_env == "admin":
            logger.info("ADMIN_USERNAME in .env is 'admin', which is already handled as the default admin. Skipping redundant assignment.")

        # 3. Populate Model table from MANAGED_OLLAMA_MODELS and assign to admin
        logger.info("Processing managed models for Model table and admin assignment...")
        
        admin_models_assigned_this_run = [] # To track models assigned to admin in this run
        llm_service_type = app.config.get('LLM_SERVICE_TYPE', 'ollama').lower()


        if llm_service_type == 'ollama':
            managed_ollama_models_from_config = app.config.get('MANAGED_OLLAMA_MODELS', [])
            if not managed_ollama_models_from_config:
                logger.warning("No managed Ollama models found in app.config (MANAGED_OLLAMA_MODELS is empty or not set). "
                               "Model table population from this list will be skipped.")
            
            for ollama_model_name in managed_ollama_models_from_config:
                model = Model.query.filter_by(ollama_model_name=ollama_model_name).first()
                if not model:
                    # Attempt to create a somewhat friendly display name
                    display_parts = ollama_model_name.split(':')[0].replace('-', ' ').replace('_', ' ')
                    display_name_generated = ' '.join(word.capitalize() for word in display_parts.split(' '))
                    tag_suffix = ""
                    if ':' in ollama_model_name:
                        tag = ollama_model_name.split(':')[-1]
                        if tag.lower() != 'latest':
                             tag_suffix = f" ({tag.capitalize()})"
                        # else: # For 'latest', no specific tag suffix or a generic one like "(Latest)"
                             # tag_suffix = " (Latest)" # Optional: if you want to explicitly mark 'latest'
                    final_display_name = display_name_generated + tag_suffix

                    model = Model(
                        display_name=final_display_name, # This is the user-facing name
                        ollama_model_name=ollama_model_name, # This is for Ollama API
                        description=f"Ollama model: {ollama_model_name}",
                        is_active=True # Managed models are active by default
                    )
                    db.session.add(model)
                    logger.info(f"Created new Model DB entry for managed model: {ollama_model_name} (Display: {final_display_name})")
                    try:
                        db.session.commit() # Commit each new model to get its ID for relationships
                    except Exception as e:
                        db.session.rollback()
                        logger.error(f"Error committing new model {ollama_model_name}: {e}", exc_info=True)
                        continue # Skip to next model
                elif not model.is_active:
                    logger.info(f"Model '{ollama_model_name}' found in DB but was inactive. Activating it as it's a managed model.")
                    model.is_active = True
                    # No immediate commit needed here, will be committed with role assignment or at the end of this section.

                # Assign to admin role if not already assigned and model object exists
                if admin_role_obj and model and model.id is not None: # Ensure model is committed or fetched with an ID
                    if model not in admin_role_obj.models:
                        admin_role_obj.models.append(model)
                        admin_models_assigned_this_run.append(model.display_name)
                        logger.info(f"Assigned model '{model.display_name}' to 'admin' role.")
                elif not model:
                    logger.warning(f"Skipped assigning model {ollama_model_name} to admin as model object was None (likely due to creation error).")
        
        elif llm_service_type == 'llamacpp':
            logger.info("LLM service is LlamaCPP. Checking/creating DB entry for its default model.")
            llamacpp_default_model_name = app.config.get('EFFECTIVE_DEFAULT_MODEL_NAME')
            if llamacpp_default_model_name:
                model = Model.query.filter_by(ollama_model_name=llamacpp_default_model_name).first()
                if not model:
                    model = Model(
                        display_name=llamacpp_default_model_name, 
                        ollama_model_name=llamacpp_default_model_name, # Using ollama_model_name field for identifier consistency
                        description=f"LlamaCPP model: {llamacpp_default_model_name}",
                        is_active=True
                    )
                    db.session.add(model)
                    logger.info(f"Created new Model DB entry for LlamaCPP default model: {llamacpp_default_model_name}")
                    try:
                        db.session.commit() # Commit to get ID
                    except Exception as e:
                        db.session.rollback()
                        logger.error(f"Error committing LlamaCPP default model {llamacpp_default_model_name} to DB: {e}", exc_info=True)
                        model = None # Ensure model is None if commit failed
                
                if admin_role_obj and model and model.id is not None:
                    if model not in admin_role_obj.models:
                        admin_role_obj.models.append(model)
                        admin_models_assigned_this_run.append(model.display_name) # CORRECTED HERE
                        logger.info(f"Assigned LlamaCPP model '{model.display_name}' to 'admin' role.") # Also ensure log uses display_name
            else:
                logger.warning("LlamaCPP service type, but no EFFECTIVE_DEFAULT_MODEL_NAME found in app.config.")
        else:
            logger.info(f"LLM service type is '{llm_service_type}'. Model DB population from managed list is primarily for Ollama.")

        if admin_models_assigned_this_run:
            logger.info(f"Models assigned/confirmed for admin role in this initialization run: {', '.join(admin_models_assigned_this_run)}")

        try:
            db.session.commit() # Commit all changes (new models, model activations, role assignments)
            logger.info("Committed all model and admin role assignment changes for initialize_rbac_data.")
        except Exception as e:
            db.session.rollback()
            logger.error(f"Final commit in initialize_rbac_data failed: {e}", exc_info=True)

        logger.info("FUNC_INIT_RBAC: Exiting initialize_rbac_data function.") # Cascade Temp Logn managed/default models complete.")
# ===========================
@app.route('/admin/rbac')
@login_required
def admin_rbac_page():
    if not current_user.has_role('admin'):
        flash("You must be an administrator to access this page.", "danger")
        return redirect(url_for('index'))
    
    if not current_user.can_access_page('admin_rbac_page'):
        flash('You do not have permission to access this page.', 'danger')
        return redirect(url_for('index'))

    try:
        all_users = User.query.all()
        all_roles = Role.query.all()
        all_models = Model.query.all()
    except Exception as e:
        logger.error(f"Error fetching data for RBAC page: {e}", exc_info=True)
        flash("Error loading RBAC management page data.", "danger")
        return redirect(url_for('index'))
        
    # Define application pages/resources for access management display
    # In a more complex app, these might come from a config or be discovered
    # Prepare current page permissions for the template
    current_page_permissions = {}
    for role in all_roles:
        current_page_permissions[role.id] = {}
        for page in MANAGED_PAGE_ENDPOINTS:
            current_page_permissions[role.id][page['endpoint']] = role.has_page_access(page['endpoint'])

    current_model_permissions = {}
    for role in all_roles:
        current_model_permissions[role.id] = {}
        # Ensure all_models is a list of Model objects, not just names or IDs
        # The admin_rbac_page already queries all_models = Model.query.all()
        for model_obj in all_models: # Iterate through the fetched Model objects
            current_model_permissions[role.id][model_obj.id] = role.has_model_access(model_obj.id)
    
    return render_template('admin_rbac.html', 
                           users=all_users, 
                           roles=all_roles, 
                           models=all_models, 
                           app_pages=MANAGED_PAGE_ENDPOINTS, 
                           current_page_permissions=current_page_permissions,
                           current_model_permissions=current_model_permissions)


@app.route('/admin/permissions/page_access/update', methods=['POST'])
@login_required
def update_page_access_permissions():
    if not current_user.has_role('admin'):
        return jsonify({'success': False, 'message': 'You do not have permission to perform this action.'}), 403

    try:
        data = request.get_json()
        permissions_to_set = data.get('permissions', []) # Expected: [{'role_id': X, 'page_endpoint': 'Y'}, ...]

        # Clear existing permissions for all managed pages
        # This is a simple approach; a more complex app might do more granular updates.
        existing_managed_endpoints = [p['endpoint'] for p in MANAGED_PAGE_ENDPOINTS]
        PagePermission.query.filter(PagePermission.page_endpoint.in_(existing_managed_endpoints)).delete(synchronize_session=False)
        # Note: synchronize_session=False is used here. If issues arise, consider 'fetch' or individual deletes.

        # Add new permissions
        for perm_data in permissions_to_set:
            role_id = perm_data.get('role_id')
            page_endpoint = perm_data.get('page_endpoint')

            if not role_id or not page_endpoint:
                logger.warning(f"Skipping invalid permission data: {perm_data}")
                continue
            
            role = Role.query.get(role_id)
            if not role:
                logger.warning(f"Role ID {role_id} not found while updating page permissions.")
                continue
            
            # Ensure the page_endpoint is one of the known manageable endpoints
            if page_endpoint not in existing_managed_endpoints:
                logger.warning(f"Attempt to set permission for unmanaged page_endpoint '{page_endpoint}'. Skipping.")
                continue

            new_permission = PagePermission(role_id=role.id, page_endpoint=page_endpoint)
            db.session.add(new_permission)
        
        db.session.commit()
        logger.info(f"Page access permissions updated by user {current_user.id}.")
        return jsonify({'success': True, 'message': 'Page access permissions updated successfully.'})

    except Exception as e:
        db.session.rollback()
        logger.error(f"Error updating page access permissions: {e}", exc_info=True)
        return jsonify({'success': False, 'message': 'Failed to update page access permissions due to a server error.'}), 500


@app.route('/admin/permissions/model_access/update', methods=['POST'])
@login_required
def update_model_access_permissions():
    if not current_user.has_role('admin'):
        return jsonify({'success': False, 'message': 'You do not have permission to perform this action.'}), 403

    try:
        data = request.get_json()
        role_id = data.get('role_id')
        model_ids_to_assign = data.get('model_ids', [])  # Expected: [1, 2, 3]

        if role_id is None:
            logger.warning("Role ID not provided in request to update_model_access_permissions.")
            return jsonify({'success': False, 'message': 'Role ID is required.'}), 400

        role = Role.query.get(role_id)
        if not role:
            logger.warning(f"Role ID {role_id} not found while updating model access permissions.")
            return jsonify({'success': False, 'message': f'Role with ID {role_id} not found.'}), 404

        # Fetch valid Model objects based on provided IDs
        valid_models = []
        if model_ids_to_assign:
            valid_models = Model.query.filter(Model.id.in_(model_ids_to_assign)).all()
            
            # Log if some provided model IDs were not found, but proceed with valid ones
            assigned_model_ids = {model.id for model in valid_models}
            invalid_ids_provided = [mid for mid in model_ids_to_assign if mid not in assigned_model_ids]
            if invalid_ids_provided:
                logger.warning(f"Invalid or non-existent model IDs {invalid_ids_provided} provided for role '{role.name}' (ID: {role.id}). These will be ignored.")

        # Update the role's associated models
        # SQLAlchemy automatically handles the changes in the 'role_models' association table
        role.models = valid_models
        
        db.session.commit()
        logger.info(f"Model access permissions for role '{role.name}' (ID: {role.id}) updated by admin user {current_user.id}. Assigned model IDs: {[m.id for m in valid_models]}.")
        return jsonify({'success': True, 'message': f'Model access permissions for role \'{role.name}\' updated successfully.'})

    except Exception as e:
        db.session.rollback()
        # Try to get role_id from data if available for logging, otherwise use a placeholder
        requested_role_id = data.get('role_id', 'N/A') if isinstance(data, dict) else 'N/A'
        logger.error(f"Error updating model access permissions for role ID {requested_role_id}: {e}", exc_info=True)
        return jsonify({'success': False, 'message': 'Failed to update model access permissions due to a server error.'}), 500



@app.route('/admin/user/assign_roles/<int:user_id>', methods=['POST'])
@login_required
def assign_user_roles(user_id):
    if not current_user.has_role('admin'):
        return jsonify({'success': False, 'message': 'You do not have permission to perform this action.'}), 403

    user = User.query.get_or_404(user_id)
    # Ensure admin cannot unwittlingly remove their own admin role if they are the only admin
    # or de-admin themselves via this mechanism if they are the user being edited.
    if user.id == current_user.id:  # Check if the user being edited is the current admin
        admin_role = Role.query.filter_by(name='admin').first()
        # If the 'admin' role exists and its ID is NOT in the list of roles to be assigned to self
        if admin_role and str(admin_role.id) not in request.form.getlist('role_ids[]'):
            # Check if there are other admins in the system
            other_admins = User.query.join(User.roles).filter(Role.name == 'admin', User.id != user.id).count()
            if other_admins == 0:
                return jsonify({'success': False, 'message': 'As the only administrator, you cannot remove your own admin role.'}), 400

    role_ids = request.form.getlist('role_ids[]') # Assuming role_ids are sent as a list
    
    # logger.info(f"Assigning roles to user {user_id}: {role_ids}") # For debugging

    # Clear existing roles
    user.roles.clear()

    # Add new roles
    if role_ids:
        for role_id_str in role_ids:
            try:
                role_id = int(role_id_str)
                role = Role.query.get(role_id)
                if role:
                    user.roles.append(role)
                else:
                    logger.warning(f"Role ID {role_id} not found while assigning roles to user {user_id}.")
            except ValueError:
                logger.warning(f"Invalid role ID {role_id_str} received for user {user_id}.")
                # Optionally, return an error here if strict validation is needed

    try:
        db.session.commit()
        # Fetch the updated roles to send back for potential UI update
        updated_role_names = [role.name for role in user.roles]
        logger.info(f"Successfully updated roles for user {user.id} to: {updated_role_names}")
        return jsonify({'success': True, 'message': 'User roles updated successfully.', 'user_id': user_id, 'new_roles': updated_role_names})
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error updating roles for user {user_id}: {e}", exc_info=True)
        return jsonify({'success': False, 'message': 'Failed to update roles due to a server error.'}), 500


@app.route('/admin/roles/create', methods=['POST'])
@login_required
def create_role():
    if not current_user.has_role('admin'):
        flash("You do not have permission to perform this action.", "danger")
        return redirect(url_for('admin_rbac_page'))

    role_name = request.form.get('role_name')
    role_description = request.form.get('role_description')

    if not role_name:
        flash('Role name is required.', 'warning')
        return redirect(url_for('admin_rbac_page'))

    existing_role = Role.query.filter_by(name=role_name).first()
    if existing_role:
        flash(f'Role "{role_name}" already exists.', 'warning')
        return redirect(url_for('admin_rbac_page'))

    try:
        new_role = Role(name=role_name, description=role_description)
        db.session.add(new_role)
        db.session.commit()
        flash(f'Role "{role_name}" created successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error creating role '{role_name}': {e}", exc_info=True)
        flash('Error creating role. Please check logs.', 'danger')

    return redirect(url_for('admin_rbac_page'))


@app.route('/admin/role/edit/<int:role_id>', methods=['GET'])
@login_required
def edit_role_page(role_id):
    if not current_user.has_role('admin'):
        flash("You do not have permission to perform this action.", "danger")
        return redirect(url_for('admin_rbac_page'))
    role = Role.query.get_or_404(role_id)
    return render_template('edit_role.html', role=role)

@app.route('/admin/role/edit/<int:role_id>', methods=['POST'])
@login_required
def edit_role_submit(role_id):
    if not current_user.has_role('admin'):
        flash("You do not have permission to perform this action.", "danger")
        return redirect(url_for('admin_rbac_page'))

    role = Role.query.get_or_404(role_id)
    new_name = request.form.get('role_name')
    new_description = request.form.get('role_description')

    if not new_name:
        flash('Role name is required.', 'warning')
        return render_template('edit_role.html', role=role) # Stay on edit page with error

    # Check if new name conflicts with an existing role (excluding the current role itself)
    existing_role_with_new_name = Role.query.filter(Role.name == new_name, Role.id != role_id).first()
    if existing_role_with_new_name:
        flash(f'Role name "{new_name}" already exists. Please choose a different name.', 'warning')
        return render_template('edit_role.html', role=role) # Stay on edit page with error

    role.name = new_name
    role.description = new_description

    try:
        db.session.commit()
        flash(f'Role "{role.name}" updated successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error updating role '{role.name}': {e}", exc_info=True)
        flash('Error updating role. Please check logs.', 'danger')
    
    return redirect(url_for('admin_rbac_page'))


@app.route('/admin/users/create', methods=['POST'])
@login_required
def create_user_from_admin():
    if not current_user.has_role('admin'):
        flash("You do not have permission to perform this action.", "danger")
        return redirect(url_for('admin_rbac_page'))

    username = request.form.get('username')
    email = request.form.get('email')
    password = request.form.get('password')

    if not username or not email or not password:
        flash('Username, email, and password are required.', 'warning')
        return redirect(url_for('admin_rbac_page'))

    if User.query.filter_by(username=username).first():
        flash(f'Username "{username}" already exists.', 'warning')
        return redirect(url_for('admin_rbac_page'))
    
    if User.query.filter_by(email=email).first():
        flash(f'Email "{email}" already registered.', 'warning')
        return redirect(url_for('admin_rbac_page'))

    try:
        new_user = User(username=username, email=email)
        new_user.set_password(password)
        new_user.confirmed = True  # Admins create confirmed users
        # Optionally, assign a default role (e.g., 'user')
        # user_role = Role.query.filter_by(name='user').first()
        # if user_role:
        #     new_user.roles.append(user_role)
        
        db.session.add(new_user)
        db.session.commit()
        flash(f'User "{username}" created successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error creating user '{username}': {e}", exc_info=True)
        flash('Error creating user. Please check logs.', 'danger')

    return redirect(url_for('admin_rbac_page'))


@app.route('/admin/user/toggle_active/<int:user_id>', methods=['POST'])
@login_required
def toggle_user_active(user_id):
    logger.info(f"TOGGLE_USER_ACTIVE: current_user ID: {current_user.id}, Username: {current_user.username}, Email: {current_user.email}")
    logger.info(f"TOGGLE_USER_ACTIVE: current_user roles: {[role.name for role in current_user.roles]}")
    if not current_user.has_role('admin'):
        return jsonify({'success': False, 'message': 'You do not have permission to perform this action.', 'user_id': user_id, 'new_status': None}), 403

    user_to_toggle = User.query.get_or_404(user_id)

    # Prevent admins from deactivating themselves if they are the only admin
    if user_to_toggle.id == current_user.id and 'admin' in [role.name for role in user_to_toggle.roles]:
        other_admins = User.query.join(User.roles).filter(Role.name == 'admin', User.id != user_to_toggle.id).count()
        if other_admins == 0:
            return jsonify({'success': False, 'message': 'As the only administrator, you cannot deactivate your own account.'}), 400

    try:
        user_to_toggle.is_active = not user_to_toggle.is_active
        db.session.commit()
        status_text = "activated" if user_to_toggle.is_active else "deactivated"
        logger.info(f"User {user_to_toggle.username} has been {status_text}.")
        return jsonify({'success': True, 'message': f'User {user_to_toggle.username} has been {status_text}.', 'is_active': user_to_toggle.is_active})
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error toggling active status for user {user_id}: {e}", exc_info=True)
        return jsonify({'success': False, 'message': 'An unexpected error occurred.', 'is_active': user_to_toggle.is_active if 'user_to_toggle' in locals() and hasattr(user_to_toggle, 'is_active') else None}), 500



@app.route('/admin/user/update', methods=['POST'])
@login_required
def update_user_details():
    logger.info(f"UPDATE_USER_DETAILS: current_user ID: {current_user.id}, Username: {current_user.username}")
    if not current_user.has_role('admin'):
        logger.warning(f"UPDATE_USER_DETAILS: Unauthorized attempt by user {current_user.id}")
        return jsonify({'success': False, 'error': 'Unauthorized access attempt.'}), 403

    data = request.get_json()
    if not data:
        logger.error("UPDATE_USER_DETAILS: No JSON data received.")
        return jsonify({'success': False, 'error': 'No data received.'}), 400

    user_id = data.get('user_id')
    new_username = data.get('username', '').strip()
    new_email = data.get('email', '').strip()

    if not user_id or not isinstance(user_id, int):
        logger.error(f"UPDATE_USER_DETAILS: Invalid or missing user_id: {user_id}")
        return jsonify({'success': False, 'error': 'Invalid or missing user ID.'}), 400
    
    if not new_username:
        logger.warning(f"UPDATE_USER_DETAILS: Username cannot be empty for user_id {user_id}.")
        return jsonify({'success': False, 'error': 'Username cannot be empty.'}), 400

    if not new_email:
        logger.warning(f"UPDATE_USER_DETAILS: Email cannot be empty for user_id {user_id}.")
        return jsonify({'success': False, 'error': 'Email cannot be empty.'}), 400

    user_to_update = User.query.get(user_id)
    if not user_to_update:
        logger.error(f"UPDATE_USER_DETAILS: User with ID {user_id} not found.")
        return jsonify({'success': False, 'error': 'User not found.'}), 404

    # Check for username conflict (if changed and new username exists for another user)
    if user_to_update.username != new_username:
        existing_user_by_username = User.query.filter(User.id != user_id, User.username == new_username).first()
        if existing_user_by_username:
            logger.warning(f"UPDATE_USER_DETAILS: Username '{new_username}' already taken by user ID {existing_user_by_username.id}.")
            return jsonify({'success': False, 'error': f'Username \"{new_username}\" is already taken.'}), 409

    # Check for email conflict (if changed and new email exists for another user)
    if user_to_update.email != new_email:
        existing_user_by_email = User.query.filter(User.id != user_id, User.email == new_email).first()
        if existing_user_by_email:
            logger.warning(f"UPDATE_USER_DETAILS: Email '{new_email}' already registered to user ID {existing_user_by_email.id}.")
            return jsonify({'success': False, 'error': f'Email \"{new_email}\" is already registered.'}), 409

    try:
        user_to_update.username = new_username
        user_to_update.email = new_email
        db.session.commit()
        logger.info(f"UPDATE_USER_DETAILS: User {user_id} ({new_username}) updated successfully.")
        return jsonify({'success': True, 'message': 'User details updated successfully.'})
    except Exception as e:
        db.session.rollback()
        logger.error(f"UPDATE_USER_DETAILS: Error updating user {user_id}: {e}", exc_info=True)
        return jsonify({'success': False, 'error': 'An internal error occurred while updating the user.'}), 500



if __name__ == '__main__':
    # Check dependencies
    system_deps = check_system_dependencies()
    if not system_deps['ffmpeg']:
        logger.error("FFmpeg is not installed. Please install FFmpeg to use this application.")
        # Don't exit in Docker, let it try to run
        # exit(1) 
    if not system_deps['whisper']:
        logger.error("Whisper models are not available. Please install faster-whisper to use this application.")
        # Don't exit in Docker, let it try to run
        # exit(1)
    
    # === Add a startup message showing the LLM service type ===
    logger.info(f"Starting AI Chat application with {llm_service_type.upper()} as the LLM service")
    
    logger.info("MAIN: Attempting to enter app_context for db.create_all()...") # Cascade Temp Log
    with app.app_context():
        logger.info("MAIN: Inside app_context. Attempting db.create_all()...") # Cascade Temp Log
        db.create_all()
        initialize_rbac_data() # Initialize roles and models
    
    # === Update app.run for Docker ===
    # Use host='0.0.0.0' to be accessible outside the container
    # Use port=5001 as exposed in Dockerfile/docker-compose.yml


app.run(debug=True, host='0.0.0.0', port=5001)