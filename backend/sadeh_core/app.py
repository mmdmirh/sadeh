# In order to use the latest version of sqlite3, we need to import it before any other libraries that might use it.
# This is a workaround for the fact that the default sqlite3 version in the system is too old for chromadb.
# see: https://docs.trychroma.com/troubleshooting#sqlite
try:
    __import__('pysqlite3')
    import sys
    sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')
except ImportError:
    print("pysqlite3 not found, using default sqlite3")

import os
from backend.db_folder.chroma_db import get_chroma_client, get_embedding_function





import sys
import io
import tempfile
import datetime
import json
import wave
import logging
import glob
import time
import threading
import tiktoken
import traceback
import re
import concurrent.futures


from flask import Flask, render_template, request, redirect, url_for, flash, Response, send_file, jsonify, session, stream_with_context
from sqlalchemy import create_engine, event, text, inspect, and_, or_, not_, desc, asc, func, exc
from sqlalchemy.orm import joinedload


from backend.db.db.extensions import db
from backend.db.db.models import User, Conversation, ChatMessage, Document, PagePermission, Role, RagDocument, RagIndex, Model, user_roles, role_models, rag_doc_in_index
from flask_login import LoginManager, login_user, logout_user, current_user, login_required, UserMixin
from flask_migrate import Migrate
from flask_mail import Mail, Message
from werkzeug.security import generate_password_hash, check_password_hash
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
import time
from flask_wtf.csrf import CSRFProtect
import subprocess
from dotenv import load_dotenv

# Load environment variables from the deploy/.env file for local development
load_dotenv(dotenv_path='deploy/.env')
from ollama import RequestError, ResponseError
import uuid
import threading
import subprocess
from pathlib import Path

# Langchain imports for RAG
# Using the newer package-specific imports to avoid deprecation warnings
from langchain_chroma import Chroma
import chromadb # Keep this for now, will remove if not needed after other changes

from langchain_community.document_loaders import PyPDFLoader, TextLoader, UnstructuredFileLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_ollama import OllamaEmbeddings

# Import our new LLM service abstraction
from backend.llm.llm_service import LLMServiceFactory

# ===========================
# Background Model Pull
import threading
import os
import ollama

def _run_pull_command(model_name, app_instance):
    with app_instance.app_context():
        try:
            app_instance.logger.info(f'Starting background pull for embedding model: {model_name}')
            ollama_host = os.getenv('OLLAMA_HOST', 'http://localhost:11434')
            client = ollama.Client(host=ollama_host)
            response = client.pull(model_name)
            # The ollama.pull method streams output, so we need to consume it
            for chunk in response:
                app_instance.logger.debug(f'Ollama pull progress for {model_name}: {chunk}')
            app_instance.logger.info(f'Successfully pulled embedding model: {model_name}')
        except ollama.ResponseError as e:
            app_instance.logger.error(f'Failed to pull embedding model {model_name} via Ollama API. Error: {e}')
        except Exception as e:
            app_instance.logger.error(f'An unexpected error occurred while pulling embedding model {model_name}: {e}')

def pull_embedding_models_in_background(app_instance):
    app_instance.logger.info("--- Running V3 of pull_embedding_models_in_background ---")
    embedding_models_str = os.getenv('OLLAMA_EMBEDDING_MODELS')
    app_instance.logger.info(f"Raw OLLAMA_EMBEDDING_MODELS value: '{embedding_models_str}'")
    
    if not embedding_models_str:
        app_instance.logger.info("No OLLAMA_EMBEDDING_MODELS set in .env, skipping background pull.")
        return
    
    # Handle a potentially malformed string that includes the variable name
    if embedding_models_str.startswith('OLLAMA_EMBEDDING_MODELS='):
        app_instance.logger.warning(f"OLLAMA_EMBEDDING_MODELS value includes variable name. Fixing: '{embedding_models_str}'")
        embedding_models_str = embedding_models_str.replace('OLLAMA_EMBEDDING_MODELS=', '')
    
    # Clean up potential "ollama pull" prefix from the model name
    embedding_models = [
        model.replace('ollama pull ', '').strip()
        for model in embedding_models_str.split(',')
    ]
    app_instance.logger.info(f"Queueing background pull for embedding models: {embedding_models}")

    for model_name in embedding_models:
        if model_name:
            thread = threading.Thread(target=_run_pull_command, args=(model_name, app_instance,))
            thread.daemon = True
            thread.start()

# ===========================
# Flask App Initialization
# ===========================
app = Flask(__name__, static_folder='../../frontend/static', template_folder='../../frontend/templates')

# === Set log level from environment variable (default to INFO) ===
import logging
log_level = os.getenv('LOG_LEVEL', 'INFO').upper()
if hasattr(logging, log_level):
    app.logger.setLevel(getattr(logging, log_level))
else:
    app.logger.setLevel(logging.INFO)
app.logger.info(f"[INIT] Log level set to {app.logger.level} ({log_level})")

# Set log level from environment variable (default to INFO)
import os, logging
log_level = os.getenv('LOG_LEVEL', 'INFO').upper()
if hasattr(logging, log_level):
    app.logger.setLevel(getattr(logging, log_level))
else:
    app.logger.setLevel(logging.INFO)
app.logger.info(f"[INIT] Log level set to {app.logger.level} ({log_level})")

# Start pulling embedding models in the background
pull_embedding_models_in_background(app)

# Speech service functionality has been removed
from flask import Blueprint # Added for Blueprint
from flask import jsonify

# Logging is now handled by Gunicorn and the Flask app app.logger.
# We will use app.logger for all application-level logging.
import logging

@app.route('/api/rag_index/progress/<int:index_id>')
def get_rag_index_progress(index_id):
    index = RagIndex.query.get(index_id)
    if not index:
        return jsonify({'status': 'not_found', 'error_message': 'Index not found.'}), 404

    # The frontend calculates percentage from current/total. We now store percentage directly.
    # We send the stored percentage as 'current_chunk' and 100 as 'total_chunks' to fit the existing UI logic.
    return jsonify({
        'status': index.indexing_status,
        'current_chunk': index.indexing_progress,
        'total_chunks': 100,
        'error_message': index.indexing_error_message
    })

def load_and_process_document(filepath, mime_type):
    """Load and process a document based on its mime type.
    
    Args:
        filepath (str): Path to the document file
        mime_type (str): MIME type of the document
        
    Returns:
        list: List of LangChain Document objects with text chunks
    """
    app.logger.info(f"Processing document: {filepath} with MIME type: {mime_type}")
    
    # Select appropriate loader based on MIME type
    try:
        if mime_type == 'application/pdf':
            loader = PyPDFLoader(filepath)
        elif mime_type == 'text/plain':
            loader = TextLoader(filepath)
        else:
            # Fall back to unstructured loader for other file types
            loader = UnstructuredFileLoader(filepath)
        
        # Load the document
        documents = loader.load()
        app.logger.info(f"Loaded {len(documents)} document parts")
        
        # Split the document into chunks
        # Use a more semantically aware splitter that prioritizes sentence boundaries
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000, 
            chunk_overlap=150, # A larger overlap helps maintain context between chunks
            separators=["\n\n", "\n", ". ", "? ", "! ", " ", ""] # Prioritize splitting on paragraphs, then sentences
        )
        text_chunks = text_splitter.split_documents(documents)
        app.logger.info(f"Split into {len(text_chunks)} chunks")
        
        return text_chunks
    except Exception as e:
        app.logger.error(f"Error processing document {filepath}: {str(e)}")
        raise

def create_or_update_rag_index(index_id: int) -> bool:
    """
    Build (or rebuild) the Chroma vector-store for a single RagIndex.
    ▸ index_id – the primary-key of the RagIndex row.
    Returns True on success, False on any fatal error (already logged).
    """
    import time
    import traceback
    start_time = time.time()
    app.logger.info("\n================= [RAG] START INDEXING =================")
    app.logger.info(f"[RAG] (re)building index_id={index_id}")
    index = RagIndex.query.get(index_id)
    if not index:
        app.logger.error(f"[RAG] ABORT: Could not find RagIndex with id={index_id}")
        return False

    try:
        # Set initial status in DB
        index.indexing_status = 'in_progress'
        index.indexing_progress = 0
        index.indexing_error_message = None
        db.session.commit()
        rag_index = index

        app.logger.info(f"[RAG] (re)building index «{rag_index.name}» (id={rag_index.id})")
        app.logger.info(f"[RAG]  - Embedding model: {rag_index.embedding_model_name}")
        app.logger.info(f"[RAG]  - Vector store path: {rag_index.vector_store_path_segment}")
        app.logger.info(f"[RAG]  - Number of documents: {len(rag_index.documents)}")

        # ────────────────────────────────────────────────── filesystem layout
        root_dir = Path(app.config["RAG_INDEX_FOLDER"])       # <-- defined at start-up
        # If this RagIndex does not yet have a folder name, create one
        if not rag_index.vector_store_path_segment:
            rag_index.vector_store_path_segment = (
                f"{secure_filename(rag_index.name.lower().replace(' ', '_'))}_"
                f"{uuid.uuid4().hex[:8]}"
            )
            db.session.commit()  # persist the new segment name

        index_path: Path = root_dir / rag_index.vector_store_path_segment
        index_path.mkdir(parents=True, exist_ok=True)

        # ─────────────────────────────────────────────── load & split docs
        all_chunks = []
        valid_documents = 0
        phantom_documents = 0
        
        # Process all documents
        for doc_idx, doc in enumerate(rag_index.documents):
            app.logger.info(f"[RAG]   → Document {doc_idx+1}/{len(rag_index.documents)}: {doc.filename} (id={doc.id})")
            try:
                doc_start = time.time()
                full_doc_filepath = Path(app.config['RAG_UPLOAD_FOLDER']) / doc.filepath
    
                # Check if the document file actually exists before trying to load it
                if not full_doc_filepath.exists():
                    app.logger.error(f"[RAG] Document file not found: {str(full_doc_filepath)} for doc_id={doc.id}. Skipping.")
                    phantom_documents += 1
                    continue
            
                valid_documents += 1
                chunks = load_and_process_document(str(full_doc_filepath), doc.mime_type)
                all_chunks.extend(chunks)
                app.logger.info(f"[RAG]     - Split into {len(chunks)} chunks in {time.time()-doc_start:.2f}s")
                
                # Debug logging for chunks
                for chunk_idx, chunk in enumerate(chunks):
                    # Try to extract text for preview
                    if isinstance(chunk, dict) and 'text' in chunk:
                        chunk_text = chunk['text']
                    elif hasattr(chunk, 'text'):
                        chunk_text = chunk.text
                    elif isinstance(chunk, str):
                        chunk_text = chunk
                    else:
                        chunk_text = str(chunk)
                    
                    chunk_preview = chunk_text[:80].replace('\n', ' ')
                    app.logger.debug(f"[RAG]       Chunk {chunk_idx+1}/{len(chunks)}: {chunk_preview}{'...' if len(chunk_text)>80 else ''}")
            except Exception as e:
                app.logger.error(f"[RAG]     FAILED to process {doc.filename}: {e}")
                app.logger.error(traceback.format_exc())

        # Check if we found any valid documents
        if valid_documents == 0:
            if phantom_documents > 0:
                app.logger.warning(f"[RAG] All {phantom_documents} documents were phantom (missing files) - index creation skipped")
                app.logger.info("================= [RAG] END INDEXING (ALL PHANTOM DOCUMENTS) =================\n")
                # Return true since this is not a fatal error, just a state where no action is needed
                return True
            else:
                app.logger.warning("[RAG] No documents in index - nothing to do")
                app.logger.info("================= [RAG] END INDEXING (NO DOCUMENTS) =================\n")
                return True

        # If we have valid documents but no chunks were produced
        if not all_chunks:
            app.logger.warning("[RAG] No chunks produced from valid documents – aborting build")
            app.logger.info("================= [RAG] END INDEXING (NO CHUNKS) =================\n")
            return False

        # ──────────────────────────────────────────────── connect to Chroma
        # The ChromaDB client is failing due to deprecated environment variables.
        # This fix unsets those variables before initializing the client to ensure
        # the modern configuration is used.
        app.logger.info(f"[RAG] Initializing ChromaDB client for index path: {index_path}")
        client = get_chroma_client(str(index_path))

        # ───────────────────────────────────────────────── EMBEDDING FUNCTION
        app.logger.info(f"[RAG] Initializing embedding function for model: {rag_index.embedding_model_name}")
        langchain_ef = None
        try:
            chroma_native_ef = get_embedding_function(
                embedding_model_name=rag_index.embedding_model_name,
                ollama_host=app.config['OLLAMA_HOST'],
                hf_token=os.getenv('HF_TOKEN')
            )
            if rag_index.embedding_model_name.startswith('ollama/'):
                ollama_model_name = rag_index.embedding_model_name.split('ollama/')[1]
                langchain_ef = OllamaEmbeddings(model=ollama_model_name, base_url=app.config['OLLAMA_HOST'])
            elif rag_index.embedding_model_name.startswith('huggingface/'):
                hf_model_name = rag_index.embedding_model_name.split('huggingface/')[1]
                langchain_ef = HuggingFaceEmbeddings(model_name=hf_model_name, model_kwargs={'device': 'cpu'}, encode_kwargs={'normalize_embeddings': True})
        except ValueError as e:
            app.logger.error(f"[RAG] Error initializing embedding function: {e}")
            raise

        if not langchain_ef:
            app.logger.error("[RAG] Failed to initialize Langchain embedding function.")
            raise RuntimeError("Failed to initialize Langchain embedding function.")

        # ───────────────────────────────────────────────── CREATE COLLECTION
        # always recreate the collection for a clean rebuild
        try:
            client.delete_collection(rag_index.name)
            app.logger.info(f"[RAG] Existing Chroma collection '{rag_index.name}' deleted for rebuild.")
        except Exception as e:
            app.logger.info(f"[RAG] No existing Chroma collection to delete: {e}")
        
        collection = client.create_collection(
            name=rag_index.name,
            embedding_function=chroma_native_ef
        )
        app.logger.info(f"[RAG] Chroma collection «{rag_index.name}» created with its embedding function.")

        # ───────────────────────────────────────────────── ADD DOCUMENTS
        vecstore = Chroma(
            client=client,
            collection_name=rag_index.name,
            embedding_function=langchain_ef,
        )
        import psutil
        app.logger.info(f"[RAG] Adding {len(all_chunks)} chunks to vector store (embedding and persisting, PER-CHUNK DEBUG MODE)...")
        add_start = time.time()
        success_count = 0
        fail_count = 0
        for idx, chunk in enumerate(all_chunks):
            chunk_start = time.time()
            try:
                # Extract chunk text and metadata for logging
                if isinstance(chunk, dict) and 'text' in chunk:
                    chunk_text = chunk['text']
                elif hasattr(chunk, 'text'):
                    chunk_text = chunk.text
                elif isinstance(chunk, str):
                    chunk_text = chunk
                else:
                    chunk_text = str(chunk)
                chunk_preview = chunk_text[:120].replace('\n', ' ')
                chunk_len = len(chunk_text)
                chunk_meta = getattr(chunk, 'metadata', getattr(chunk, 'meta', {}))
                app.logger.info(f"[RAG] Embedding chunk {idx+1}/{len(all_chunks)} (len={chunk_len}): '{chunk_preview}{'...' if chunk_len>120 else ''}' | meta: {chunk_meta}")
                # Log memory usage before embedding
                try:
                    process = psutil.Process()
                    mem_info = process.memory_info()
                    app.logger.debug(f"[RAG] [Mem] RSS={mem_info.rss/1024/1024:.2f}MB, VMS={mem_info.vms/1024/1024:.2f}MB, Chunks so far: {success_count}")
                except Exception as e_mem:
                    app.logger.debug(f"[RAG] [Mem] Unable to fetch memory info: {e_mem}")
                # Actually embed and persist this chunk
                vecstore.add_documents([chunk])
                success_count += 1
                elapsed = time.time() - chunk_start
                app.logger.info(f"[RAG]   → Successfully embedded chunk {idx+1}/{len(all_chunks)} in {elapsed:.2f}s (total success: {success_count})")
                # Update progress in DB
                index.indexing_progress = int(((idx + 1) / len(all_chunks)) * 100)
                db.session.commit()
            except Exception as e:
                fail_count += 1
                app.logger.error(f"[RAG]   ✗ Error embedding chunk {idx+1}/{len(all_chunks)}: {e}", exc_info=True)
                app.logger.error(f"[RAG]   ✗ Chunk content: {repr(chunk)[:300]}")
                # Update error in DB
                index.indexing_status = 'failed'
                index.indexing_error_message = str(e)
                db.session.commit()
        total_time = time.time() - add_start
        app.logger.info(f"[RAG] Finished embedding loop: {success_count} succeeded, {fail_count} failed, total time: {total_time:.2f}s")
        if fail_count == 0:
            # Final success update
            index.indexing_status = 'success'
            index.indexing_progress = 100
            db.session.commit()
        else:
            # Final failure update
            index.indexing_status = 'failed'
            index.indexing_error_message = f"Failed to embed {fail_count} out of {len(all_chunks)} chunks."
            db.session.commit()


        app.logger.info(f"[RAG] ✅ finished index «{rag_index.name}» "
                        f"({len(all_chunks)} chunks) in {time.time()-start_time:.2f}s")
        app.logger.info("================= [RAG] END INDEXING (SUCCESS) =================\n")
        return True

    except Exception as e:
        name = index.name if 'index' in locals() and index else 'unknown'
        app.logger.error(f"[RAG] Fatal while building «{name}»: {e}")
        app.logger.error(traceback.format_exc())
        # Final failure update in DB
        if 'index' in locals() and index:
            index.indexing_status = 'failed'
            index.indexing_error_message = f"A fatal error occurred during indexing: {str(e)}"
            db.session.commit()
        app.logger.info("================= [RAG] END INDEXING (ERROR) =================\n")
        return False

# Define application pages/endpoints that can have access controlled
MANAGED_PAGE_ENDPOINTS = [
    {'endpoint': 'index', 'display_name': 'Home Page', 'description': 'The main landing page of the application.'},
    {'endpoint': 'chat', 'display_name': 'Chat Interface', 'description': 'The primary chat functionality.'},
    {'endpoint': 'admin_rbac_page', 'display_name': 'Access Control', 'description': 'This admin page for managing roles, users, and permissions.'},
    # {'endpoint': 'profile', 'display_name': 'User Profile', 'description': 'User profile viewing and editing page.'}, # Example for future
    # {'endpoint': 'settings', 'display_name': 'User Settings', 'description': 'User-specific application settings.'} # Example for future
]



# Ensure static directories exist
if not os.path.exists('static'):
    os.makedirs('static')
if not os.path.exists('static/css'): # Corrected line
    os.makedirs('static/css')
if not os.path.exists('static/js'):
    os.makedirs('static/js')
if not os.path.exists('static/uploads'):
    os.makedirs('static/uploads')

# Define and create RAG_INDEX_FOLDER
RAG_INDEX_FOLDER = Path(app.root_path) / 'rag_indexes'
if not os.path.exists(RAG_INDEX_FOLDER):
    os.makedirs(RAG_INDEX_FOLDER)

app.config['RAG_INDEX_FOLDER'] = str(RAG_INDEX_FOLDER)   # make it visible everywhere


app.secret_key = os.environ.get('SECRET_KEY', 'you-will-never-guess')
csrf = CSRFProtect(app)

# Retrieve individual MySQL settings from .env
mysql_user = os.environ.get('MYSQL_USER')
mysql_password = os.environ.get('MYSQL_PASSWORD')
mysql_database = os.environ.get('MYSQL_DATABASE')
# Use the Docker service name 'mysql' or the env var if set
mysql_host = os.environ.get('MYSQL_HOST', 'mysql')
mysql_port = os.environ.get('MYSQL_PORT', '3306') # Internal Docker port

# Build the SQLAlchemy connection string dynamically.
app.logger.info(f"Attempting to connect to MySQL with host: '{mysql_host}' and port: '{mysql_port}'")
connection_str = f"mysql+pymysql://{mysql_user}:{mysql_password}@{mysql_host}:{mysql_port}/{mysql_database}"
app.config['SQLALCHEMY_DATABASE_URI'] = connection_str
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_POOL_RECYCLE'] = 280 # Recycle connections to prevent timeouts

# Initialize SQLAlchemy with the app instance
db.init_app(app)
migrate = Migrate(app, db, directory='backend/db/db/migrations', compare_type=True)

# Configure Flask-Mail (read SMTP settings from environment)
app.config['MAIL_SERVER'] = os.environ.get('MAIL_SERVER', 'smtp.example.com')
app.config['MAIL_PORT'] = int(os.environ.get('MAIL_PORT', 587))
app.config['MAIL_USE_TLS'] = os.environ.get('MAIL_USE_TLS', 'true').lower() in ['true', '1']
app.config['MAIL_USERNAME'] = os.environ.get('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.environ.get('MAIL_PASSWORD')
app.config['MAIL_DEFAULT_SENDER'] = os.environ.get('MAIL_DEFAULT_SENDER')
# Define RAG_UPLOAD_FOLDER early to be used in both places
RAG_UPLOAD_FOLDER = Path(app.root_path) / 'rag_documents'
app.config['RAG_UPLOAD_FOLDER'] = RAG_UPLOAD_FOLDER
app.config['RAG_DOCUMENTS_DIR'] = str(RAG_UPLOAD_FOLDER)  # Use the same path for consistency

# Add Ollama configuration
app.config['OLLAMA_HOST'] = os.environ.get('OLLAMA_HOST', 'http://localhost:11434')
mail = Mail(app)

# Configure Flask-Login
login_manager = LoginManager(app)
login_manager.login_view = 'index'


# === Initialize LLM Service and Manage Models ===
llm_service_type = os.environ.get('LLM_SERVICE', 'ollama').lower()
app.logger.info(f"Initializing LLM service of type: {llm_service_type}")
llm_service = LLMServiceFactory.create_service() # This might call test_connection

# Store successfully managed models and the effective default model in app.config
app.config['MANAGED_OLLAMA_MODELS'] = []
app.config['EFFECTIVE_DEFAULT_MODEL_NAME'] = None
DEFAULT_MODEL_NAME = None # Will be determined by the logic below

# Check if database tables exist and set up error handling
import sqlalchemy.exc

def is_database_initialized():
    """Check if essential database tables exist"""
    try:
        # Try to query a table that should exist if the database is initialized
        db.session.execute(db.select(User).limit(1))
        return True
    except sqlalchemy.exc.ProgrammingError as e:
        app.logger.warning(f"Database tables not fully initialized: {str(e)}")
        return False
    except Exception as e:
        app.logger.error(f"Error checking database status: {str(e)}")
        return False

# Check database status
with app.app_context():
    app.config['DATABASE_INITIALIZED'] = is_database_initialized()

if llm_service_type == 'ollama':
    if llm_service and hasattr(llm_service, 'test_connection') and llm_service.test_connection(): # Ensure service is responsive
        ollama_models_env = os.environ.get('OLLAMA_MODELS')
        if ollama_models_env:
            managed_model_names_from_env = [name.strip() for name in ollama_models_env.split(',') if name.strip()]
            app.logger.info(f"Target Ollama models from OLLAMA_MODELS env var: {managed_model_names_from_env}")

            current_ollama_server_models = llm_service.list_models() # Get models currently on server
            app.logger.info(f"Models currently on Ollama server: {current_ollama_server_models}")

            successfully_managed_models_list = []
            # Check if we're likely in a Docker container
            is_docker = os.path.exists('/.dockerenv') or os.environ.get('RUNNING_IN_DOCKER') == 'true'
            if is_docker:
                app.logger.info("Running in Docker environment - will handle DNS resolution errors")
                
            for model_name in managed_model_names_from_env:
                is_on_server = False
                for server_model in current_ollama_server_models:
                    if server_model == model_name or server_model.startswith(model_name + ":"):
                        is_on_server = True
                        break
                
                if is_on_server:
                    app.logger.info(f"Model '{model_name}' is already available on Ollama server.")
                    successfully_managed_models_list.append(model_name)
                else:
                    app.logger.info(f"Model '{model_name}' not found on Ollama server. Attempting to pull...")
                    try:
                        if llm_service.pull_model(model_name): # pull_model returns True on success
                            app.logger.info(f"Successfully pulled model '{model_name}'.") 
                            successfully_managed_models_list.append(model_name)
                        else:
                            app.logger.warning(f"Failed to pull model '{model_name}'. It will not be available through this managed list.")
                    except Exception as e:
                        error_msg = str(e).lower()
                        # Check for DNS-related errors which are common in Docker environments
                        if is_docker and ('name resolution' in error_msg or 'no such host' in error_msg or 'dns' in error_msg):
                            app.logger.error(f"Docker DNS resolution error while pulling model '{model_name}': {str(e)}")
                            app.logger.warning(f"Adding '{model_name}' to managed models list despite pull failure (DNS issue)")
                            # In Docker with DNS issues, still add the model to the list so the app can function
                            # This allows the app to work when network connectivity is restored without requiring restart
                            successfully_managed_models_list.append(model_name)
                        else:
                            app.logger.error(f"Error pulling model '{model_name}': {str(e)}")
                            app.logger.warning(f"Model '{model_name}' will not be available.")
                            # For non-DNS errors, don't add the model to the managed list
            
            app.config['MANAGED_OLLAMA_MODELS'] = successfully_managed_models_list
            app.logger.info(f"Successfully managed Ollama models (available/pulled): {app.config['MANAGED_OLLAMA_MODELS']}")
        else:
            app.logger.warning("OLLAMA_MODELS environment variable not set. No specific Ollama models will be pre-managed/pulled.")
            # If OLLAMA_MODELS is not set, MANAGED_OLLAMA_MODELS remains empty.
            # initialize_rbac_data will then not add any models from this list.

        # Determine DEFAULT_MODEL_NAME for Ollama
        env_default_model_from_env = os.environ.get('DEFAULT_MODEL_NAME')
        effective_default = None

        if env_default_model_from_env:
            if env_default_model_from_env in app.config['MANAGED_OLLAMA_MODELS']:
                effective_default = env_default_model_from_env
            else:
                app.logger.warning(f"DEFAULT_MODEL_NAME '{env_default_model_from_env}' from .env is not in the list of successfully managed models ({app.config['MANAGED_OLLAMA_MODELS']}).")
                if app.config['MANAGED_OLLAMA_MODELS']:
                    effective_default = app.config['MANAGED_OLLAMA_MODELS'][0]
                    app.logger.warning(f"Falling back to the first managed model as default: '{effective_default}'.")
                else:
                    app.logger.error("No managed Ollama models available. Cannot set a default model from managed list.")
        elif app.config['MANAGED_OLLAMA_MODELS']:
            effective_default = app.config['MANAGED_OLLAMA_MODELS'][0]
            app.logger.info(f"DEFAULT_MODEL_NAME not set in .env. Using the first managed model as default: '{effective_default}'.")
        else:
            app.logger.warning("DEFAULT_MODEL_NAME not set in .env and no managed Ollama models available. Default model not set from managed list.")
        
        app.config['EFFECTIVE_DEFAULT_MODEL_NAME'] = effective_default
        DEFAULT_MODEL_NAME = effective_default 
    else:
        app.logger.error("Ollama service is not available or failed connection test. Cannot manage Ollama models or set default.")
        app.config['MANAGED_OLLAMA_MODELS'] = [] # Ensure it's empty
        app.config['EFFECTIVE_DEFAULT_MODEL_NAME'] = None
        DEFAULT_MODEL_NAME = None

elif llm_service_type == 'llamacpp':
    DEFAULT_MODEL_NAME = os.environ.get('LLAMACPP_MODEL', "llama-2-7b-chat.Q4_K_M.gguf")
    app.config['EFFECTIVE_DEFAULT_MODEL_NAME'] = DEFAULT_MODEL_NAME
    # LlamaCPP models are not currently managed via OLLAMA_MODELS env var
    app.config['MANAGED_OLLAMA_MODELS'] = [] # Ensure this is empty for non-ollama services
else:
    app.logger.error(f"Unsupported LLM_SERVICE_TYPE '{llm_service_type}'. No default model configured through this logic.")
    DEFAULT_MODEL_NAME = None
    app.config['EFFECTIVE_DEFAULT_MODEL_NAME'] = None
    app.config['MANAGED_OLLAMA_MODELS'] = []

if DEFAULT_MODEL_NAME:
    app.logger.info(f"Global DEFAULT_MODEL_NAME set to: {DEFAULT_MODEL_NAME}")
else:
    app.logger.warning("Global DEFAULT_MODEL_NAME could not be determined based on configuration and service availability.")

# Add a template filter for converting newlines to <br> tags
@app.template_filter('nl2br')
def nl2br(value):
    # First trim the string to remove leading/trailing whitespace
    if not value:
        return value
    value = value.strip()
    # Replace newlines with <br> tags
    value = value.replace('\n', '<br>')
    # Remove any double <br> tags that might cause excessive spacing
    while '<br><br><br>' in value:
        value = value.replace('<br><br><br>', '<br><br>')
    return value


# Print all registered routes for debugging (always runs)
print("\n=== Registered Flask Routes ===")
try:
    for rule in app.url_map.iter_rules():
        print(f"{rule.endpoint}: {rule}")
except Exception as e:
    print(f"Error printing routes: {e}")
print("==============================\n")

# ===========================
# Association Tables for RBAC
# ===========================
# Association table for User and Role (many-to-many)


@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))

@app.route('/', methods=['GET', 'POST'])
def index():
    if current_user.is_authenticated:
        return redirect(url_for('chat'))

    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        user = User.query.filter_by(email=email).first()

        if not user or not user.check_password(password):
            flash("Invalid email or password.", 'danger')
            return redirect(url_for('index'))

        if not user.confirmed:
            flash("Please confirm your email before logging in.", 'warning')
            return redirect(url_for('index'))
        
        if not user.is_active:
            flash("Your account is inactive. Please contact your administrator.", 'danger')
            return redirect(url_for('index'))

        login_user(user)
        next_page = request.args.get('next')
        if next_page and next_page.startswith('/'):
            return redirect(next_page)
        else:
            return redirect(url_for('chat'))
            
    return render_template('index.html')

@app.route('/ping', methods=['GET', 'HEAD'])
def ping():
    """Simple health check endpoint"""
    try:
        # Check database connection
        db.session.execute('SELECT 1')
        response = "pong"
        return response, 200, {'Content-Type': 'text/plain', 'Content-Length': str(len(response))}
    except Exception as e:
        app.logger.error(f"Health check failed: {str(e)}")
        response = "database connection error"
        return response, 500, {'Content-Type': 'text/plain', 'Content-Length': str(len(response))}

# ===========================
# Ollama model listing at startup
# ===========================
def load_and_ensure_llm_models():
    """Loads model list, ensures default (and specified) models are pulled if missing."""
    with app.app_context(): # Ensure all operations run within app context
        # Use the centrally managed list of Ollama models from app.config
        managed_ollama_models = app.config.get('MANAGED_OLLAMA_MODELS', [])
        app.logger.info(f"LOAD_AND_ENSURE: Using MANAGED_OLLAMA_MODELS from app.config: {managed_ollama_models}")

        models_to_ensure = set(managed_ollama_models) # Start with all managed models
        
        # DEFAULT_MODEL_NAME should already be in managed_ollama_models if valid,
        # but adding it here ensures it's considered if somehow missed or if it's a LlamaCPP model not in OLLAMA_MODELS.
        if DEFAULT_MODEL_NAME: 
            models_to_ensure.add(DEFAULT_MODEL_NAME)
        
        app.logger.info(f"LOAD_AND_ENSURE: Final models_to_ensure (after adding default if needed): {list(models_to_ensure)}")

        available_models = []
        try:
            all_active_db_models = Model.query.filter_by(is_active=True).all()
            active_db_model_names = {model.ollama_model_name for model in all_active_db_models}
            app.logger.info(f"Active models from DB: {active_db_model_names}")

            if llm_service_type == 'ollama':
                try:
                    initial_server_models_list = llm_service.list_models()
                    initial_server_models_set = set(initial_server_models_list)
                    app.logger.info(f"LOAD_AND_ENSURE: Initial models on Ollama server: {initial_server_models_set}")

                    models_pulled_this_run = set()
                    # Ensure managed models are pulled if not on server
                    for model_name_to_pull in managed_ollama_models: # managed_ollama_models is from app.config
                        if model_name_to_pull not in initial_server_models_set:
                            try:
                                app.logger.info(f"LOAD_AND_ENSURE: Model '{model_name_to_pull}' (managed) not on server. Attempting to pull...")
                                llm_service.pull_model(model_name_to_pull)
                                app.logger.info(f"LOAD_AND_ENSURE: Successfully pulled model '{model_name_to_pull}'.")
                                models_pulled_this_run.add(model_name_to_pull)
                            except Exception as e:
                                app.logger.error(f"LOAD_AND_ENSURE: Failed to pull managed model '{model_name_to_pull}': {e}")
                    
                    final_server_models_set = initial_server_models_set.union(models_pulled_this_run)
                    app.logger.info(f"LOAD_AND_ENSURE: Final models on Ollama server (after pulls): {final_server_models_set}")

                    # Populate available_models from all models now on the server that are also active in DB
                    for server_model_name in final_server_models_set:
                        if server_model_name in active_db_model_names:
                            db_model_obj = next((m for m in all_active_db_models if m.ollama_model_name == server_model_name), None)
                            if db_model_obj: # Should be true if server_model_name in active_db_model_names
                                available_models.append({
                                    'id': db_model_obj.id,
                                    'name': db_model_obj.display_name or db_model_obj.ollama_model_name,
                                    'ollama_model_name': db_model_obj.ollama_model_name,
                                    'is_default': db_model_obj.ollama_model_name == DEFAULT_MODEL_NAME
                                })
                except RequestError as re:
                    app.logger.error(f"Ollama RequestError when ensuring models: {re}. This might happen if Ollama is not running or not reachable.")
                    # Fallback: only use models already in DB if Ollama is down and they were in models_to_ensure
                    for db_model in all_active_db_models:
                        if db_model.ollama_model_name in models_to_ensure:
                            available_models.append({
                                'id': db_model.id,
                                'name': db_model.display_name or db_model.ollama_model_name,
                                'ollama_model_name': db_model.ollama_model_name,
                                'is_default': db_model.ollama_model_name == DEFAULT_MODEL_NAME
                            })

            elif llm_service_type == 'llamacpp':
                if DEFAULT_MODEL_NAME in active_db_model_names:
                    db_model_obj = next((m for m in all_active_db_models if m.ollama_model_name == DEFAULT_MODEL_NAME), None)
                    if db_model_obj:
                        available_models.append({
                            'id': db_model_obj.id,
                            'name': db_model_obj.display_name, # Corrected to display_name
                            'ollama_model_name': db_model_obj.ollama_model_name,
                            'is_default': True
                        })
                else:
                    app.logger.warning(f"LlamaCPP model '{DEFAULT_MODEL_NAME}' is not marked active in the database.")

            # Consolidate fallback for empty available_models
            if not available_models and DEFAULT_MODEL_NAME and DEFAULT_MODEL_NAME in active_db_model_names:
                app.logger.info(f"No specific models made it to available_models list, but default '{DEFAULT_MODEL_NAME}' is active. Adding it.")
                db_model_obj = next((m for m in all_active_db_models if m.ollama_model_name == DEFAULT_MODEL_NAME), None)
                if db_model_obj:
                    available_models.append({
                        'id': db_model_obj.id,
                        'name': db_model_obj.display_name, # Corrected to display_name
                        'ollama_model_name': db_model_obj.ollama_model_name,
                        'is_default': True
                    })
            
            if DEFAULT_MODEL_NAME:
                available_models.sort(key=lambda x: x['ollama_model_name'] != DEFAULT_MODEL_NAME)

            return available_models

        except Exception as e:
            app.logger.exception(f"Error during model loading and pulling process (within app_context): {e}")
            # Fallback logic if an error occurs even within the app_context
            if DEFAULT_MODEL_NAME:
                # Check if DEFAULT_MODEL_NAME exists in the database as a last resort
                try:
                    default_db_model = Model.query.filter_by(ollama_model_name=DEFAULT_MODEL_NAME, is_active=True).first()
                    if default_db_model:
                        app.logger.warning(f"Proceeding with fallback default model for UI due to error: {default_db_model.ollama_model_name} (from DB)")
                        available_models.append({
                            'id': default_db_model.id,
                            'name': default_db_model.display_name,  
                            'ollama_model_name': default_db_model.ollama_model_name, 
                            'is_default': True,
                            'description': default_db_model.description
                        })
                except Exception as db_e:
                    app.logger.error(f"Could not even fetch default model from DB during fallback: {db_e}")
                
                app.logger.warning(f"Proceeding with fallback default model name (string only) for UI due to error: {DEFAULT_MODEL_NAME}")
                # Return a structure consistent with what the chat route expects if possible, even if it's just the name
                return [{'name': DEFAULT_MODEL_NAME, 'ollama_model_name': DEFAULT_MODEL_NAME, 'is_default': True, 'id': None}]
            return []

# Only initialize models if we are not running a DB migration command
if 'db' not in sys.argv and 'migrations' not in sys.argv:
    llm_models = load_and_ensure_llm_models()
    app.config['LLM_MODELS'] = llm_models if llm_models else ([DEFAULT_MODEL_NAME] if DEFAULT_MODEL_NAME else [])
    app.logger.info(f"Using models for dropdown: {app.config['LLM_MODELS']}")
else:
    app.logger.info("Skipping LLM model loading during DB migration to prevent race conditions.")
    app.config['LLM_MODELS'] = []

# ===========================
# Audio Processing Placeholders (functionality removed)
# ===========================

def check_whisper_model_exists(model_name="base"):
    """Placeholder function - Whisper functionality has been removed"""
    app.logger.warning("Whisper functionality has been removed from this version")
    return False

def check_ffmpeg_installed():
    """Placeholder function - FFmpeg check"""
    app.logger.warning("FFmpeg check - audio processing functionality has been removed")
    return False

def recognize_audio(file_path, language=None):
    """Placeholder function - Audio recognition has been removed"""
    app.logger.warning("Audio recognition functionality has been removed")
    return "Audio recognition is not available in this version"

def convert_audio_format(input_path):
    """Placeholder function - Audio conversion has been removed"""
    app.logger.warning("Audio conversion functionality has been removed")
    return None

def detect_language(audio_file_path):
    """Placeholder function - Language detection has been removed"""
    app.logger.warning("Language detection functionality has been removed")
    return 'en'  # Default to English

# ===========================
# Routes for Authentication
# ===========================
    
@app.route('/confirm/<token>')
def confirm_email(token):
    try:
        user_id = int(token.split('-')[0])
    except Exception:
        flash("Invalid confirmation token.")
        return redirect(url_for('index'))
    user = db.session.get(User, user_id)
    if user:
        user.confirmed = True
        db.session.commit()
        flash("Your account has been confirmed!")
    return redirect(url_for('index'))



@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('You have been logged out.', 'info')
    return redirect(url_for('index'))
    
@app.route('/reset', methods=['GET', 'POST'])
def reset_request():
    if request.method == 'POST':
        email = request.form['email']
        user = User.query.filter_by(email=email).first()
        if user:
            token = f"{user.id}-reset-token"  # Replace with secure token generation.
            reset_url = url_for('reset_password', token=token, _external=True)
            msg = Message("Password Reset", recipients=[email])
            msg.body = f"Reset your password by clicking on the link: {reset_url}"
            mail.send(msg)
            flash("Password reset email sent.")
        else:
            flash("Email not found.")
    return render_template('reset_password_coming_soon.html')

@app.route('/reset_password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    try:
        user_id = int(token.split('-')[0])
    except Exception:
        flash("Invalid reset token.")
        return redirect(url_for('reset_request'))
    user = db.session.get(User, user_id)
    if request.method == 'POST':
        new_password = request.form['password']
        user.set_password(new_password)
        db.session.commit()
        flash("Your password has been updated.")
        return redirect(url_for('login'))
    return render_template('reset_password.html')

# ===========================
# Routes for Chat and Conversations
# ===========================
# =================================================================
# ================== Helper Functions =============================
# =================================================================

def get_available_models_for_user(user):
    """
    Fetches all models (base and RAG) from the database that the given user has permission to access.
    Returns a list of model dictionaries.
    """
    if not user.is_authenticated:
        return []
        
    # If database is not initialized yet, use models from app.config
    if not app.config.get('DATABASE_INITIALIZED', False):
        app.logger.warning("Database not fully initialized, using managed models from config")
        managed_models = app.config.get('MANAGED_OLLAMA_MODELS', [])
        # Return all managed models without permissions check since DB isn't ready
        return [{
            'id': idx + 1,  # Temporary IDs
            'name': model_name,
            'ollama_model_name': model_name,
            'description': f"Ollama model: {model_name}"
        } for idx, model_name in enumerate(managed_models)]
    
    # Normal flow when database is available
    try:
        # 1. Get base model names from config (from Ollama)
        llm_models_config = app.config.get('LLM_MODELS', [])
        processed_ollama_names = []
        if isinstance(llm_models_config, list):
            for item in llm_models_config:
                if isinstance(item, str):
                    processed_ollama_names.append(item)
                elif isinstance(item, dict):
                    model_name = item.get('ollama_model_name')
                    if model_name and isinstance(model_name, str):
                        processed_ollama_names.append(model_name)

        # 2. Fetch corresponding models from DB
        # Base models must be in the list from Ollama
        base_models = Model.query.filter(
            Model.is_active==True,
            Model.is_rag_model==False,
            Model.ollama_model_name.in_(processed_ollama_names)
        ).all()
        
        # RAG models are just fetched if active
        rag_models = Model.query.filter(
            Model.is_active==True,
            Model.is_rag_model==True
        ).all()
        
        db_models_to_check = base_models + rag_models
        
        # 3. Filter by user permissions
        accessible_models = []
        for model_obj in db_models_to_check:
            try:
                if user.can_access_model(model_obj.id):
                    accessible_models.append({
                        'id': model_obj.id, 
                        'name': model_obj.display_name, 
                        'ollama_model_name': model_obj.ollama_model_name, 
                        'description': model_obj.description
                    })
            except Exception as e:
                app.logger.warning(f"Error checking model access for model {model_obj.id}: {str(e)}")
        app.logger.debug(f"get_available_models_for_user: Found {len(accessible_models)} models for user {user.id}")
        return accessible_models
    except Exception as e:
        app.logger.error(f"Error retrieving models: {str(e)}")
        # Fallback to managed models from config
        managed_models = app.config.get('MANAGED_OLLAMA_MODELS', [])
        return [{
            'id': idx + 1,  # Temporary IDs
            'name': model_name,
            'ollama_model_name': model_name,
            'description': f"Ollama model: {model_name}"
        } for idx, model_name in enumerate(managed_models)]

def get_default_model():
    """
    Retrieves the default model from the database.
    Returns a model dictionary or None.
    """
    default_model_name = app.config.get('EFFECTIVE_DEFAULT_MODEL_NAME')
    if not default_model_name:
        app.logger.warning("get_default_model: No EFFECTIVE_DEFAULT_MODEL_NAME configured.")
        return None
        
    default_db_model = Model.query.filter_by(ollama_model_name=default_model_name, is_active=True).first()
    if default_db_model:
        app.logger.info(f"get_default_model: Found default model '{default_model_name}' in DB.")
        return {
            'id': default_db_model.id,
            'name': default_db_model.display_name,
            'ollama_model_name': default_db_model.ollama_model_name,
            'description': default_db_model.description
        }
    app.logger.warning(f"get_default_model: Default model '{default_model_name}' not found or not active in DB.")
    return None


@app.route('/chat', methods=['GET', 'POST'])
@login_required
def chat():
    if not current_user.can_access_page('chat'):
        flash('You do not have permission to access this page.', 'danger')
        return redirect(url_for('index'))
    # Check if user has a preferred LLM service
    user_service = session.get('user_llm_service')
    global llm_service
    global llm_service_type
    
    # If user has a preference that's different from current service, switch
    if user_service and user_service != llm_service_type:
        try:
            app.logger.info(f"Switching to user's preferred LLM service: {user_service}")
            
            # Try to create the service and test it
            temp_service = LLMServiceFactory.create_service_by_type(user_service)
            available_models = temp_service.list_models()
            
            # If we get here, the service is working
            llm_service = temp_service
            llm_service_type = user_service
            
            # Update the application's cached models list
            app.config['LLM_MODELS'] = available_models if available_models else [DEFAULT_MODEL_NAME]
            app.logger.info(f"Switched to user's preferred service {user_service} with models: {app.config['LLM_MODELS']}")
        except Exception as e:
            app.logger.exception(f"Failed to switch to user's preferred service {user_service}: {e}")
            # If we can't use the preferred service, clear the preference
            session.pop('user_llm_service', None)
    
    # Continue with the existing chat route logic
    conversation_id = request.args.get('conversation_id', None)
    # Use the new helper function to get all models this user can access.
    available_models = get_available_models_for_user(current_user)
    app.logger.debug(f"Chat Route: get_available_models_for_user returned {len(available_models)} models.")

    # Fallback to default if user has no models assigned and a default exists.
    if not available_models:
        default_model = get_default_model()
        if default_model and current_user.can_access_model(default_model['id']):
            available_models = [default_model]
            app.logger.info(f"User {current_user.id} has no specific models, falling back to accessible default: {default_model['ollama_model_name']}")
        else:
            app.logger.warning(f"User {current_user.id} has no accessible models, including the default. Model selection will be empty.") 

    # Create a set of ollama_model_names that are currently available and accessible to the user
    accessible_ollama_model_names = {model_dict['ollama_model_name'] for model_dict in available_models if 'ollama_model_name' in model_dict}
    app.logger.debug(f"Chat Route: Accessible ollama_model_names (set, after fallback logic): {accessible_ollama_model_names}")
    app.logger.debug(f"Chat Route: Configured EFFECTIVE_DEFAULT_MODEL_NAME: {app.config.get('EFFECTIVE_DEFAULT_MODEL_NAME')}")

    if conversation_id:
        conversation = Conversation.query.filter_by(id=conversation_id, user_id=current_user.id).first_or_404()
        
        model_needs_fallback = False
        if not conversation.selected_model: # No model selected yet for this existing conversation
            model_needs_fallback = True
            app.logger.info(f"Conversation {conversation.id} has no model selected.")
        elif conversation.selected_model not in accessible_ollama_model_names: # Selected model is no longer available/accessible
            model_needs_fallback = True
            app.logger.warning(f"Conversation {conversation.id} had model '{conversation.selected_model}' which is not available or accessible.")

        if model_needs_fallback:
            if available_models: # Check if there are any models to fall back to
                fallback_model_dict = available_models[0]
                new_model_name = fallback_model_dict.get('ollama_model_name')
                conversation.selected_model = new_model_name
                app.logger.info(f"Falling back/setting model for conversation {conversation.id} to '{new_model_name}'.")
                db.session.add(conversation) # Mark for update
            else: # No models available to fall back to
                conversation.selected_model = None
                app.logger.warning(f"No accessible models to fall back to for conversation {conversation.id}. Setting selected_model to None.")
                db.session.add(conversation) # Mark for update
    else: # No conversation_id, so it's a new session or fetching the latest conversation
        conversation = Conversation.query.filter_by(user_id=current_user.id).order_by(Conversation.created_at.desc()).first()
        if conversation: # Only process if there's an existing conversation
            model_needs_fallback_for_latest = False
            if not conversation.selected_model:
                model_needs_fallback_for_latest = True
                app.logger.info(f"Latest conversation {conversation.id} has no model selected.")
            elif conversation.selected_model not in accessible_ollama_model_names:
                model_needs_fallback_for_latest = True
                app.logger.warning(f"Latest conversation {conversation.id} had model '{conversation.selected_model}' which is not available or accessible.")

            if model_needs_fallback_for_latest:
                if available_models:
                    fallback_model_dict = available_models[0]
                    new_model_name = fallback_model_dict.get('ollama_model_name')
                    conversation.selected_model = new_model_name
                    app.logger.info(f"Falling back/setting model for latest conversation {conversation.id} to '{new_model_name}'.")
                    db.session.add(conversation)
                else:
                    conversation.selected_model = None
                    app.logger.warning(f"No accessible models to fall back to for latest conversation {conversation.id}. Setting selected_model to None.")
                    db.session.add(conversation)


    all_conversations = Conversation.query.filter_by(user_id=current_user.id).order_by(Conversation.created_at.desc()).all()
    messages = []
    if conversation:
        messages = ChatMessage.query.filter_by(conversation_id=conversation.id).order_by(ChatMessage.created_at).all()

    app.logger.debug(f"Chat Route: Final 'available_models' for template: {available_models}")
    app.logger.debug(f"Chat Route: Final 'conversation.selected_model' for UI: {conversation.selected_model if conversation else 'No conversation object'}")
    
    # Pass the LLM service type to the template
    cache_id = int(time.time())
    return render_template('chat.html', 
                           messages=messages, 
                           conversation=conversation, 
                           all_conversations=all_conversations, 
                           models=available_models, 
                           selected_model=conversation.selected_model if conversation else None,
                           cache_id=cache_id,
                           llm_service_type=llm_service_type)

@app.route('/conversation/new', methods=['POST'])
@login_required
def new_conversation():
    # Get all available models for the user (including RAG models)
    available_models = get_available_models_for_user(current_user)
    available_model_names = {model['ollama_model_name'] for model in available_models}

    if not available_model_names:
        app.logger.error("NEW_CONV: No models available for user. Cannot create new conversation.")
        flash("Cannot start a new chat: No AI models are currently available for your account.", "danger")
        return redirect(url_for('chat'))

    form_model_name = request.form.get('model')
    final_selected_model_name = None

    if form_model_name and form_model_name in available_model_names:
        final_selected_model_name = form_model_name
        app.logger.info(f"NEW_CONV: Model '{form_model_name}' from form is valid and selected.")
    else:
        if form_model_name:
            app.logger.warning(f"NEW_CONV: Model '{form_model_name}' from form is not in user's available list. Falling back.")
        else:
            app.logger.info("NEW_CONV: No model provided in form. Falling back.")
        
        # Fallback to the default model or the first available one
        default_model = get_default_model()
        if default_model and default_model['ollama_model_name'] in available_model_names:
            final_selected_model_name = default_model['ollama_model_name']
            app.logger.info(f"NEW_CONV: Falling back to default model '{final_selected_model_name}'.")
        elif available_model_names:
            final_selected_model_name = list(available_model_names)[0]
            app.logger.info(f"NEW_CONV: Falling back to first available model '{final_selected_model_name}'.")
        else:
            # This case is handled by the check at the beginning, but as a safeguard:
            app.logger.error("NEW_CONV: Critical fallback error. No models available.")
            flash("Error: No models available to create a conversation.", "danger")
            return redirect(url_for('chat'))

    app.logger.info(f"Creating new conversation for user {current_user.id} with model '{final_selected_model_name}'")
    conversation = Conversation(
        user_id=current_user.id,
        title="New Conversation", # Title can be set later based on first message
        selected_model=final_selected_model_name # This is now guaranteed to be an ollama_model_name string
    )
    db.session.add(conversation)
    db.session.commit()
    return redirect(url_for('chat', conversation_id=conversation.id))

@app.route('/conversation/<int:conversation_id>/rename', methods=['POST'])
@login_required
def rename_conversation(conversation_id):
    conversation = db.session.get(Conversation, conversation_id)
    # Check ownership
    if conversation.user_id != current_user.id:
        flash("Unauthorized access")
        return redirect(url_for('chat'))
    new_title = request.form.get('title', 'Untitled Conversation')
    conversation.title = new_title
    db.session.commit()
    
    flash("Conversation renamed successfully")
    return redirect(url_for('chat', conversation_id=conversation_id))

@app.route('/conversation/<int:conversation_id>/delete', methods=['POST'])
@login_required
def delete_conversation(conversation_id):
    conversation = db.session.get(Conversation, conversation_id)
    # Check ownership
    if conversation.user_id != current_user.id:
        flash("Unauthorized access")
        return redirect(url_for('chat'))
    else:
        db.session.delete(conversation)
        db.session.commit()
        flash("Conversation deleted")
    return redirect(url_for('chat'))

@app.route('/switch_model', methods=['POST'])
@login_required
def switch_model():
    conversation_id = request.form['conversation_id']
    new_model = request.form['model']
    conversation = db.session.get(Conversation, conversation_id)
    if conversation and conversation.user_id == current_user.id:
        conversation.selected_model = new_model
        db.session.commit()
        flash("Model switched successfully.")
    return redirect(url_for('chat'))

@app.route('/edit_message/<int:message_id>', methods=['POST'])
@login_required
def edit_message(message_id):
    new_content = request.form['content']
    message = db.session.get(ChatMessage, message_id)
    if message and message.conversation.user_id == current_user.id:
        message.content = new_content
        db.session.commit()
        flash("Message updated.")
    return redirect(url_for('chat'))

# ===========================
# File and Voice Upload Routes
# ===========================
def extract_text_from_document(document):
    """Extract text content from a document based on its MIME type"""
    try:
        mime_type = document.mime_type
        
        # For plain text files
        if (mime_type == 'text/plain'):
            return document.data.decode('utf-8')
        # For PDF files
        elif mime_type == 'application/pdf':
            try:
                import io
                from pdfminer.high_level import extract_text
                pdf_file = io.BytesIO(document.data)
                text = extract_text(pdf_file)
                return text
            except ImportError:
                app.logger.warning("pdfminer.six not installed. Cannot extract PDF text.")
                return "PDF text extraction requires pdfminer.six. Please install it with: pip install pdfminer.six"
                
        # For docx files
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            try:
                import io
                import docx
                docx_file = io.BytesIO(document.data)
                doc = docx.Document(docx_file)
                return "\n".join([para.text for para in doc.paragraphs])
            except ImportError:
                app.logger.warning("python-docx not installed. Cannot extract DOCX text.")
                return "DOCX text extraction requires python-docx. Please install it with: pip install python-docx"
        
        # For other formats, return a message
        else:
            return f"Content extraction not supported for {mime_type}. Using filename only."
    except Exception as e:
        app.logger.exception(f"Error extracting text from document: {str(e)}")
        return f"Error extracting text: {str(e)}"

@app.route('/upload_document', methods=['POST'])
@login_required
def upload_document():
    file = request.files['file']
    conversation_id = request.form['conversation_id']
    if file:
        try:
            # Get the conversation
            conversation = db.session.get(Conversation, conversation_id)
            if not conversation or conversation.user_id != current_user.id:
                flash("Unauthorized access")
                return redirect(url_for('chat'))
            
            # Save document to database
            doc = Document(
                conversation_id=conversation_id,
                filename=file.filename,
                data=file.read(),
                mime_type=file.mimetype
            )
            
            # Enable document mode for this conversation
            conversation.document_mode = True
            
            db.session.add(doc)
            # Commit here to ensure doc and conversation mode are saved before system message
            db.session.commit() 
            
            # Extract document text for context
            # file.seek(0)  # Reset file pointer
            # doc_text = extract_text_from_document(doc)
            
            # Create a system message to indicate document context mode
            system_message = ChatMessage(
                conversation_id=conversation_id,
                sender='ai',
                content=f"📄 Document '{file.filename}' has been uploaded. My responses will now be based only on knowledge from this document."
            )
            db.session.add(system_message)
            db.session.commit()
            
            flash("Document uploaded successfully. AI will now respond based on document content.")
        except Exception as e:
            app.logger.exception(f"Error uploading document: {str(e)}")
            flash(f"Error uploading document: {str(e)}")
            db.session.rollback()
    return redirect(url_for('chat', conversation_id=conversation_id))

@app.route('/upload_voice', methods=['POST'])
@login_required 
def upload_voice():
    if not check_whisper_model_exists():
        return jsonify({'success': False, 'error': 'Voice transcription service not available.'}), 500
        
    if 'voice' not in request.files:
        return jsonify({'success': False, 'error': 'No voice file part'}), 400
        
    file = request.files['voice']
    conversation_id = request.form.get('conversation_id')
    language = request.form.get('language', 'english')  # Default to English if not provided

    if file.filename == '':
        return jsonify({'success': False, 'error': 'No selected voice file'}), 400

    if not conversation_id:
        return jsonify({'success': False, 'error': 'Missing conversation ID'}), 400

    conversation = db.session.get(Conversation, conversation_id)
    if not conversation or conversation.user_id != current_user.id:
        return jsonify({'success': False, 'error': 'Conversation not found or unauthorized'}), 404

    # Use a temporary file for processing
    temp_audio_path = None
    transcription_result = None
    detected_language = None
    error_message = None

    try:
        # Save blob to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".webm") as temp_audio:
            file.save(temp_audio.name)
            temp_audio_path = temp_audio.name
            app.logger.info(f"Saved temporary voice file to {temp_audio_path}")

        # Transcribe using Whisper
        app.logger.info(f"Transcribing {temp_audio_path} with language hint: {language}")
        # Let Whisper detect language unless a specific one is strongly needed
        whisper_options = {"language": language if language in ['english', 'persian'] else None} 
        transcription_result = speech_service.transcribe_audio(temp_audio_path, **whisper_options)
        transcribed_text = transcription_result['text'].strip()
        detected_language = transcription_result.get('language', language)  # Use detected or fallback to hint
        app.logger.info(f"Transcription successful. Detected language: {detected_language}. Text: {transcribed_text}")

        if not transcribed_text:
            raise ValueError("Transcription resulted in empty text.")

        # --- Save User Message (Transcription) ---
        # Prefix with an indicator that it came from voice
        user_message_content = f"🎤: {transcribed_text}"
        user_message = ChatMessage(
            conversation_id=conversation.id,
            sender='user',
            content=user_message_content,
        )
        app.logger.info(f"Saved user transcription message with ID: {user_message.id}")

        # --- Get AI Response ---
        history = ChatMessage.query.filter_by(conversation_id=conversation.id).order_by(ChatMessage.created_at).all()
        formatted_history = [{"role": msg.sender, "content": msg.content} for msg in history]
        
        # Use the transcribed text as the latest user prompt
        # No need to include the "🎤: " prefix for the AI model context
        latest_prompt = transcribed_text 

        try:
            app.logger.info(f"Sending prompt to {llm_service_type} model {conversation.selected_model}: {latest_prompt}")
            # Use the new llm_service abstraction
            response = llm_service.chat(
                model=conversation.selected_model,
                messages=formatted_history
            )
            ai_response_text = response['message']['content']
            app.logger.info(f"Received AI response: {ai_response_text}")

            # --- Save AI Message ---
            ai_message = ChatMessage(
                conversation_id=conversation.id,
                sender='ai',
                content=ai_response_text,
            )
            db.session.add(ai_message)
            db.session.commit()
            app.logger.info(f"Saved AI response message with ID: {ai_message.id}")

            # --- Prepare JSON Response ---
            return jsonify({
                'success': True,
                'transcription': user_message_content, 
                'message_id': user_message.id,  # ID of the saved user message
                'ai_response': ai_response_text,
                'detected_language': detected_language,
            })

        except Exception as e:
            app.logger.error(f"Error getting AI response: {e}")
            error_message = f"Error getting AI response: {e}"
            # Still return success=True because transcription worked, but include error for AI part
            return jsonify({
                'success': True,  # Transcription succeeded
                'transcription': user_message_content,
                'message_id': user_message.id,
                'ai_response': None,  # Indicate AI response failed
                'error': error_message,  # Provide error detail
                'detected_language': detected_language,
            })

    except Exception as e:
        app.logger.error(f"Error processing voice file: {e}")
        error_message = f"Error processing voice: {e}"
        # Return success=False as the core voice processing failed
        return jsonify({'success': False, 'error': error_message, 'transcription': error_message}), 500
    finally:
        # Clean up temporary file
        if temp_audio_path and os.path.exists(temp_audio_path):
            try:
                os.remove(temp_audio_path)
                app.logger.info(f"Removed temporary voice file: {temp_audio_path}")
            except Exception as e:
                app.logger.error(f"Error removing temporary file {temp_audio_path}: {e}")

# Add a new function to call the AI model directly from backend
def call_ai_model(model_name, prompt):
    """Call the AI model synchronously and return the full response using the configured LLM service"""
    # Force use of the fixed model for voice assistant
    fixed_model = os.environ.get('DEFAULT_VOICE_MODEL', "llama2") # Consider making this configurable

    if not isinstance(prompt, str):
        app.logger.error(f"call_ai_model received non-string prompt: {type(prompt)}")
        raise TypeError("Prompt must be a string")

    app.logger.info(f"Calling {llm_service_type} model {fixed_model} with prompt: {prompt[:50]}...")

    # Detect if the prompt contains Persian text
    is_persian = any('\u0600' <= c <= '\u06FF' for c in prompt)

    # Add language instruction for Persian
    if is_persian and "پاسخ به زبان فارسی" not in prompt:
        prompt = "لطفا به سوال زیر به زبان فارسی پاسخ دهید:\n\n" + prompt
        app.logger.info("Added Persian language instruction to prompt")

    try:
        app.logger.info(f"Sending prompt to {llm_service_type} model {fixed_model}: {prompt}")
        # Use the new llm_service abstraction
        response = llm_service.chat(
            model=fixed_model,
            messages=[{'role': 'user', 'content': prompt}]
        )
        ai_response_text = response['message']['content']
        app.logger.info(f"Received AI response: {ai_response_text[:50]}...")
        return ai_response_text
    except Exception as e:
        app.logger.exception(f"Error calling {llm_service_type} API: {e}")
        raise Exception(f"Failed to get response from AI model via API: {e}")

# Add a route to get the voice recording
@app.route('/voice_recording/<int:recording_id>', methods=['GET'])
@login_required 
def get_voice_recording(recording_id):
    """Return the voice recording audio file"""
    temp_file_path = None
    try:
        # Get the document
        voice_doc = db.session.get(Document, recording_id)
        
        # Check if the voice belongs to a conversation owned by the current user
        conversation = db.session.get(Conversation, voice_doc.conversation_id)
        if not conversation or conversation.user_id != current_user.id:
            return "Unauthorized", 403
        
        # Create a temporary file to serve
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
            tmp.write(voice_doc.data)
            temp_file_path = tmp.name
        
        # Send the file without using the unsupported after_request parameter
        return send_file(
            temp_file_path,
            mimetype=voice_doc.mime_type,
            as_attachment=False,
            download_name=voice_doc.filename
        )
    except Exception as e:
        app.logger.exception(f"Error retrieving voice recording: {str(e)}")
        return f"Error retrieving voice recording: {str(e)}", 500
    finally:
        # Clean up the temporary file in a background thread to ensure
        # it happens after the file is served
        if temp_file_path and os.path.exists(temp_file_path):
            def cleanup_temp_file():
                try:
                    # Add a small delay to ensure file serving completes
                    time.sleep(1)
                    if os.path.exists(temp_file_path):
                        os.remove(temp_file_path)
                except Exception as e:
                    app.logger.warning(f"Failed to remove temporary file {temp_file_path}: {e}")
            
            import threading
            cleanup_thread = threading.Thread(target=cleanup_temp_file)
            cleanup_thread.daemon = True
            cleanup_thread.start()

# Update voice help route if needed, or remove if template is removed
@app.route('/voice_help')
def voice_help():
    """Provides detailed instructions for setting up voice recognition"""
    # Update this template or remove the route if the template is removed
    return render_template('voice_help.html') # Make sure this template exists and is updated
# ===========================
# Endpoint to Call the AI Model and Stream Response
# ===========================
def format_sse(data: str, event: str = None) -> str:
    """Formats a string according to the Server-Sent Events protocol."""
    msg = f'data: {data}\n\n'
    if event is not None:
        msg = f'event: {event}\n{msg}'
    return msg

def stream_llm_response(model_name, messages_history):
    """Streams response from the configured LLM service with robust error handling."""
    app.logger.info(f"--> Entering stream_llm_response for model: {model_name}")
    app.logger.info(f"--> Messages history count: {len(messages_history)}")
    app.logger.info(f"--> First message: {str(messages_history[0])[:200]}..." if messages_history else "No messages in history")
    
    stream_generator = None
    sent_any_chunk = False
    try:
        app.logger.info("--> Calling llm_service.stream_chat...")
        stream_generator = llm_service.stream_chat(model_name, messages_history)
        app.logger.info("--> Got generator object from llm_service.stream_chat.")
        app.logger.info("--> Starting the generator iteration loop")
        
        i = 0
        try:
            for i, chunk in enumerate(stream_generator):
                app.logger.debug(f"Raw chunk from llm_service (type: {type(chunk)}): {str(chunk)[:150]}")
                
                if isinstance(chunk, str) and chunk.startswith('data:'):
                    app.logger.debug("--> Chunk is a pre-formatted SSE string. Passing through.")
                    yield chunk if chunk.endswith('\n\n') else chunk + '\n\n'
                    sent_any_chunk = True
                elif isinstance(chunk, dict):
                    app.logger.debug("--> Chunk is a dictionary. Formatting into SSE.")
                    text_content = chunk.get('message', {}).get('content', '') or chunk.get('response', '')
                    if text_content:
                        sse_chunk = f"data: {json.dumps({'text': text_content})}\n\n"
                        yield sse_chunk
                        sent_any_chunk = True
                    else:
                        app.logger.warning(f"--> Received a dictionary chunk with no text content: {chunk}")
                else:
                    app.logger.warning(f"--> Received an unexpected chunk type: {type(chunk)}. Chunk: {str(chunk)[:150]}")
            app.logger.info(f"Exited streaming loop after {i+1} chunks.")
        except Exception as e:
            llm_service_type = app.config.get('LLM_SERVICE', 'unknown')
            app.logger.error(f"Error during {llm_service_type} stream: {e}", exc_info=True)
            error_payload = json.dumps({
                "error": f"Error during {llm_service_type} stream: {e}",
                "text": f"⚠️ An error occurred while streaming the response."
            })
            yield format_sse(error_payload, 'error')
            sent_any_chunk = True
    finally:
        app.logger.info(f"--> Exiting stream_llm_response for model: {model_name}")
        if not sent_any_chunk:
            app.logger.warning("No chunks were yielded from llm_service.stream_chat; sending empty response message.")
            empty_text = json.dumps({"error": "No response from model.", "text": "⚠️ No response from model."})
            yield f"data: {empty_text}\n\n"

@app.route('/call_model', methods=['POST'])
@login_required
def call_model():
    app.logger.info("""
    ####################################################
    #                  CALL_MODEL START                #
    ####################################################
    """)
    app.logger.info(f"Received /call_model request from user {current_user.id}")
    conversation_id = request.form['conversation_id']
    prompt = request.form['prompt']
    app.logger.info(f"Request details - conversation_id: {conversation_id}, prompt: {prompt[:50]}...")
    
    # Get conversation
    conversation = db.session.get(Conversation, conversation_id)
    if not conversation or conversation.user_id != current_user.id:
        app.logger.warning(f"Unauthorized access attempt to conversation {conversation_id}")
        return jsonify({"error": "Unauthorized"}), 403
    
    # Prepare message history
    messages_history = []
    for msg in conversation.messages:
        messages_history.append({
            'role': 'user' if msg.sender == 'user' else 'assistant',
            'content': msg.content
        })
    messages_history.append({'role': 'user', 'content': prompt})
    
    model_name = conversation.selected_model if hasattr(conversation, 'selected_model') and conversation.selected_model else DEFAULT_MODEL_NAME
    app.logger.info(f"Using model: {model_name}")
    
    # Save the user message to the DB
    user_message = ChatMessage(
        conversation_id=conversation_id,
        sender='user',
        content=prompt
    )
    db.session.add(user_message)
    db.session.commit()
    app.logger.info(f"Saved user message with ID {user_message.id}")
    
    # Check if selected model is a RAG model
    # ---> START RAG DEBUGGING LOGS
    app.logger.info(f"[RAG_DEBUG] About to query for model with ollama_model_name: '{model_name}'")
    selected_model = Model.query.filter_by(ollama_model_name=model_name).first()
    if selected_model:
        app.logger.info(f"[RAG_DEBUG] Found model in DB: ID={selected_model.id}, Name='{selected_model.display_name}', Is_RAG={selected_model.is_rag_model}")
    else:
        app.logger.warning(f"[RAG_DEBUG] Could not find a model in the DB with ollama_model_name: '{model_name}'")
    # ---> END RAG DEBUGGING LOGS
    
    # Store the original model name in case we need to switch models for RAG
    original_model_name = model_name 
    rag_context = None
    
    # Process RAG model outside the generator to avoid scoping issues
    if selected_model and selected_model.is_rag_model:
        app.logger.info(f"Processing RAG model request for model {selected_model.display_name} (id: {selected_model.id})")
        app.logger.info(f"RAG index ID: {selected_model.rag_index_id}, Base model ID: {selected_model.base_rag_model_id}")
        
        if not selected_model.rag_index_id:
            app.logger.error(f"RAG model {selected_model.display_name} has no associated RAG index!")
            # Add a system message to inform that the RAG model is misconfigured
            error_system_message = {
                "role": "system",
                "content": "The selected model is configured as a RAG model but has no associated RAG index. Please contact an administrator to fix this configuration."
            }
            messages_history.insert(0, error_system_message)
        else:
            app.logger.info(f"Processing RAG model request with index ID {selected_model.rag_index_id}")
        try:
            # Get the RAG index
            rag_index = RagIndex.query.get(selected_model.rag_index_id)
            if not rag_index:
                app.logger.error(f"RAG index with ID {selected_model.rag_index_id} not found")
                raise Exception(f"RAG index not found")
            
            # Get vector store path
            vector_store_path = os.path.join(
                app.config.get('RAG_INDEX_FOLDER', './rag_indexes'),
                rag_index.vector_store_path_segment
            )
            app.logger.info(f"RAG vector store path: {vector_store_path}")
            
            # Check if vector store exists
            if not os.path.exists(vector_store_path):
                app.logger.error(f"Vector store path does not exist: {vector_store_path}")
                raise Exception(f"Vector store not found")
            
            # Get the base model for this RAG model
            base_model = None
            if selected_model.base_rag_model_id:
                base_model = Model.query.get(selected_model.base_rag_model_id)
                app.logger.info(f"Using base model: {base_model.ollama_model_name if base_model else 'None'}")
            
            # Initialize vector store and perform semantic search
            import chromadb
            from chromadb.utils import embedding_functions
            from chromadb.config import Settings
            import httpx

            # Use the same embedding model that was used to create the index
            embedding_model_name = rag_index.embedding_model_name
            app.logger.info(f"Using embedding model: {embedding_model_name}")
            
            embedding_function = None
            try:
                # Try to use ollama embeddings if available
                ollama_ef = embedding_functions.OllamaEmbeddingFunction(
                    url=f"{os.environ.get('OLLAMA_HOST', 'http://local-ollama:11434')}/api/embeddings",
                    model_name=embedding_model_name
                )
                embedding_function = ollama_ef
                app.logger.info("Using Ollama for embeddings")
            except Exception as e:
                # Fall back to HuggingFace embeddings
                app.logger.warning(f"Failed to initialize Ollama embeddings: {e}. Falling back to HuggingFace.")
                hf_ef = embedding_functions.HuggingFaceEmbeddingFunction(
                    api_key=os.environ.get('HF_API_KEY', None),
                    model_name=embedding_model_name
                )
                embedding_function = hf_ef
                app.logger.info("Using HuggingFace for embeddings")

            # Connect to the existing Chroma collection
            app.logger.info(f"Connecting to ChromaDB at path: {vector_store_path}")
            try:
                client = chromadb.PersistentClient(path=vector_store_path, settings=Settings(anonymized_telemetry=False))
                
                # List collections to verify they exist
                collections = client.list_collections()
                collection_names = [c.name for c in collections]
                app.logger.info(f"Available collections: {collection_names}")
                
                collection_name_to_use = None
                if rag_index.name in collection_names:
                    collection_name_to_use = rag_index.name
                elif "documents" in collection_names:
                    collection_name_to_use = "documents"

                if collection_name_to_use:
                    app.logger.info(f"Using '{collection_name_to_use}' collection for RAG")
                    collection = client.get_collection(name=collection_name_to_use, embedding_function=embedding_function)
                else:
                    app.logger.error(f"No valid collection found in ChromaDB at {vector_store_path}")
                    raise Exception("RAG collection not found")
                
                # Check if collection has documents
                collection_count = collection.count()
                app.logger.info(f"Document collection contains {collection_count} chunks")
                
                if collection_count == 0:
                    app.logger.warning("RAG collection is empty! No document chunks available.")
            except ValueError as e:
                # This is likely an embedding function conflict.
                if "Embedding function conflict" in str(e):
                    app.logger.error(f"Embedding function conflict in RAG index '{rag_index.name}': {e}")
                    # Create a specific, user-friendly error message for this case
                    error_system_message = {
                        "role": "system",
                        "content": "The RAG model is misconfigured. The embedding function used for querying does not match the one used to create the index. An administrator must rebuild the RAG index to resolve this issue."
                    }
                    messages_history.insert(0, error_system_message)
                    # Skip the rest of the RAG processing
                    rag_context = "SKIP_RAG_PROCESSING"
                else:
                    # Re-raise other ValueErrors
                    raise
            except Exception as e:
                app.logger.error(f"Error connecting to ChromaDB: {e}")
                raise
            
            # If there was an embedding function conflict, we skip the query
            if rag_context == "SKIP_RAG_PROCESSING":
                # We've already added the error message, so we just need to make sure we fall through
                # to the non-RAG response path.
                results = None
            else:
                # Get user's last query (the prompt)
                query = prompt
                app.logger.info(f"Performing semantic search with query: {query[:100]}...")
                
                # Stage 1: Initial retrieval of a larger set of documents
                initial_k = 20
                results = collection.query(
                    query_texts=[query],
                    n_results=initial_k
                )

            if results and 'documents' in results and results['documents'] and results['documents'][0]:
                doc_chunks = results['documents'][0]
                
                # Stage 2: Reranking using the LLM
                app.logger.info(f"Stage 1 (Vector Search) retrieved {len(doc_chunks)} documents. Starting Stage 2 (LLM Reranking)...")

                def _get_rerank_score(doc_chunk):
                    # This function will be executed in parallel for each document chunk.
                    rerank_prompt = (
                        f"You are a relevance scoring assistant. Your task is to evaluate a document's relevance to a user query and provide a score between 0.0 and 1.0.\n"
                        f"IMPORTANT: You must respond with ONLY a single floating point number. Do not include any other text, explanation, or punctuation.\n\n"
                        f"--- EXAMPLE ---\n"
                        f"User Query: What is the capital of France?\n"
                        f"Document: Paris is a beautiful city known for the Eiffel Tower.\n"
                        f"Your Response: 0.9\n"
                        f"---------------\n\n"
                        f"--- TASK ---\n"
                        f"User Query: {query}\n"
                        f"Document: {doc_chunk}\n"
                        f"Your Response:"
                    )
                    # Switch to a faster model for reranking to prevent timeouts.
                    response = llm_service.chat('tinyllama:1.1b-chat-v1-q2_K', [{'role': 'user', 'content': rerank_prompt}], stream=False)
                    score = 0.0
                    try:
                        response_text = response.get('message', {}).get('content', '').strip()
                        match = re.search(r"(\d+\.?\d*)", response_text)
                        if match:
                            parsed_score = float(match.group(1))
                            if 1.0 < parsed_score <= 10.0:
                                score = parsed_score / 10.0
                            elif 0.0 <= parsed_score <= 1.0:
                                score = parsed_score
                    except (ValueError, TypeError, AttributeError):
                        pass # Errors are logged in the main thread if needed
                    return (doc_chunk, score)

                reranked_docs = []
                # Use a ThreadPoolExecutor to run reranking in parallel, limiting workers to prevent overwhelming Ollama.
                # Reduce workers to 2 to prevent Ollama from timing out under load.
                # Reduce workers to 1 to make reranking sequential and prevent timeouts.
                with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
                    future_to_doc = {executor.submit(_get_rerank_score, doc): doc for doc in doc_chunks}
                    for future in concurrent.futures.as_completed(future_to_doc):
                        try:
                            reranked_docs.append(future.result())
                        except Exception as exc:
                            app.logger.error(f'Reranking generated an exception: {exc}')

                # Sort documents by the new LLM-generated score
                reranked_docs.sort(key=lambda x: x[1], reverse=True)

                # Select the top N documents after reranking
                final_k = 5
                final_docs = [doc for doc, score in reranked_docs[:final_k]]
                
                app.logger.info(f"Stage 2 (LLM Reranking) complete. Selected top {len(final_docs)} documents.")
                for i, (doc, score) in enumerate(reranked_docs[:final_k]):
                    app.logger.info(f"  Reranked #{i+1}: (LLM Score: {score:.2f}) {doc[:100]}...")

                # Build context from the top reranked documents
                rag_context = "\n\n".join([f"Document chunk: {doc}" for doc in final_docs])

                # ---> START RAG CONTEXT LOGGING
                app.logger.info(f"\n\n[RAG_CONTEXT_PROVIDED]======================================\n{rag_context}\n========================================================\n")
                # ---> END RAG CONTEXT LOGGING

                # Create a new prompt that includes the RAG context
                rag_system_message = {
                    "role": "system",
                    "content": f"You are a helpful assistant. Answer questions based ONLY on the following context. " +
                              f"If the context doesn't contain relevant information, say 'I don't have information about that in my knowledge base.'\n\n" +
                              f"CONTEXT:\n{rag_context}"
                }
                
                # Insert system message at beginning and use base model name instead of RAG model
                messages_history.insert(0, rag_system_message)
                
                # Use the base model if available, otherwise use the RAG model (which might fail)
                if base_model:
                    model_name = base_model.ollama_model_name
                    app.logger.info(f"Switched to base model: {model_name} for RAG processing")
                else:
                    # Make sure we still have a valid model_name
                    model_name = original_model_name
                    app.logger.info(f"Using original model: {model_name} for RAG processing")
            else:
                app.logger.warning("Semantic search returned no results")
                # Add a system message to inform that no context was found
                rag_system_message = {
                    "role": "system",
                    "content": "You are a helpful assistant. No relevant information was found in the knowledge base for this query. " +
                              "Please respond that you don't have information about this in your knowledge base."
                }
                messages_history.insert(0, rag_system_message)
                
                # Use the base model if available
                if base_model:
                    model_name = base_model.ollama_model_name
                else:
                    # Make sure we still have a valid model_name
                    model_name = original_model_name
        except httpx.ReadTimeout as e:
            app.logger.error(f"Timeout during RAG query: {e}")
            error_system_message = {
                "role": "system",
                "content": "The request to the document index timed out. This could be due to a network issue or the embedding model taking too long to respond. Please try again later or contact an administrator."
            }
            messages_history.insert(0, error_system_message)
            # Use the base model if available
            if base_model:
                model_name = base_model.ollama_model_name
            else:
                model_name = original_model_name
        except Exception as e:
            app.logger.exception(f"Error during RAG processing: {e}")
            # Fallback to regular non-RAG response but with a warning
            error_system_message = {
                "role": "system",
                "content": f"WARNING: There was an error retrieving information from the RAG index: {str(e)}. " +
                          "Please inform the user that there was an error processing their request with RAG and they should contact the administrator."
            }
            messages_history.insert(0, error_system_message)
            # Use the base model if available
            if base_model:
                model_name = base_model.ollama_model_name
            else:
                # Make sure we still have a valid model_name
                model_name = original_model_name
    
    # Log the final model name after all potential changes
    app.logger.info(f"Final model name for generation: {model_name}")
    
    def response_wrapper():
        app.logger.info("Starting response_wrapper generator")
        full_response = ""
        ai_message_id = None
        user_id = current_user.id if hasattr(current_user, 'id') and current_user.id else 0
        conv_id = conversation_id
        generator_key = f"user_{user_id}_conv_{conv_id}"
        active_response_generators[generator_key] = False
        try:
            app.logger.info("Preparing message history for LLM API")
            app.logger.info(f"Final streaming model: {model_name}")
            chunk_count = 0
            try:
                app.logger.info("--> Entering stream_llm_response from response_wrapper...")
                
                # Stream the response using the prepared model_name and messages_history
                # All RAG processing is already done outside this function
                yielded_any = False
                app.logger.info(f"Starting stream with model: {model_name}")
                for chunk in stream_llm_response(model_name, messages_history):
                    app.logger.info(f"Yielding chunk #{chunk_count+1}: {str(chunk)[:60]}")
                    # Accumulate bot response text from each chunk
                    # Each chunk should be like: 'data: {"text": "..."}\n\n'
                    if chunk.startswith('data: '):
                        try:
                            data = json.loads(chunk[6:].strip())
                            if 'text' in data:
                                full_response += data['text']
                        except Exception as e:
                            app.logger.warning(f"Failed to parse streamed chunk for accumulation: {e}")
                    yield chunk
                    yielded_any = True
                    chunk_count += 1
                app.logger.info(f"Exited streaming loop after {chunk_count} chunks.")
                if not yielded_any:
                    app.logger.warning("No chunks were yielded from stream_llm_response; yielding fallback error chunk.")
                    error_text = json.dumps({"error": "No response from model.", "text": "⚠️ No response from model."})
                    yield f"data: {error_text}\n\n"
                # After streaming is done, save the bot message to DB
                if full_response.strip():
                    try:
                        ai_message = ChatMessage(
                            conversation_id=conversation_id,
                            sender='ai',
                            content=full_response
                        )
                        db.session.add(ai_message)
                        db.session.commit()
                        app.logger.info(f"Saved AI message with ID {ai_message.id}")
                        ai_message_id = ai_message.id
                    except Exception as e:
                        app.logger.error(f"Failed to save AI message to DB: {e}")
                        db.session.rollback()
                # Extract conversation topic after saving AI message
                try:
                    topic_prompt = "Please provide a concise conversation topic (single word or phrase, max 6 characters) for the conversation above."
                    # Build topic messages with full conversation context and AI response
                    topic_messages = messages_history + [
                        {"role": "assistant", "content": full_response},
                        {"role": "user", "content": topic_prompt}
                    ]
                    topic_response = llm_service.chat(model_name, topic_messages, stream=False)
                    app.logger.info(f"Topic extraction raw response: {topic_response}")
                    # Robustly extract topic
                    topic = None
                    resp = topic_response
                    if isinstance(resp, dict):
                        # Try OpenAI-style choices
                        if 'choices' in resp and isinstance(resp.get('choices'), list) and resp['choices']:
                            choice = resp['choices'][0]
                            msg_obj = choice.get('message') or choice
                            if isinstance(msg_obj, dict):
                                topic = msg_obj.get('content') or msg_obj.get('text')
                        # Fallback to top-level message
                        elif 'message' in resp and isinstance(resp['message'], dict):
                            topic = resp['message'].get('content') or resp['message'].get('text')
                        # Fallback to text
                        elif 'text' in resp:
                            topic = resp.get('text')
                    elif isinstance(resp, str):
                        topic = resp
                    # Normalize and truncate topic
                    topic = (topic or "").strip().strip('"').strip()
                    topic = topic[:191]  # Truncate to match DB schema (VARCHAR(191))
                    # Update DB title if found
                    if topic:
                        conv = db.session.get(Conversation, conversation_id)
                        conv.title = topic
                        db.session.commit()
                    else:
                        app.logger.warning("Topic extraction returned empty topic")
                    app.logger.info(f"Emitting topic SSE: {topic} for conversation {conversation_id}")
                    # Always emit a topic SSE event
                    yield f"data: {json.dumps({'topic': topic})}\n\n"
                except Exception as e:
                    app.logger.error(f"Error extracting topic: {e}")
                    yield f"data: {json.dumps({'topic': ''})}\n\n"
            except Exception as e:
                app.logger.error(f"Exception in streaming loop: {traceback.format_exc()}")
                error_text = json.dumps({"error": str(e), "text": f"⚠️ {str(e)}"})
                yield f"data: {error_text}\n\n"
        finally:
            app.logger.info("Exiting response_wrapper generator")
    app.logger.info("Returning streaming response")
    # Ensure correct mimetype for SSE and use stream_with_context
    return Response(stream_with_context(response_wrapper()), mimetype="text/event-stream")

# Add a dictionary to track active response generators
active_response_generators = {}

@app.route('/stop_response', methods=['POST'])
@login_required
def stop_response():
    """Stop an active AI response for a conversation"""
    conversation_id = request.form.get('conversation_id')
    app.logger.info(f"Request to stop response for conversation {conversation_id}")
    
    if not conversation_id:
        return jsonify({"success": False, "error": "No conversation ID provided"}), 400
        
    # Check permission (user must own the conversation)
    conversation = db.session.get(Conversation, conversation_id)
    if not conversation or conversation.user_id != current_user.id:
        return jsonify({"success": False, "error": "Unauthorized"}), 403
    
    # Set the flag to stop the generator for this conversation
    generator_key = f"user_{current_user.id}_conv_{conversation_id}"
    if generator_key in active_response_generators:
        active_response_generators[generator_key] = True
        app.logger.info(f"Set stop flag for generator {generator_key}")
        return jsonify({"success": True, "message": "Response generation stopping"})
    else:
        return jsonify({"success": False, "error": "No active response found"}), 404

@app.route('/toggle_document_mode/<int:conversation_id>', methods=['POST'])
@login_required
def toggle_document_mode(conversation_id):
    conversation = db.session.get(Conversation, conversation_id)
    
    # Check ownership
    if conversation.user_id != current_user.id:
        return jsonify({"success": False, "error": "Unauthorized"}), 403
        
    try:
        # Toggle document mode
        conversation.document_mode = not conversation.document_mode
        
        # Add a system message to indicate the change
        message_text = ""
        if conversation.document_mode:
            # Find the latest document
            latest_doc = Document.query.filter(
                Document.conversation_id == conversation_id,
                Document.mime_type != 'audio/wav' # Exclude voice recordings
            ).order_by(Document.uploaded_at.desc()).first()
            if latest_doc:
                message_text = f"📄 Document mode enabled. My responses will now be based only on document: '{latest_doc.filename}'."
            else:
                message_text = "📄 Document mode enabled, but no documents found. Please upload a document."
        else:
            message_text = "📄 Document mode disabled. I'll now use my general knowledge to answer your questions."
        
        system_message = ChatMessage(
            conversation_id=conversation_id,
            sender='ai',
            content=message_text
        )
        db.session.add(system_message)
        db.session.commit()
        
        return jsonify({
            "success": True,
            "document_mode": conversation.document_mode,
            "message": message_text
        })
    except Exception as e:
        app.logger.exception(f"Error toggling document mode: {str(e)}")
        db.session.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

# Add a new route to update conversation title directly
@app.route('/conversation/<int:conversation_id>/update_title', methods=['POST'])
@login_required
def update_conversation_title(conversation_id):
    """Update a conversation title and save it to the database"""
    try:
        conversation = db.session.get(Conversation, conversation_id)
        
        # Check ownership
        if conversation.user_id != current_user.id:
            return jsonify({"success": False, "error": "Unauthorized"}), 403
        
        data = request.get_json()
        if not data or 'title' not in data:
            return jsonify({"success": False, "error": "No title provided"}), 400
        
        title = data['title'].strip()
        if not title:
            title = "Untitled Conversation"
        
        # Update the title in the database
        conversation.title = title
        db.session.commit()
        
        app.logger.info(f"Title updated for conversation {conversation_id}: '{title}'")
        return jsonify({"success": True, "title": title})
    except Exception as e:
        app.logger.exception(f"Error updating conversation title: {str(e)}")
        db.session.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

# Add routes for text-to-speech capabilities
@app.route('/voice_for_message/<int:message_id>', methods=['GET'])
@login_required
def voice_for_message(message_id):
    """Check if a voice recording exists for a message and return it"""
    try:
        message = db.session.get(ChatMessage, message_id)
        
        # Check if this message belongs to a conversation owned by the current user
        conversation = db.session.get(Conversation, message.conversation_id)
        if not conversation or conversation.user_id != current_user.id:
            return "Unauthorized", 403
        
        # Check if this is a voice response message
        if message.sender == 'ai' and message.content.startswith('VOICE_RESPONSE:'):
            parts = message.content.split(':')
            if len(parts) >= 3:
                language = parts[1]
                content = ':'.join(parts[2:])
                
                # Check if we already have a voice recording for this AI message
                # We'll search for documents with a filename containing the message_id
                voice_doc = Document.query.filter(
                    Document.conversation_id == message.conversation_id,
                    Document.filename.like(f'ai_voice_response_%_{message_id}.wav')
                ).first()
                
                if voice_doc:
                    # Create a temporary file to serve
                    temp_file_path = None
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as tmp:
                        tmp.write(voice_doc.data)
                        temp_file_path = tmp.name
                    
                    # Function to clean up temporary file
                    def cleanup_temp_file():
                        try:
                            # Add a small delay to ensure file serving completes
                            time.sleep(1)
                            if os.path.exists(temp_file_path):
                                os.remove(temp_file_path)
                        except Exception as e:
                            app.logger.warning(f"Failed to remove temporary file {temp_file_path}: {e}")
                    
                    # Start cleanup thread
                    cleanup_thread = threading.Thread(target=cleanup_temp_file)
                    cleanup_thread.daemon = True
                    cleanup_thread.start()
                    
                    # Return the file
                    return send_file(
                        temp_file_path,
                        mimetype="audio/wav",
                        as_attachment=False
                    )
        
        # If we get here, no voice recording was found
        return "No voice recording found for this message", 404
    
    except Exception as e:
        app.logger.exception(f"Error retrieving voice for message: {e}")
        return f"Error retrieving voice: {e}", 500

@app.route('/synthesize_for_message', methods=['POST'])
@login_required
def synthesize_for_message():
    """Generate text-to-speech for a message and store it"""
    try:
        text = request.form['text']
        language = request.form.get('language', 'english')
        message_id = request.form.get('message_id')
        
        if not text.strip():
            return "No text provided", 400
        
        # Generate speech
        speech_file = synthesize_speech(text, language) # This function needs to be defined or imported
        if not speech_file:

            return "Failed to synthesize speech", 500
        
        # Find the associated message if message_id was provided
        if message_id:
            message = db.session.get(ChatMessage, int(message_id))
            if message:
                conversation_id = message.conversation_id
            else:
                return "Message not found", 404
        else:
            # If no specific message, use the active conversation
            conversation_id = request.form.get('conversation_id')
            if not conversation_id:
                return "No conversation ID provided", 400
        
        # Check user authorization for this conversation
        conversation = db.session.get(Conversation, conversation_id)
        if not conversation or conversation.user_id != current_user.id:
            return "Unauthorized", 403
        
        # Read the generated audio file
        with open(speech_file, 'rb') as audio_file:
            audio_data = audio_file.read()
        
        # Save the speech as a document with reference to the message
        filename = f"ai_voice_response_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
        if message_id:
            filename += f"_{message_id}"
        filename += ".wav"
        
        voice_doc = Document(
            conversation_id=conversation_id,
            filename=filename,
            data=audio_data,
            mime_type="audio/wav"
        )
        
        db.session.add(voice_doc)
        db.session.commit()
        
        # Clean up the temporary file
        try:
            os.remove(speech_file)
        except:
            pass
        
        # Return the audio
        return send_file(
            io.BytesIO(audio_data),
            mimetype="audio/wav",
            as_attachment=False
        )
    except Exception as e:
        app.logger.exception(f"Error synthesizing speech: {e}")
        return f"Error synthesizing speech: {e}", 500

# Add a function to check system dependencies at startup
def check_system_dependencies():
    deps = {
        "ffmpeg": check_ffmpeg_installed(),
        "whisper": check_whisper_model_exists() # Check if base model can load
    }
    return deps

# === Add Placeholder for synthesize_speech ===
def synthesize_speech(text, language):
    """Placeholder function for text-to-speech synthesis."""
    app.logger.warning(f"Placeholder synthesize_speech called for language '{language}'. Text: {text[:50]}...")
    # In a real implementation, this would call Bark via speech_service
    # and return the path to the generated audio file.
    # For now, return None to indicate failure.
    # Example call (if speech_service had this method):
    # return speech_service.synthesize_speech(text, language)
    return None
# === End Placeholder ===

# Create a session variable to store the user's LLM service preference
@app.route('/switch_llm_service', methods=['POST'])
@login_required
def switch_llm_service():
    """Switch the LLM service (Ollama or Llama.cpp) for the current user"""
    try:
        new_service = request.form.get('llm_service', 'ollama').lower()
        
        # Validate the service type
        if new_service not in ['ollama', 'llamacpp']:
            return jsonify({"success": False, "error": f"Invalid LLM service type: {new_service}"}), 400
        
        app.logger.info(f"User {current_user.id} switching LLM service to {new_service}")
        
        # Store the user's preference in session
        session['user_llm_service'] = new_service
        
        # Create a new LLM service instance
        global llm_service
        global llm_service_type
        
        # Create temp service to test connectivity
        try:
            # Use the factory to create the new service
            temp_service = LLMServiceFactory.create_service_by_type(new_service)
            # Test listing models to ensure connectivity
            available_models = temp_service.list_models()
            
            if not available_models:
                app.logger.warning(f"No models found for {new_service} service")
                
            # If we get here, the service is working
            llm_service = temp_service
            llm_service_type = new_service
            
            # Update the application's cached models list
            app.config['LLM_MODELS'] = available_models if available_models else [DEFAULT_MODEL_NAME]
            app.logger.info(f"Successfully switched to {new_service} service with models: {app.config['LLM_MODELS']}")
            
            return jsonify({
                "success": True, 
                "service": new_service,
                "models": app.config['LLM_MODELS']
            })
            
        except Exception as e:
            app.logger.exception(f"Error switching to {new_service} service: {e}")
            return jsonify({"success": False, "error": f"Failed to connect to {new_service} service: {str(e)}"}), 500
            
    except Exception as e:
        app.logger.exception(f"Error in switch_llm_service: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/test_ollama', methods=['GET'])
def test_ollama():
    """
    Test connectivity to Ollama service and return detailed diagnostics.
    """
    results = {
        "status": "Running diagnostics...",
        "ollama_host": os.environ.get("OLLAMA_HOST", "http://ollama:11434"),
        "llm_service_type": os.environ.get("LLM_SERVICE", "ollama"),
        "tests": []
    }
    
    try:
        # Test 1: Basic connectivity via requests
        results["tests"].append({"name": "Basic connectivity test"})
        try:
            import requests
            ollama_url = os.environ.get("OLLAMA_HOST", "http://ollama:11434")
            health_url = f"{ollama_url}/api/tags"
            app.logger.info(f"Testing basic connectivity to Ollama at {health_url}")
            response = requests.get(health_url, timeout=5)
            results["tests"][-1]["status"] = f"Success ({response.status_code})"
            results["tests"][-1]["details"] = f"Connected to {health_url}"
            
            # Include the first 500 chars of the response for verification
            response_data = response.json()
            results["tests"][-1]["response_preview"] = str(response_data)[:500]
        except Exception as e:
            results["tests"][-1]["status"] = "Failed"
            results["tests"][-1]["details"] = f"Error: {str(e)}"
            app.logger.exception(f"Basic connectivity test failed: {e}")
        
        # Test 2: Try to list models
        results["tests"].append({"name": "List models test"})
        try:
            models = llm_service.list_models()
            results["tests"][-1]["status"] = "Success"
            results["tests"][-1]["details"] = f"Found {len(models)} models"
            results["tests"][-1]["models"] = models
        except Exception as e:
            results["tests"][-1]["status"] = "Failed"
            results["tests"][-1]["details"] = f"Error: {str(e)}"
            app.logger.exception(f"List models test failed: {e}")
        
        # Test 3: Simple completion without streaming
        results["tests"].append({"name": "Simple completion test"})
        try:
            import requests
            import json
            
            ollama_url = os.environ.get("OLLAMA_HOST", "http://ollama:11434")
            api_url = f"{ollama_url}/api/chat"
            payload = {
                "model": DEFAULT_MODEL_NAME,
                "messages": [{"role": "user", "content": "Hello, say hi in one word"}],
                "stream": False
            }
            app.logger.info(f"Testing simple completion to {api_url}")
            response = requests.post(api_url, json=payload, timeout=10)
            
            if response.status_code == 200:
                response_data = response.json()
                results["tests"][-1]["status"] = "Success"
                results["tests"][-1]["details"] = "Completion received"
                results["tests"][-1]["response"] = str(response_data)[:500]
            else:
                results["tests"][-1]["status"] = "Failed"
                results["tests"][-1]["details"] = f"Error: Status {response.status_code}"
                results["tests"][-1]["response"] = response.text[:500]
        except Exception as e:
            results["tests"][-1]["status"] = "Failed"
            results["tests"][-1]["details"] = f"Error: {str(e)}"
            app.logger.exception(f"Simple completion test failed: {e}")
        
        # Overall status
        failed_tests = [t for t in results["tests"] if t.get("status", "").startswith("Failed")]
        if failed_tests:
            results["status"] = f"Failed ({len(failed_tests)}/{len(results['tests'])} tests failed)"
        else:
            results["status"] = "Success (all tests passed)"
            
    except Exception as e:
        results["status"] = "Error running diagnostics"
        results["error"] = str(e)
        app.logger.exception(f"Error in /test_ollama endpoint: {e}")
    
    return jsonify(results)

# Main entry point
def initialize_rbac_data():
    app.logger.info("FUNC_INIT_RBAC: Entered initialize_rbac_data function.")
    """
    Initializes default roles and populates the Model table
    from available Ollama models if they don't already exist.
    Also assigns all found models to the 'admin' role.
    """
    # Skip RBAC initialization if database isn't ready
    if not app.config.get('DATABASE_INITIALIZED', False):
        app.logger.warning("Skipping RBAC data initialization as database tables are not fully ready")
        app.logger.info(f"Managed models from config: {app.config.get('MANAGED_OLLAMA_MODELS', [])}")
        return
        
    # Proceed with initialization if database is ready
    try:
        with app.app_context():
            app.logger.info("FUNC_INIT_RBAC: Attempting to process roles and models...") # Cascade Temp Log (was: Initializing RBAC data (roles and models)...)

            # 1. Create default roles
            default_roles_data = {
                "admin": "Administrator with full access to all models and system settings.",
                "user": "Standard user with access to a default set of models."
            }
            admin_role_obj = None
            user_role_obj = None

            for role_name, role_desc in default_roles_data.items():
                role = Role.query.filter_by(name=role_name).first()
                if not role:
                    role = Role(name=role_name, description=role_desc)
                    db.session.add(role)
                    app.logger.info(f"Created role: {role_name}")
                if role_name == "admin":
                    admin_role_obj = role
                elif role_name == "user":
                    user_role_obj = role
            
            try:
                db.session.commit() # Commit roles first
            except Exception as e:
                db.session.rollback()
                app.logger.error(f"Error committing roles: {e}", exc_info=True)
                return # Cannot proceed without roles

            if not admin_role_obj:
                app.logger.error("Admin role could not be found or created. Cannot proceed.")
                return

            # 2. Create and assign 'Admin' role to the default admin user (admin@admin.com)
            default_admin_email = "admin@admin.com"
            admin_user_obj = User.query.filter_by(email=default_admin_email).first()
            if not admin_user_obj:
                app.logger.info(f"Default admin user '{default_admin_email}' not found. Creating new admin user.")
                hashed_password = generate_password_hash("admin")
                admin_user_obj = User(
                    username="admin",
                    email=default_admin_email,
                    firstname="admin",
                    lastname="admin",
                    password_hash=hashed_password,
                    confirmed=True, 
                    is_active=True
                )
                db.session.add(admin_user_obj)
            try:
                db.session.commit() # Commit new admin user first
                app.logger.info(f"Successfully created default admin user '{default_admin_email}'.") 
            except Exception as e:
                db.session.rollback()
                app.logger.error(f"Error creating default admin user '{default_admin_email}': {e}", exc_info=True)
                # Don't return, try to proceed with other admin if configured
            
            # Ensure the default admin user is active if they exist
            if admin_user_obj and not admin_user_obj.is_active:
                app.logger.info(f"Ensuring default admin user '{default_admin_email}' is active.")
                admin_user_obj.is_active = True
                # Commit this change before proceeding with role assignment
                try:
                    db.session.commit()
                    app.logger.info(f"Default admin user '{default_admin_email}' set to active.")
                except Exception as e:
                    db.session.rollback()
                    app.logger.error(f"Error setting default admin user '{default_admin_email}' to active: {e}", exc_info=True)

            # Assign Admin role to the default admin user if they exist and don't have it
            if admin_user_obj and admin_role_obj not in admin_user_obj.roles:
                admin_user_obj.roles.append(admin_role_obj)
                try:
                    db.session.commit()
                    app.logger.info(f"Assigned 'admin' role to default admin user '{default_admin_email}'.") 
                except Exception as e:
                    db.session.rollback()
                    app.logger.error(f"Error assigning 'admin' role to default admin user '{default_admin_email}': {e}", exc_info=True)
            elif admin_user_obj:
                app.logger.info(f"Default admin user '{default_admin_email}' already has 'admin' role or was just created with it.")

            # 3. Assign 'Admin' role to the admin user from .env (if different from default and exists)
            admin_username_env = app.config.get('ADMIN_USERNAME')
            if admin_username_env and admin_username_env != "admin": # Check if .env admin is set and different from default 'admin'
                env_admin_user = User.query.filter_by(username=admin_username_env).first()
                if env_admin_user:
                    if admin_role_obj not in env_admin_user.roles:
                        env_admin_user.roles.append(admin_role_obj)
                        try:
                            db.session.commit()
                            app.logger.info(f"Assigned 'admin' role to .env admin user '{admin_username_env}'.") 
                        except Exception as e:
                            db.session.rollback()
                            app.logger.error(f"Error assigning 'admin' role to .env admin user '{admin_username_env}': {e}", exc_info=True)
                    else:
                        app.logger.info(f".env admin user '{admin_username_env}' already has 'admin' role.")
                else:
                    app.logger.warning(f"Admin user '{admin_username_env}' specified in .env not found in database. Cannot assign 'admin' role.")
            elif not admin_username_env:
                app.logger.info("ADMIN_USERNAME not set in .env file. Skipping .env admin role assignment.")
            elif admin_username_env == "admin":
                app.logger.info("ADMIN_USERNAME in .env is 'admin', which is already handled as the default admin. Skipping redundant assignment.")

            # 4. Populate Model table from MANAGED_OLLAMA_MODELS and assign to admin
            app.logger.info("Processing managed models for Model table and admin assignment...")
            
            admin_models_assigned_this_run = [] # To track models assigned to admin in this run
            llm_service_type = app.config.get('LLM_SERVICE_TYPE', 'ollama').lower()


            if llm_service_type == 'ollama':
                managed_ollama_models_from_config = app.config.get('MANAGED_OLLAMA_MODELS', [])
                if not managed_ollama_models_from_config:
                    app.logger.warning("No managed Ollama models found in app.config (MANAGED_OLLAMA_MODELS is empty or not set). "
                                   "Model table population from this list will be skipped.")
                
                for ollama_model_name in managed_ollama_models_from_config:
                    model = Model.query.filter_by(ollama_model_name=ollama_model_name).first()
                    if not model:
                        # Attempt to create a somewhat friendly display name
                        display_parts = ollama_model_name.split(':')[0].replace('-', ' ').replace('_', ' ')
                        display_name_generated = ' '.join(word.capitalize() for word in display_parts.split(' '))
                        tag_suffix = ""
                        if ':' in ollama_model_name:
                            tag = ollama_model_name.split(':')[-1]
                            if tag.lower() != 'latest':
                                 tag_suffix = f" ({tag.capitalize()})"
                            # else: # For 'latest', no specific tag suffix or a generic one like "(Latest)"
                                 # tag_suffix = " (Latest)" # Optional: if you want to explicitly mark 'latest'
                        final_display_name = display_name_generated + tag_suffix

                        model = Model(
                            display_name=final_display_name, # This is the user-facing name
                            ollama_model_name=ollama_model_name, # This is for Ollama API
                            description=f"Ollama model: {ollama_model_name}",
                            is_active=True # Managed models are active by default
                        )
                        db.session.add(model)
                        app.logger.info(f"Created new Model DB entry for managed model: {ollama_model_name} (Display: {final_display_name})")
                        try:
                            db.session.commit() # Commit each new model to get its ID for relationships
                        except Exception as e:
                            db.session.rollback()
                            app.logger.error(f"Error committing new model {ollama_model_name}: {e}", exc_info=True)
                            continue # Skip to next model
                    elif not model.is_active:
                        app.logger.info(f"Model '{ollama_model_name}' found in DB but was inactive. Activating it as it's a managed model.")
                        model.is_active = True
                        # No immediate commit needed here, will be committed with role assignment or at the end of this section.

                    # Assign to admin role if not already assigned and model object exists
                    if admin_role_obj and model and model.id is not None: # Ensure model is committed or fetched with an ID
                        if model not in admin_role_obj.models:
                            admin_role_obj.models.append(model)
                            admin_models_assigned_this_run.append(model.display_name)
                            app.logger.info(f"Assigned model '{model.display_name}' to 'admin' role.")
                    elif not model:
                        app.logger.warning(f"Skipped assigning model {ollama_model_name} to admin as model object was None (likely due to creation error).")

            
            elif llm_service_type == 'llamacpp':
                app.logger.info("LLM service is LlamaCPP. Checking/creating DB entry for its default model.")
                llamacpp_default_model_name = app.config.get('EFFECTIVE_DEFAULT_MODEL_NAME')
                if llamacpp_default_model_name:
                    model = Model.query.filter_by(ollama_model_name=llamacpp_default_model_name).first()
                    if not model:
                        model = Model(
                            display_name=llamacpp_default_model_name, 
                            ollama_model_name=llamacpp_default_model_name, # Using ollama_model_name field for identifier consistency
                            description=f"LlamaCPP model: {llamacpp_default_model_name}",
                            is_active=True
                        )
                        db.session.add(model)
                        app.logger.info(f"Created new Model DB entry for LlamaCPP default model: {llamacpp_default_model_name}")
                        try:
                            db.session.commit() # Commit to get ID
                        except Exception as e:
                            db.session.rollback()
                            app.logger.error(f"Error committing LlamaCPP default model {llamacpp_default_model_name} to DB: {e}", exc_info=True)
                            model = None # Ensure model is None if commit failed
                    
                    if admin_role_obj and model and model.id is not None:
                        if model not in admin_role_obj.models:
                            admin_role_obj.models.append(model)
                            admin_models_assigned_this_run.append(model.display_name) # CORRECTED HERE
                            app.logger.info(f"Assigned LlamaCPP model '{model.display_name}' to 'admin' role.") # Also ensure log uses display_name
                else:
                    app.logger.warning("LlamaCPP service type, but no EFFECTIVE_DEFAULT_MODEL_NAME found in app.config.")
            else:
                app.logger.info(f"LLM service type is '{llm_service_type}'. Model DB population from managed list is primarily for Ollama.")

            if admin_models_assigned_this_run:
                app.logger.info(f"Models assigned/confirmed for admin role in this initialization run: {', '.join(admin_models_assigned_this_run)}")

            try:
                db.session.commit() # Commit all changes (new models, model activations, role assignments)
                app.logger.info("Committed all model and admin role assignment changes for initialize_rbac_data.")
            except Exception as e:
                db.session.rollback()
                app.logger.error(f"Final commit in initialize_rbac_data failed: {e}", exc_info=True)

            app.logger.info("FUNC_INIT_RBAC: Exiting initialize_rbac_data function.")
    except Exception as e:
        app.logger.error(f"Error in initialize_rbac_data: {str(e)}")
        # Set the flag in app config to indicate database is not fully initialized
        app.config['DATABASE_INITIALIZED'] = False
# ===========================
@app.route('/admin/rag')
@login_required
def admin_rag_page():
    if not current_user.has_role('admin'):
        flash('You do not have permission to access this page.', 'danger')
        return redirect(url_for('index'))
    
    indexes = RagIndex.query.options(joinedload(RagIndex.documents)).order_by(RagIndex.created_at.desc()).all()
    base_models = Model.query.filter_by(is_rag_model=False, is_active=True).order_by(Model.display_name).all()
    documents = RagDocument.query.order_by(RagDocument.uploaded_at.desc()).all()
    documents_json = json.dumps([{'id': doc.id, 'filename': doc.filename} for doc in documents])
    
    # Get Ollama models for embeddings
    ollama_models = app.config.get('MANAGED_OLLAMA_MODELS', [])
    if not ollama_models and 'OLLAMA_MODELS' in os.environ:
        # Fallback to env var if app.config doesn't have the models
        ollama_models = [model.strip() for model in os.environ.get('OLLAMA_MODELS').split(',')]
    
    # Add 'ollama/' prefix to all models to match the format expected by the embedding function
    available_embedding_models = [f"ollama/{model}" for model in ollama_models]
    
    # If no models available, add a default one
    if not available_embedding_models:
        available_embedding_models = ['ollama/tinyllama:1.1b-chat-v1-q2_K']
    return render_template('admin_rag.html', indexes=indexes, base_models=base_models, available_embedding_models=available_embedding_models, documents=documents, documents_json=documents_json)

@app.route('/admin/rag/index/create', methods=['POST'])
@login_required
def create_rag_index():
    if not current_user.has_role('admin'):
        flash('You do not have permission to perform this action.', 'danger')
        return redirect(url_for('admin_rag_page'))

    try:
        name = request.form.get('name')
        base_model_id_str = request.form.get('base_model_id')
        embedding_model_name = request.form.get('embedding_model_name')

        if not name or not base_model_id_str or not embedding_model_name:
            flash('All fields are required.', 'danger')
            return redirect(url_for('admin_rag_page'))

        try:
            base_model_id = int(base_model_id_str)
        except (ValueError, TypeError):
            flash('Invalid Base Model ID.', 'danger')
            return redirect(url_for('admin_rag_page'))

        # Verify that the base model exists and is not a RAG model
        base_model = Model.query.filter_by(id=base_model_id, is_rag_model=False).first()
        if not base_model:
            flash('The selected base model is invalid or is already a RAG model.', 'danger')
            return redirect(url_for('admin_rag_page'))

        # Create a unique path for the vector store
        vector_store_path_segment = f"rag_{name.lower().replace(' ', '_')}_{int(time.time())}"
        vector_store_path = os.path.join(app.config.get('RAG_INDEX_FOLDER', './rag_indexes'), vector_store_path_segment)
        os.makedirs(vector_store_path, exist_ok=True)

        # Create the RAG index
        new_index = RagIndex(
            name=name,
            base_model_id=base_model_id,
            embedding_model_name=embedding_model_name,
            vector_store_path_segment=vector_store_path_segment,
            created_by_id=current_user.id
        )
        db.session.add(new_index)
        db.session.flush()  # Flush to get the new_index.id

        # Create the associated RAG model
        new_model = Model(
            ollama_model_name=f"{base_model.ollama_model_name}-rag-{new_index.id}",  # More descriptive ollama_model_name
            display_name=name,
            description=f"RAG model based on index '{name}' using '{base_model.display_name}'",
            is_rag_model=True,
            rag_index_id=new_index.id,
            base_rag_model_id=base_model.id
        )
        db.session.add(new_model)
        db.session.commit()

        flash(f'RAG Index and Model "{name}" created successfully.', 'success')
        app.logger.info(f"RAG Index '{name}' (ID: {new_index.id}) and Model (ID: {new_model.id}) created by user {current_user.id}")

    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error creating RAG index: {e}", exc_info=True)
        flash('An unexpected error occurred while creating the RAG index.', 'danger')

    return redirect(url_for('admin_rag_page'))

@app.route('/admin/rag/index/delete', methods=['POST'])
@login_required
def delete_rag_index():
    """Delete a RAG index and its associated model."""
    if not current_user.has_role('admin'):
        return jsonify({'success': False, 'message': 'You do not have permission to perform this action.'}), 403
    
    try:
        data = request.get_json()
        index_id = data.get('index_id')
        
        if not index_id:
            return jsonify({'success': False, 'message': 'Index ID is required.'}), 400
            
        # Get the RAG index
        rag_index = RagIndex.query.get(index_id)
        if not rag_index:
            return jsonify({'success': False, 'message': f'RAG index with ID {index_id} not found.'}), 404
        
        # Get the associated RAG model if it exists
        rag_model = Model.query.filter_by(rag_index_id=index_id).first()
        
        # Store names for logging
        index_name = rag_index.name
        model_name = rag_model.display_name if rag_model else 'None'
        
        # Delete the associated RAG model if it exists
        if rag_model:
            # First remove role associations to prevent foreign key constraint errors
            # Get all roles that have this model
            role_model_assoc = db.session.query(role_models).filter(role_models.c.model_id == rag_model.id).all()
            for role_id, _ in role_model_assoc:
                role = Role.query.get(role_id)
                if role and rag_model in role.models:
                    role.models.remove(rag_model)
            
            db.session.delete(rag_model)
            app.logger.info(f"Deleted RAG model '{model_name}' (ID: {rag_model.id}) associated with index '{index_name}'")
        
        # Delete the RAG index
        db.session.delete(rag_index)
        
        db.session.commit()
        
        # Try to delete the vector store directory if possible
        try:
            if rag_index.vector_store_path_segment and app.config.get('RAG_VECTOR_STORES_DIR'):
                vector_store_path = os.path.join(app.config['RAG_VECTOR_STORES_DIR'], rag_index.vector_store_path_segment)
                if os.path.exists(vector_store_path):
                    try:
                        import shutil
                        shutil.rmtree(vector_store_path)
                        app.logger.info(f"Deleted vector store directory: {vector_store_path}")
                    except Exception as ve:
                        app.logger.error(f"Error deleting vector store directory: {ve}")
        except Exception as path_err:
            app.logger.warning(f"Could not construct vector store path: {path_err}")
        
        app.logger.info(f"Admin user {current_user.id} ({current_user.username}) deleted RAG index '{index_name}' (ID: {index_id})")
        return jsonify({'success': True, 'message': f'RAG index "{index_name}" and its model have been deleted successfully.'})
        
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error deleting RAG index: {e}", exc_info=True)
        return jsonify({'success': False, 'message': 'Failed to delete RAG index due to a server error.'}), 500

    name = request.form.get('name')
    base_model_id = request.form.get('base_model_id')
    embedding_model_name = request.form.get('embedding_model_name')

    if not all([name, base_model_id, embedding_model_name]):
        flash('All fields are required.', 'danger')
        return redirect(url_for('admin_rag_page'))

    if RagIndex.query.filter_by(name=name).first():
        flash(f'An index with the name "{name}" already exists.', 'danger')
        return redirect(url_for('admin_rag_page'))

    base_model = Model.query.get(base_model_id)
    if not base_model:
        flash('Invalid base model selected.', 'danger')
        return redirect(url_for('admin_rag_page'))

    try:
        vector_store_path_segment = f"{secure_filename(name.lower().replace(' ', '_'))}_{uuid.uuid4().hex[:8]}"
        
        new_index = RagIndex(
            name=name,
            base_model_id=base_model.id,
            vector_store_path_segment=vector_store_path_segment,
            embedding_model_name=embedding_model_name,
            created_by_id=current_user.id
        )
        db.session.add(new_index)
        db.session.flush()  # So we can access new_index.id before commit

        rag_index_id = new_index.id
        rag_model_display_name = f"{name} (RAG)"
        rag_model_ollama_name = f"rag_{vector_store_path_segment}"

        new_rag_model = Model(
            ollama_model_name=rag_model_ollama_name,
            display_name=rag_model_display_name,
            description=f"RAG model based on '{base_model.display_name}' with the '{name}' document index.",
            is_active=True,
            is_rag_model=True,
            base_rag_model_id=base_model.id,
            rag_index_id=rag_index_id
        )
        db.session.add(new_rag_model)
        
        # Auto-assign the RAG model to the admin role
        admin_role = Role.query.filter_by(name='admin').first()
        if admin_role:
            # Check if relationship already exists
            if new_rag_model not in admin_role.models:
                admin_role.models.append(new_rag_model)
                app.logger.info(f"Automatically assigned RAG model '{rag_model_display_name}' to admin role")
        
        db.session.commit()

        # ✅ Background thread safely using only the ID
        def background_indexing(idx):
            try:
                with app.app_context():
                    index_to_update = RagIndex.query.get(idx)
                    if index_to_update:
                        success = create_or_update_rag_index(idx)
                        if not success:
                            app.logger.error(f"Failed to initialize empty RAG index {index_to_update.name}")
            except Exception as e:
                app.logger.error(f"Error in background indexing thread during index creation: {e}", exc_info=True)

        thread = threading.Thread(target=background_indexing, args=(rag_index_id,))
        thread.daemon = True
        thread.start()

        app.logger.info(f"Started background initialization for new index {name} (ID: {rag_index_id})")
        flash(f'Successfully created RAG index "{name}" and associated model "{rag_model_display_name}". Background vector store initialization started.', 'success')

    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error creating RAG index: {e}", exc_info=True)
        flash('An unexpected error occurred while creating the index.', 'danger')

    return redirect(url_for('admin_rag_page'))


@app.route('/admin/rag/document/upload', methods=['POST'])
@login_required
def upload_rag_document():
    if not current_user.has_role('admin'):
        flash('You do not have permission to perform this action.', 'danger')
        return redirect(url_for('admin_rag_page'))

    if 'document' not in request.files:
        flash('No file part', 'danger')
        return redirect(url_for('admin_rag_page'))
    
    file = request.files['document']
    if file.filename == '':
        flash('No selected file', 'danger')
        return redirect(url_for('admin_rag_page'))

    if file:
        try:
            filename = secure_filename(file.filename)
            stored_filename = f"{uuid.uuid4().hex}_{filename}"
            rag_docs_dir = app.config['RAG_DOCUMENTS_DIR']
            os.makedirs(rag_docs_dir, exist_ok=True)
            filepath = os.path.join(rag_docs_dir, stored_filename)
            file.save(filepath)
            
            filesize = os.path.getsize(filepath)
            mime_type = file.mimetype
            
            new_doc = RagDocument(
                filename=filename,
                stored_filename=stored_filename,
                filepath=stored_filename,
                filesize=filesize,
                mime_type=mime_type,
                uploaded_by_id=current_user.id
            )
            db.session.add(new_doc)
            db.session.commit()
            flash(f'Document "{filename}" uploaded successfully.', 'success')
        except Exception as e:
            db.session.rollback()
            app.logger.error(f"Error uploading RAG document: {e}")
            flash('An unexpected error occurred while uploading the document.', 'danger')

    return redirect(url_for('admin_rag_page'))

@app.route('/admin/rag/index/<int:index_id>/documents', methods=['GET'])
@login_required
def get_index_documents(index_id):
    if not current_user.has_role('admin'):
        return jsonify({
            'success': False,
            'message': 'You do not have permission to perform this action.'
        }), 403
    
    # Find the RAG index
    index = RagIndex.query.get_or_404(index_id)
    
    # Get all documents associated with this index
    documents = [{
        'id': doc.id,
        'filename': doc.filename,
        'filesize': doc.filesize
    } for doc in index.documents]
    
    return jsonify({
        'success': True,
        'documents': documents
    })

@app.route('/admin/rag/index/<int:index_id>/update_documents', methods=['POST'])
@login_required
def update_index_documents(index_id):
    if not current_user.has_role('admin'):
        return jsonify({
            'success': False,
            'message': 'You do not have permission to perform this action.'
        }), 403
    
    # Find the RAG index
    index = RagIndex.query.get_or_404(index_id)
    
    # Get the document IDs from the request
    data = request.get_json()
    document_ids = data.get('document_ids', [])
    
    try:
        # Clear existing documents and add the selected ones
        index.documents = []
        
        if document_ids:
            documents = RagDocument.query.filter(RagDocument.id.in_(document_ids)).all()
            index.documents = documents
        
        db.session.commit()
        
        # Create a background thread to handle the indexing process
        # This prevents the HTTP request from timing out during long indexing operations
        def background_indexing(idx_id):
            try:
                with app.app_context():
                    # Re-fetch the index from the database in this thread
                    success = create_or_update_rag_index(idx_id)
                    if not success:
                        app.logger.error(f"Failed to create/update RAG index id={idx_id}")
            except Exception as e:
                app.logger.error(f"Error in background indexing thread: {e}")
                
        if document_ids:  # Only start indexing if there are documents to process
            thread = threading.Thread(target=background_indexing, args=(index_id,))
            thread.daemon = True
            thread.start()
            app.logger.info(f"Started background indexing for index {index.name} (ID: {index.id})")
        
        return jsonify({
            'success': True,
            'message': f'Successfully updated documents for index "{index.name}". Re-indexing has been started in the background.'
        })
        
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error updating index documents: {e}")
        return jsonify({
            'success': False,
            'message': 'An error occurred while updating the index documents'
        }), 500

@app.route('/admin/rbac')
@login_required
def admin_rbac_page():
    if not current_user.has_role('admin'):
        flash("You must be an administrator to access this page.", "danger")
        return redirect(url_for('index'))
    
    if not current_user.can_access_page('admin_rbac_page'):
        flash('You do not have permission to access this page.', 'danger')
        return redirect(url_for('index'))

    try:
        all_users = User.query.all()
        all_roles = Role.query.all()
        
        # Get all models - both base and RAG-enabled models
        # We need to include all models for admin to manage permissions
        all_models = Model.query.all()
        
        # Sort models to group base models first, then RAG models
        # and within those groups, sort alphabetically by display name
        all_models.sort(key=lambda m: (m.is_rag_model, m.display_name.lower()))
        
        # Ensure all RAG models are assigned to the admin role
        admin_role = Role.query.filter_by(name='admin').first()
        if admin_role:
            rag_models = Model.query.filter_by(is_rag_model=True).all()
            for rag_model in rag_models:
                if rag_model not in admin_role.models:
                    admin_role.models.append(rag_model)
                    app.logger.info(f"Added missing RAG model '{rag_model.display_name}' to admin role")
            if rag_models:
                db.session.commit()
                app.logger.info(f"Ensured all {len(rag_models)} RAG models are assigned to admin role")
                
        app.logger.info(f"Loaded {len(all_models)} models for RBAC page, including RAG models")
    except Exception as e:
        app.logger.error(f"Error fetching data for RBAC page: {e}", exc_info=True)
        flash("Error loading RBAC management page data.", "danger")
        return redirect(url_for('index'))
        
    # Define application pages/resources for access management display
    # In a more complex app, these might come from a config or be discovered
    # Prepare current page permissions for the template
    current_page_permissions = {}
    for role in all_roles:
        current_page_permissions[role.id] = {}
        for page in MANAGED_PAGE_ENDPOINTS:
            current_page_permissions[role.id][page['endpoint']] = role.has_page_access(page['endpoint'])

    current_model_permissions = {}
    for role in all_roles:
        current_model_permissions[role.id] = {}
        # Ensure all_models is a list of Model objects, not just names or IDs
        # The admin_rbac_page already queries all_models = Model.query.all()
        for model_obj in all_models: # Iterate through the fetched Model objects
            current_model_permissions[role.id][model_obj.id] = role.has_model_access(model_obj.id)
    
    return render_template('admin_rbac.html', 
                           users=all_users, 
                           roles=all_roles, 
                           models=all_models, 
                           app_pages=MANAGED_PAGE_ENDPOINTS, 
                           current_page_permissions=current_page_permissions,
                           current_model_permissions=current_model_permissions)


@app.route('/admin/permissions/page_access/update', methods=['POST'])
@login_required
def update_page_access_permissions():
    if not current_user.has_role('admin'):
        return jsonify({'success': False, 'message': 'You do not have permission to perform this action.'}), 403

    try:
        data = request.get_json()
        permissions_to_set = data.get('permissions', []) # Expected: [{'role_id': X, 'page_endpoint': 'Y'}, ...]

        # Clear existing permissions for all managed pages
        # This is a simple approach; a more complex app might do more granular updates.
        existing_managed_endpoints = [p['endpoint'] for p in MANAGED_PAGE_ENDPOINTS]
        PagePermission.query.filter(PagePermission.page_endpoint.in_(existing_managed_endpoints)).delete(synchronize_session=False)
        # Note: synchronize_session=False is used here. If issues arise, consider 'fetch' or individual deletes.

        # Add new permissions
        for perm_data in permissions_to_set:
            role_id = perm_data.get('role_id')
            page_endpoint = perm_data.get('page_endpoint')

            if not role_id or not page_endpoint:
                app.logger.warning(f"Skipping invalid permission data: {perm_data}")
                continue
            
            role = Role.query.get(role_id)
            if not role:
                app.logger.warning(f"Role ID {role_id} not found while updating page permissions.")
                continue
            
            # Ensure the page_endpoint is one of the known manageable endpoints
            if page_endpoint not in existing_managed_endpoints:
                app.logger.warning(f"Attempt to set permission for unmanaged page_endpoint '{page_endpoint}'. Skipping.")
                continue

            new_permission = PagePermission(role_id=role.id, page_endpoint=page_endpoint)
            db.session.add(new_permission)
        
        db.session.commit()
        app.logger.info(f"Page access permissions updated by user {current_user.id}.")
        return jsonify({'success': True, 'message': 'Page access permissions updated successfully.'})

    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error updating page access permissions: {e}", exc_info=True)
        return jsonify({'success': False, 'message': 'Failed to update page access permissions due to a server error.'}), 500


@app.route('/admin/permissions/model_access/update', methods=['POST'])
@login_required
def update_model_access_permissions():
    if not current_user.has_role('admin'):
        return jsonify({'success': False, 'message': 'You do not have permission to perform this action.'}), 403

    try:
        data = request.get_json()
        role_id = data.get('role_id')
        model_ids_to_assign = data.get('model_ids', [])  # Expected: [1, 2, 3]

        if role_id is None:
            app.logger.warning("Role ID not provided in request to update_model_access_permissions.")
            return jsonify({'success': False, 'message': 'Role ID is required.'}), 400

        role = Role.query.get(role_id)
        if not role:
            app.logger.warning(f"Role ID {role_id} not found while updating model access permissions.")
            return jsonify({'success': False, 'message': f'Role with ID {role_id} not found.'}), 404

        # Fetch valid Model objects based on provided IDs
        valid_models = []
        if model_ids_to_assign:
            valid_models = Model.query.filter(Model.id.in_(model_ids_to_assign)).all()
            
            # Log if some provided model IDs were not found, but proceed with valid ones
            assigned_model_ids = {model.id for model in valid_models}
            invalid_ids_provided = [mid for mid in model_ids_to_assign if mid not in assigned_model_ids]
            if invalid_ids_provided:
                app.logger.warning(f"Invalid or non-existent model IDs {invalid_ids_provided} provided for role '{role.name}' (ID: {role.id}). These will be ignored.")

        # Update the role's associated models
        # SQLAlchemy automatically handles the changes in the 'role_models' association table
        role.models = valid_models
        
        db.session.commit()
        app.logger.info(f"Model access permissions for role '{role.name}' (ID: {role.id}) updated by admin user {current_user.id}. Assigned model IDs: {[m.id for m in valid_models]}.")
        return jsonify({'success': True, 'message': f'Model access permissions for role \'{role.name}\' updated successfully.'})

    except Exception as e:
        db.session.rollback()
        # Try to get role_id from data if available for logging, otherwise use a placeholder
        requested_role_id = data.get('role_id', 'N/A') if isinstance(data, dict) else 'N/A'
        app.logger.error(f"Error updating model access permissions for role ID {requested_role_id}: {e}", exc_info=True)
        return jsonify({'success': False, 'message': 'Failed to update model access permissions due to a server error.'}), 500

@app.route('/admin/models/rag/delete', methods=['POST'])
@login_required
def delete_rag_model():
    """Delete a RAG model from the system."""
    if not current_user.has_role('admin'):
        return jsonify({'success': False, 'message': 'You do not have permission to perform this action.'}), 403
    
    try:
        data = request.get_json()
        model_id = data.get('model_id')
        
        if not model_id:
            return jsonify({'success': False, 'message': 'Model ID is required.'}), 400
            
        # Get the RAG model
        model = Model.query.get(model_id)
        if not model:
            return jsonify({'success': False, 'message': f'Model with ID {model_id} not found.'}), 404
        
        # Check if it's actually a RAG model
        if not model.is_rag_model:
            return jsonify({'success': False, 'message': f'Model {model.display_name} is not a RAG model.'}), 400
        
        # Store names for logging
        model_name = model.display_name
        
        # First remove role associations to prevent foreign key constraint errors
        # Get all roles that have this model
        role_model_assoc = db.session.query(role_models).filter(role_models.c.model_id == model.id).all()
        for role_id, _ in role_model_assoc:
            role = Role.query.get(role_id)
            if role and model in role.models:
                role.models.remove(model)
        
        # Get the associated RAG index if it exists
        rag_index = RagIndex.query.get(model.rag_index_id) if model.rag_index_id else None
        index_name = rag_index.name if rag_index else 'None'
        
        # Delete the model
        db.session.delete(model)
        db.session.commit()
        
        app.logger.info(f"Admin user {current_user.id} ({current_user.username}) deleted RAG model '{model_name}' (ID: {model_id})")
        return jsonify({'success': True, 'message': f'RAG model "{model_name}" has been deleted successfully.'})
        
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error deleting RAG model: {e}", exc_info=True)
        return jsonify({'success': False, 'message': 'Failed to delete RAG model due to a server error.'}), 500

@app.route('/admin/permissions/model_access/batch_update', methods=['POST'])
@login_required
def update_model_access_permissions_batch():
    """Batch update model access permissions for multiple roles at once."""
    if not current_user.has_role('admin'):
        return jsonify({'success': False, 'message': 'You do not have permission to perform this action.'}), 403

    try:
        data = request.get_json()
        if not data or not isinstance(data, dict):
            app.logger.warning("No valid data provided for batch model permission update.")
            return jsonify({'success': False, 'message': 'No valid data provided.'}), 400
        
        # Data format: { role_id: { model_ids: [1, 2, 3] }, ... }
        updated_roles = []
        update_stats = {'success': 0, 'failed': 0, 'skipped': 0}
        
        for role_id_str, role_data in data.items():
            try:
                # Convert role_id to integer (it comes as string key from JSON)
                role_id = int(role_id_str)
                role = Role.query.get(role_id)
                
                if not role:
                    app.logger.warning(f"Role ID {role_id} not found during batch update.")
                    update_stats['skipped'] += 1
                    continue
                
                model_ids = role_data.get('model_ids', [])
                if not model_ids or not isinstance(model_ids, list):
                    app.logger.warning(f"No model IDs provided for role ID {role_id} in batch update.")
                    update_stats['skipped'] += 1
                    continue
                
                # Fetch valid models
                valid_models = Model.query.filter(Model.id.in_(model_ids)).all()
                
                # Check for invalid model IDs
                assigned_model_ids = {model.id for model in valid_models}
                invalid_ids = [mid for mid in model_ids if mid not in assigned_model_ids]
                if invalid_ids:
                    app.logger.warning(f"Invalid model IDs {invalid_ids} for role '{role.name}' (ID: {role_id}) will be ignored.")
                
                # Update role's models
                role.models = valid_models
                updated_roles.append(role.name)
                update_stats['success'] += 1
                
                app.logger.info(f"Batch update: Model permissions for role '{role.name}' updated, assigned {len(valid_models)} models.")
                
            except Exception as role_error:
                app.logger.error(f"Error updating permissions for role ID {role_id_str}: {str(role_error)}")
                update_stats['failed'] += 1
                # Continue with other roles despite errors
        
        # Commit all changes at once
        db.session.commit()
        
        # Format success message
        if updated_roles:
            roles_message = "Updated model permissions for roles: " + ", ".join(updated_roles)
            app.logger.info(f"Batch update completed by user {current_user.id}: {update_stats}")
            return jsonify({
                'success': True, 
                'message': roles_message,
                'stats': update_stats
            })
        else:
            return jsonify({
                'success': True, 
                'message': 'No role permissions were updated.',
                'stats': update_stats
            })

    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error in batch update of model permissions: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'message': 'Failed to update permissions due to a server error.'}), 500



@app.route('/admin/user/assign_roles/<int:user_id>', methods=['POST'])
@login_required
def assign_user_roles(user_id):
    if not current_user.has_role('admin'):
        return jsonify({'success': False, 'message': 'You do not have permission to perform this action.'}), 403

    user = User.query.get_or_404(user_id)
    # Ensure admin cannot unwittlingly remove their own admin role if they are the only admin
    # or de-admin themselves via this mechanism if they are the user being edited.
    if user.id == current_user.id:  # Check if the user being edited is the current admin
        admin_role = Role.query.filter_by(name='admin').first()
        # If the 'admin' role exists and its ID is NOT in the list of roles to be assigned to self
        if admin_role and str(admin_role.id) not in request.form.getlist('role_ids[]'):
            # Check if there are other admins in the system
            other_admins = User.query.join(User.roles).filter(Role.name == 'admin', User.id != user.id).count()
            if other_admins == 0:
                return jsonify({'success': False, 'message': 'As the only administrator, you cannot remove your own admin role.'}), 400

    role_ids = request.form.getlist('role_ids[]') # Assuming role_ids are sent as a list
    
    # app.logger.info(f"Assigning roles to user {user_id}: {role_ids}") # For debugging

    # Clear existing roles
    user.roles.clear()

    # Add new roles
    if role_ids:
        for role_id_str in role_ids:
            try:
                role_id = int(role_id_str)
                role = Role.query.get(role_id)
                if role:
                    user.roles.append(role)
                else:
                    app.logger.warning(f"Role ID {role_id} not found while assigning roles to user {user_id}.")
            except ValueError:
                app.logger.warning(f"Invalid role ID {role_id_str} received for user {user_id}.")
                # Optionally, return an error here if strict validation is needed

    try:
        db.session.commit()
        # Fetch the updated roles to send back for potential UI update
        updated_role_names = [role.name for role in user.roles]
        app.logger.info(f"Successfully updated roles for user {user.id} to: {updated_role_names}")
        return jsonify({'success': True, 'message': 'User roles updated successfully.', 'user_id': user_id, 'new_roles': updated_role_names})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error updating roles for user {user_id}: {e}", exc_info=True)
        return jsonify({'success': False, 'message': 'Failed to update roles due to a server error.'}), 500


@app.route('/admin/roles/create', methods=['POST'])
@login_required
def create_role():
    if not current_user.has_role('admin'):
        flash("You do not have permission to perform this action.", "danger")
        return redirect(url_for('admin_rbac_page'))

    role_name = request.form.get('role_name')
    role_description = request.form.get('role_description')

    if not role_name:
        flash('Role name is required.', 'warning')
        return redirect(url_for('admin_rbac_page'))

    existing_role = Role.query.filter_by(name=role_name).first()
    if existing_role:
        flash(f'Role "{role_name}" already exists.', 'warning')
        return redirect(url_for('admin_rbac_page'))

    try:
        new_role = Role(name=role_name, description=role_description)
        db.session.add(new_role)
        db.session.commit()
        flash(f'Role "{role_name}" created successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error creating role '{role_name}': {e}", exc_info=True)
        flash('Error creating role. Please check logs.', 'danger')

    return redirect(url_for('admin_rbac_page'))


@app.route('/admin/role/edit/<int:role_id>', methods=['GET'])
@login_required
def edit_role_page(role_id):
    if not current_user.has_role('admin'):
        flash("You do not have permission to perform this action.", "danger")
        return redirect(url_for('admin_rbac_page'))
    role = Role.query.get_or_404(role_id)
    return render_template('edit_role.html', role=role)

@app.route('/admin/role/edit/<int:role_id>', methods=['POST'])
@login_required
def edit_role_submit(role_id):
    if not current_user.has_role('admin'):
        flash("You do not have permission to perform this action.", "danger")
        return redirect(url_for('admin_rbac_page'))

    role = Role.query.get_or_404(role_id)
    new_name = request.form.get('role_name')
    new_description = request.form.get('role_description')

    if not new_name:
        flash('Role name is required.', 'warning')
        return render_template('edit_role.html', role=role) # Stay on edit page with error

    # Check if new name conflicts with an existing role (excluding the current role itself)
    existing_role_with_new_name = Role.query.filter(Role.name == new_name, Role.id != role_id).first()
    if existing_role_with_new_name:
        flash(f'Role name "{new_name}" already exists. Please choose a different name.', 'warning')
        return render_template('edit_role.html', role=role) # Stay on edit page with error

    role.name = new_name
    role.description = new_description

    try:
        db.session.commit()
        flash(f'Role "{role.name}" updated successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error updating role '{role.name}': {e}", exc_info=True)
        flash('Error updating role. Please check logs.', 'danger')
    
    return redirect(url_for('admin_rbac_page'))


@app.route('/admin/users/create', methods=['POST'])
@login_required
def create_user_from_admin():
    if not current_user.has_role('admin'):
        flash("You do not have permission to perform this action.", "danger")
        return redirect(url_for('admin_rbac_page'))

    username = request.form.get('username')
    email = request.form.get('email')
    password = request.form.get('password')

    if not username or not email or not password:
        flash('Username, email, and password are required.', 'warning')
        return redirect(url_for('admin_rbac_page'))

    if User.query.filter_by(username=username).first():
        flash(f'Username "{username}" already exists.', 'warning')
        return redirect(url_for('admin_rbac_page'))
    
    if User.query.filter_by(email=email).first():
        flash(f'Email "{email}" already registered.', 'warning')
        return redirect(url_for('admin_rbac_page'))

    try:
        new_user = User(username=username, email=email)
        new_user.set_password(password)
        new_user.confirmed = True  # Admins create confirmed users
        # Optionally, assign a default role (e.g., 'user')
        # user_role = Role.query.filter_by(name='user').first()
        # if user_role:
        #     new_user.roles.append(user_role)
        
        db.session.add(new_user)
        db.session.commit()
        flash(f'User "{username}" created successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error creating user '{username}': {e}", exc_info=True)
        flash('Error creating user. Please check logs.', 'danger')

    return redirect(url_for('admin_rbac_page'))


@app.route('/admin/user/toggle_active/<int:user_id>', methods=['POST'])
@login_required
def toggle_user_active(user_id):
    app.logger.info(f"TOGGLE_USER_ACTIVE: current_user ID: {current_user.id}, Username: {current_user.username}, Email: {current_user.email}")
    app.logger.info(f"TOGGLE_USER_ACTIVE: current_user roles: {[role.name for role in current_user.roles]}")
    if not current_user.has_role('admin'):
        return jsonify({'success': False, 'message': 'You do not have permission to perform this action.', 'user_id': user_id, 'new_status': None}), 403

    user_to_toggle = User.query.get_or_404(user_id)

    # Prevent admins from deactivating themselves if they are the only admin
    if user_to_toggle.id == current_user.id and 'admin' in [role.name for role in user_to_toggle.roles]:
        other_admins = User.query.join(User.roles).filter(Role.name == 'admin', User.id != user_to_toggle.id).count()
        if other_admins == 0:
            return jsonify({'success': False, 'message': 'As the only administrator, you cannot deactivate your own account.'}), 400

    try:
        user_to_toggle.is_active = not user_to_toggle.is_active
        db.session.commit()
        status_text = "activated" if user_to_toggle.is_active else "deactivated"
        app.logger.info(f"User {user_to_toggle.username} has been {status_text}.")
        return jsonify({'success': True, 'message': f'User {user_to_toggle.username} has been {status_text}.', 'is_active': user_to_toggle.is_active})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error toggling active status for user {user_id}: {e}", exc_info=True)
        return jsonify({'success': False, 'message': 'An unexpected error occurred.', 'is_active': user_to_toggle.is_active if 'user_to_toggle' in locals() and hasattr(user_to_toggle, 'is_active') else None}), 500



@app.route('/admin/user/update', methods=['POST'])
@login_required
def update_user_details():
    app.logger.info(f"UPDATE_USER_DETAILS: current_user ID: {current_user.id}, Username: {current_user.username}")
    if not current_user.has_role('admin'):
        app.logger.warning(f"UPDATE_USER_DETAILS: Unauthorized attempt by user {current_user.id}")
        return jsonify({'success': False, 'error': 'Unauthorized access attempt.'}), 403

    data = request.get_json()
    if not data:
        app.logger.error("UPDATE_USER_DETAILS: No JSON data received.")
        return jsonify({'success': False, 'error': 'No data received.'}), 400

    user_id = data.get('user_id')
    new_username = data.get('username', '').strip()
    new_email = data.get('email', '').strip()

    if not user_id or not isinstance(user_id, int):
        app.logger.error(f"UPDATE_USER_DETAILS: Invalid or missing user_id: {user_id}")
        return jsonify({'success': False, 'error': 'Invalid or missing user ID.'}), 400
    
    if not new_username:
        app.logger.warning(f"UPDATE_USER_DETAILS: Username cannot be empty for user_id {user_id}.")
        return jsonify({'success': False, 'error': 'Username cannot be empty.'}), 400

    if not new_email:
        app.logger.warning(f"UPDATE_USER_DETAILS: Email cannot be empty for user_id {user_id}.")
        return jsonify({'success': False, 'error': 'Email cannot be empty.'}), 400

    user_to_update = User.query.get(user_id)
    if not user_to_update:
        app.logger.error(f"UPDATE_USER_DETAILS: User with ID {user_id} not found.")
        return jsonify({'success': False, 'error': 'User not found.'}), 404

    # Check for username conflict (if changed and new username exists for another user)
    if user_to_update.username != new_username:
        existing_user_by_username = User.query.filter(User.id != user_id, User.username == new_username).first()
        if existing_user_by_username:
            app.logger.warning(f"UPDATE_USER_DETAILS: Username '{new_username}' already taken by user ID {existing_user_by_username.id}.")
            return jsonify({'success': False, 'error': f'Username \"{new_username}\" is already taken.'}), 409

    # Check for email conflict (if changed and new email exists for another user)
    if user_to_update.email != new_email:
        existing_user_by_email = User.query.filter(User.id != user_id, User.email == new_email).first()
        if existing_user_by_email:
            app.logger.warning(f"UPDATE_USER_DETAILS: Email '{new_email}' already registered to user ID {existing_user_by_email.id}.")
            return jsonify({'success': False, 'error': f'Email \"{new_email}\" is already registered.'}), 409

    try:
        user_to_update.username = new_username
        user_to_update.email = new_email
        db.session.commit()
        app.logger.info(f"UPDATE_USER_DETAILS: User {user_id} ({new_username}) updated successfully.")
        return jsonify({'success': True, 'message': 'User details updated successfully.'})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"UPDATE_USER_DETAILS: Error updating user {user_id}: {e}", exc_info=True)
        return jsonify({'success': False, 'error': 'An internal error occurred while updating the user.'}), 500



def init_db():
    """Initialize the database with Flask-Migrate"""
    with app.app_context():
        # This will create the database tables if they don't exist
        # and apply any pending migrations
        from flask_migrate import upgrade
        
        # db.create_all() is removed because Flask-Migrate now handles schema creation.
        # The 'flask db upgrade' command in the startup script will create/update tables.
        
        # Apply any pending migrations
        # Initialize RBAC data
        initialize_rbac_data()
        load_and_ensure_llm_models()
        # If you need to run any data migrations, you can add them here
        # For example:
        # migrate_data()


@app.cli.command('init-db')
def init_db_command():
    """Initializes the database."""
    init_db()


def check_database_connection():
    """Check if we can connect to the database"""
    max_retries = 10
    retry_count = 0
    
    while retry_count < max_retries:
        try:
            # Try to execute a simple query
            db.session.execute('SELECT 1')
            app.logger.info("Successfully connected to the database")
            return True
        except Exception as e:
            retry_count += 1
            app.logger.warning(f"Database connection failed (attempt {retry_count}/{max_retries}): {str(e)}")
            if retry_count >= max_retries:
                app.logger.error("Max retries reached. Could not connect to the database.")
                return False
            time.sleep(5)  # Wait before retrying


if __name__ == '__main__':
    # Check dependencies
    system_deps = check_system_dependencies()
    if not system_deps['ffmpeg']:
        app.logger.error("FFmpeg is not installed. Please install FFmpeg to use this application.")
        # Don't exit in Docker, let it try to run
        # exit(1) 
    if not system_deps['whisper']:
        app.logger.error("Whisper models are not available. Please install faster-whisper to use this application.")
        # Don't exit in Docker, let it try to run
        # exit(1)
    
    app.logger.info(f"Starting AI Chat application with {llm_service_type.upper()} as the LLM service")
    
    # Wait for database to be ready
    if not check_database_connection():
        app.logger.error("Failed to connect to the database. Exiting...")
        exit(1)
    
    # Initialize database and apply migrations
    try:
        app.logger.info("Initializing database...")
        init_db()
        app.logger.info("Database initialization completed successfully")
    except Exception as e:
        app.logger.error(f"Failed to initialize database: {str(e)}")
        exit(1)
    